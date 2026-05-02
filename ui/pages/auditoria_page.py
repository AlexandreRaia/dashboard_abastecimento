"""
Página de Auditoria de Frota.

Fonte de dados: relatorio.db (READ-ONLY) — tabela abastecimentos.
Resultados do pipeline gravados em: auditoria_resultados.db.

Estrutura:
- Sidebar: filtros (data, combustível, secretaria, marca, modelos,
           motorista, placa, sigma) + botão Aplicar Auditoria
- Main: stats + gráfico de dispersão (sempre visíveis)
- Tabs: Visão Geral | Ocorrências | Notificações | Exportar | Manual
"""
from __future__ import annotations

import base64
import json
import smtplib
import ssl
from datetime import date, datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from io import BytesIO
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st

from agents import OrchestradorAuditoria
from agents.config import (
    MODEL_CANONICAL_DISPLAY,
    TANK_CAPACITY,
    THRESHOLDS as _CFG_THRESHOLDS,
    normalizar_texto as _normalizar_modelo,
)
from config import settings
from core.utils.formatters import currency
from infrastructure.repositories.abastecimento_repo import load_abastecimentos

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------
MAX_RENDER_ROWS = 5_000
_CONSUMO_MAX = float(_CFG_THRESHOLDS.get("consumo_max_valido", 30.0))
_ASSETS_DIR = Path(__file__).resolve().parent.parent.parent / "assets"
_MANUAL_HTML = _ASSETS_DIR / "manual_auditoria.html"
_MODELO_RELATORIO = Path(__file__).resolve().parent.parent.parent / "Modelo_relatorio"

_FUEL_DISPLAY = {"Gasolina": "GASOLINA", "Álcool": "ALCOOL", "Diesel S10": "DIESEL S10"}

_CONFIG_JSON = Path(__file__).resolve().parent.parent.parent / "config.json"


def _load_smtp_cfg() -> dict:
    """Lê configurações SMTP salvas em config.json (secção 'smtp')."""
    try:
        data = json.loads(_CONFIG_JSON.read_text(encoding="utf-8"))
        cfg  = data.get("smtp", {})
        if cfg.get("senha_b64"):
            cfg["senha"] = base64.b64decode(cfg["senha_b64"]).decode()
        cfg.pop("senha_b64", None)
        return cfg
    except Exception:
        return {}


def _lookup_contato(condutor: str, unidade: str) -> str:
    """Retorna e-mail cadastrado para o condutor ou, se não houver, para a secretaria."""
    try:
        data = json.loads(_CONFIG_JSON.read_text(encoding="utf-8"))
        contatos = data.get("contatos", {})
        email = contatos.get("condutores", {}).get(condutor, "")
        if not email:
            email = contatos.get("secretarias", {}).get(unidade, "")
        return email or ""
    except Exception:
        return ""


def _save_smtp_cfg(host: str, port: int, remetente: str,
                   senha: str, usar_ssl: bool) -> None:
    """Persiste configurações SMTP em config.json."""
    try:
        data = json.loads(_CONFIG_JSON.read_text(encoding="utf-8"))
        data["smtp"] = {
            "host":      host.strip(),
            "port":      port,
            "remetente": remetente.strip(),
            "senha_b64": base64.b64encode(senha.replace(" ", "").strip().encode()).decode(),
            "usar_ssl":  usar_ssl,
        }
        _CONFIG_JSON.write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass  # falha silenciosa — envio ainda funciona


def _save_contato(tipo: str, chave: str, email: str) -> None:
    """Salva/atualiza e-mail de contato (secretaria ou condutor) em config.json."""
    try:
        data = json.loads(_CONFIG_JSON.read_text(encoding="utf-8"))
        if "contatos" not in data:
            data["contatos"] = {"secretarias": {}, "condutores": {}}
        section = "secretarias" if tipo == "secretaria" else "condutores"
        data["contatos"].setdefault(section, {})[chave.strip()] = email.strip()
        _CONFIG_JSON.write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass

# ---------------------------------------------------------------------------
# Cache — leitura do banco via repositório normalizado
# ---------------------------------------------------------------------------

@st.cache_data(show_spinner="Carregando banco de abastecimentos…", ttl=300)
def _load_all(_db_path: str) -> pd.DataFrame:
    """Carrega todos os abastecimentos via repositório — dados já normalizados."""
    return load_abastecimentos(Path(_db_path))


def _date_bounds(df: pd.DataFrame) -> tuple[date, date]:
    dates = df["data_hora"].dropna()
    return dates.min().date(), dates.max().date()

# ---------------------------------------------------------------------------
# Helpers CSS
# ---------------------------------------------------------------------------

def _apply_style() -> None:
    st.markdown(
        """
        <style>
        .block-container { padding-top: 0.35rem !important; padding-bottom: 1.2rem !important; }
        .painel-executivo {
            background: linear-gradient(135deg,#111d2b 0%,#0f1826 100%);
            border:1px solid rgba(142,163,190,0.18);border-radius:22px;
            padding:1.2rem 1.25rem;margin-bottom:1rem;
        }
        .painel-grid {
            display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));
            gap:.85rem;margin-top:.85rem;
        }
        .painel-card {
            background:linear-gradient(150deg,#162436 0%,#111d2b 100%);
            border:1px solid rgba(142,163,190,0.22);
            border-radius:18px;padding:1rem;box-shadow:0 6px 16px rgba(0,0,0,0.3);
        }
        .painel-card.alerta {
            background:linear-gradient(150deg,#2d1a1e 0%,#1e1015 100%);
            border-color:rgba(239,68,68,0.35);
        }
        .painel-card .rotulo {
            display:block;color:#8ea3be;font-size:.82rem;margin-bottom:.45rem;
            text-transform:uppercase;letter-spacing:.04em;
        }
        .painel-card .valor {
            display:block;color:#e7eef8;font-size:1.9rem;font-weight:700;line-height:1.05;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

# ---------------------------------------------------------------------------
# Funções de renderização
# ---------------------------------------------------------------------------

def _render_stats_and_scatter(
    df_filt: pd.DataFrame,
    sigma_mult: float,
    modelos_selecionados: list[str] | None = None,
    period: tuple | None = None,
) -> None:
    """KPIs de consumo + gráfico de dispersão com outliers em vermelho."""
    consumo = df_filt["km_por_litro"]
    consumo_ok = consumo[(consumo > 0) & (consumo <= _CONSUMO_MAX)].dropna()

    if consumo_ok.empty:
        st.warning("Sem dados de consumo válidos para o período/filtros selecionados.")
        return

    media  = consumo_ok.mean()
    desvio = consumo_ok.std()
    cmax   = consumo_ok.max()
    cmin   = consumo_ok.min()
    limiar = media - sigma_mult * desvio

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Máx KM/L",      f"{cmax:.2f}")
    c2.metric("Mín KM/L",      f"{cmin:.2f}")
    c3.metric("Média KM/L",    f"{media:.2f}")
    c4.metric("Desvio Padrão", f"{desvio:.2f}")

    # Prepara df para o scatter
    df_plot = df_filt[
        (df_filt["km_por_litro"] > 0) & (df_filt["km_por_litro"] <= _CONSUMO_MAX)
    ].copy()
    excluidos = int((df_filt["km_por_litro"] > _CONSUMO_MAX).sum())
    if excluidos:
        st.caption(
            f"⚠️ {excluidos} registro(s) com km/L > {_CONSUMO_MAX:.0f} excluídos do gráfico "
            "(hodômetro inválido)."
        )

    df_plot["outlier"] = df_plot["km_por_litro"] < limiar
    df_plot["data_str"] = df_plot["data_hora"].dt.strftime("%d/%m/%Y %H:%M")

    fig = px.scatter(
        df_plot,
        x="km_por_litro",
        y="litros",
        color="outlier",
        color_discrete_map={True: "#ef4444", False: "#38bdf8"},
        hover_data={
            "placa": True,
            "modelo": True,
            "condutor": True,
            "km_por_litro": ":.2f",
            "litros": ":.1f",
            "data_str": True,
            "outlier": False,
        },
        labels={
            "km_por_litro": "KM/L",
            "litros": "Litros abastecidos",
            "data_str": "Data/Hora",
            "outlier": "Outlier",
        },
        title=(
            f"Dispersão KM/L × Litros abastecidos  "
            f"({len(df_plot):,} registros"
            + (
                f" · {period[0].strftime('%d/%m/%Y')} – {period[1].strftime('%d/%m/%Y')}"
                if period else ""
            )
            + ")"
        ),
        template="plotly_dark",
    )
    # Linhas horizontais de capacidade máxima do tanque por modelo
    for k, cap in TANK_CAPACITY.items():
        if not modelos_selecionados or any(
            k in _normalizar_modelo(m) for m in modelos_selecionados
        ):
            display = MODEL_CANONICAL_DISPLAY.get(k, k.capitalize())
            fig.add_hline(
                y=cap,
                line_dash="dot",
                line_color="#facc15",
                annotation_text=f"{display} {cap}L",
                annotation_position="bottom right",
                annotation_font_color="#facc15",
            )
    fig.update_layout(
        xaxis_title="KM/L",
        yaxis_title="Litros",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(15,24,38,0.6)",
        margin=dict(l=8, r=8, t=40, b=8),
        legend_title_text="",
    )
    fig.for_each_trace(lambda t: t.update(name="Outlier" if t.name == "True" else "Normal"))
    st.plotly_chart(fig, use_container_width=True)


def _render_painel_executivo(resultado: dict) -> None:
    ocs  = resultado.get("ocorrencias", [])
    n    = len(ocs)
    tot  = len(resultado.get("df_auditado", [])) or n
    grav = sum(1 for o in ocs if isinstance(o, dict) and o.get("gravidade_final") == "ALTA")
    st.markdown(
        f"""
        <div class="painel-executivo">
            <b style="font-size:1.2rem;">📋 Painel Executivo — Agentes</b>
            <div class="painel-grid">
                <div class="painel-card">
                    <span class="rotulo">Registros auditados</span>
                    <span class="valor">{tot:,}</span>
                </div>
                <div class="painel-card alerta">
                    <span class="rotulo">Ocorrências totais</span>
                    <span class="valor">{n}</span>
                </div>
                <div class="painel-card alerta">
                    <span class="rotulo">Ocorrências graves</span>
                    <span class="valor">{grav}</span>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Alertas de qualidade de dados
    falhas = [
        f for f in resultado.get("relatorio_qualidade", [])
        if f.get("tipo") != "VALIDACAO_CONCLUIDA"
        and f.get("gravidade") in ("MEDIA", "ALTA")
    ]
    if falhas:
        with st.expander(f"⚠️ {len(falhas)} problema(s) de qualidade detectado(s)", expanded=True):
            for f in falhas:
                icon = "🔴" if f.get("gravidade") == "ALTA" else "🟡"
                st.markdown(f"{icon} **{f.get('tipo','').replace('_',' ')}** — {f.get('detalhe','')}")

    with st.expander("📋 Log do pipeline", expanded=False):
        for linha in resultado.get("log", []):
            st.code(linha, language=None)


from infrastructure.repositories.resolucoes_repo import (
    get_resolucoes, salvar_resolucao, remover_resolucao,
)

_AUDIT_DB = Path(__file__).resolve().parent.parent.parent / "auditoria_resultados.db"

_GRAV_ORDER = {"ALTA": 2, "MEDIA": 1, "BAIXA": 0}
_GRAV_BADGE = {"ALTA": "🔴 ALTA", "MEDIA": "🟡 MÉDIA", "BAIXA": "🟢 BAIXA"}

def _grav_badge(val: str) -> str:
    return _GRAV_BADGE.get(str(val).upper(), val)


def _render_ocorrencias(resultado: dict) -> None:
    ocs = resultado.get("ocorrencias", resultado.get("alertas", []))
    if not ocs:
        st.success("✅ Nenhuma ocorrência detectada.")
        return

    # Carrega resoluções gravadas
    resolucoes = get_resolucoes(_AUDIT_DB)

    df_oc = pd.DataFrame(ocs) if isinstance(ocs, list) else ocs.copy()
    col_grav = next(
        (c for c in ("gravidade_final", "gravidade_inicial", "gravidade") if c in df_oc.columns),
        None,
    )

    # Marca status de cada ocorrência
    has_id = "id_ocorrencia" in df_oc.columns
    if has_id:
        df_oc["_status"] = df_oc["id_ocorrencia"].map(
            lambda oid: resolucoes.get(oid, {}).get("status", "PENDENTE")
        )
    else:
        df_oc["_status"] = "PENDENTE"

    # ── Filtros ──────────────────────────────────────────────────────────────
    f1, f2 = st.columns([2, 2])
    with f1:
        if col_grav:
            opts_g = ["TODAS"] + sorted(df_oc[col_grav].dropna().unique().tolist())
            sel_g  = st.radio("Gravidade", opts_g, horizontal=True, key="aud_grav_filter")
        else:
            sel_g = "TODAS"
    with f2:
        sel_s = st.radio("Status", ["PENDENTES", "RESOLVIDAS", "TODAS"],
                         horizontal=True, key="aud_status_filter")

    df_view = df_oc.copy()
    if sel_g != "TODAS" and col_grav:
        df_view = df_view[df_view[col_grav] == sel_g]
    if sel_s == "PENDENTES":
        df_view = df_view[df_view["_status"] == "PENDENTE"]
    elif sel_s == "RESOLVIDAS":
        df_view = df_view[df_view["_status"] != "PENDENTE"]

    n_pend = int((df_oc["_status"] == "PENDENTE").sum())
    n_res  = int((df_oc["_status"] != "PENDENTE").sum())
    c1, c2, c3 = st.columns(3)
    c1.metric("Total", len(df_oc))
    c2.metric("🔴 Pendentes", n_pend)
    c3.metric("✅ Resolvidas", n_res)

    st.divider()

    # ── Tabela ───────────────────────────────────────────────────────────────
    _COLS_ORDER = [
        "data_hora", "placa", "condutor", "modelo", "unidade",
        "codigo_regra", "gravidade_final",
        "km_anterior", "km_atual", "km_rodados", "km_esperado",
        "litros", "km_l", "media", "desvio",
        "estabelecimento", "produto", "descricao_tecnica",
    ]
    _COLS_CFG = {
        "data_hora":        st.column_config.DatetimeColumn("Data/Hora",    format="DD/MM/YYYY HH:mm"),
        "placa":            st.column_config.TextColumn("Placa"),
        "condutor":         st.column_config.TextColumn("Condutor"),
        "modelo":           st.column_config.TextColumn("Modelo"),
        "unidade":          st.column_config.TextColumn("Secretaria"),
        "codigo_regra":     st.column_config.TextColumn("Regra"),
        "gravidade_final":  st.column_config.TextColumn("Gravidade"),
        "km_anterior":      st.column_config.NumberColumn("KM Anterior",   format="%.0f"),
        "km_atual":         st.column_config.NumberColumn("KM Atual",      format="%.0f"),
        "km_rodados":       st.column_config.NumberColumn("KM Rodados",    format="%.1f"),
        "km_esperado":      st.column_config.NumberColumn("KM Esperado",   format="%.1f"),
        "litros":           st.column_config.NumberColumn("Litros",        format="%.2f L"),
        "km_l":             st.column_config.NumberColumn("KM/L",          format="%.2f"),
        "media":            st.column_config.NumberColumn("Média KM/L",    format="%.2f"),
        "desvio":           st.column_config.NumberColumn("Desvio KM/L",   format="%.2f"),
        "estabelecimento":  st.column_config.TextColumn("Posto"),
        "produto":          st.column_config.TextColumn("Combustível"),
        "descricao_tecnica": st.column_config.TextColumn("Descrição"),
    }

    cols_pres = [c for c in _COLS_ORDER if c in df_view.columns]
    cfg_pres  = {k: v for k, v in _COLS_CFG.items() if k in cols_pres}
    df_show   = df_view[cols_pres].copy()

    if col_grav and col_grav in df_show.columns:
        df_show = df_show.sort_values(
            col_grav,
            key=lambda s: s.map(lambda g: _GRAV_ORDER.get(g, -1)),
            ascending=False,
        )
        df_show[col_grav] = df_show[col_grav].map(_grav_badge)

    n_show = min(len(df_show), MAX_RENDER_ROWS)
    if len(df_show) > MAX_RENDER_ROWS:
        st.caption(f"Exibindo os primeiros {MAX_RENDER_ROWS} de {len(df_show)}.")
    st.dataframe(df_show.head(n_show).reset_index(drop=True),
                 use_container_width=True, hide_index=True, column_config=cfg_pres)

    # ── Painel de resolução ──────────────────────────────────────────────────
    if not has_id:
        return

    st.divider()
    st.markdown("#### ✏️ Registrar resolução")

    pendentes = df_oc[df_oc["_status"] == "PENDENTE"]
    if pendentes.empty:
        st.success("Todas as ocorrências já foram resolvidas.")
        return

    # Monta opções: "REGRA — Placa — Data"
    def _label(row):
        dt = row.get("data_hora")
        dt_str = dt.strftime("%d/%m/%Y %H:%M") if hasattr(dt, "strftime") else str(dt)
        return f"[{row.get('codigo_regra','?')}] {row.get('placa','?')} — {dt_str} — {row.get('condutor','?')}"

    opcoes = {_label(row): row["id_ocorrencia"]
              for _, row in pendentes.iterrows()}

    sel_label = st.selectbox("Selecionar ocorrência pendente",
                             list(opcoes.keys()), key="aud_oc_sel")
    id_sel = opcoes[sel_label]

    col_status, col_obs = st.columns([1, 3])
    with col_status:
        novo_status = st.radio("Resolução", ["JUSTIFICADA", "DESCARTADA"],
                               key="aud_res_status")
    with col_obs:
        obs = st.text_area("Observação (opcional)", height=80, key="aud_res_obs",
                           placeholder="Descreva a justificativa ou motivo do descarte…")

    if st.button("💾 Gravar resolução", type="primary", key="aud_res_salvar"):
        salvar_resolucao(_AUDIT_DB, id_sel, novo_status, obs)
        st.success(f"Ocorrência marcada como **{novo_status}**.")
        st.rerun()

    # Seção para re-abrir ocorrências já resolvidas
    resolvidas = df_oc[df_oc["_status"] != "PENDENTE"]
    if not resolvidas.empty:
        with st.expander(f"🔄 Re-abrir ocorrência resolvida ({len(resolvidas)})"):
            opcoes_res = {_label(row) + f" [{row['_status']}]": row["id_ocorrencia"]
                          for _, row in resolvidas.iterrows()}
            sel_res = st.selectbox("Ocorrência", list(opcoes_res.keys()),
                                   key="aud_reopen_sel")
            if st.button("↩️ Re-abrir como pendente", key="aud_reopen_btn"):
                remover_resolucao(_AUDIT_DB, opcoes_res[sel_res])
                st.success("Ocorrência re-aberta como pendente.")
                st.rerun()


def _render_alertas_operacionais(df: pd.DataFrame) -> None:
    """Alertas operacionais: múltiplos abastecimentos no mesmo dia / intervalo < 1h."""
    if df.empty or "data_hora" not in df.columns:
        st.info("Sem dados para verificar alertas operacionais.")
        return

    df_t = df.copy()
    df_t["_data_dia"] = df_t["data_hora"].dt.date

    # ── 1. Múltiplos abastecimentos no mesmo dia ──────────────────────────────
    same_day = (
        df_t.groupby(["placa", "_data_dia"])
        .filter(lambda g: len(g) > 1)
        .sort_values(["placa", "data_hora"])
    )

    # ── 2. Intervalo < 1 hora entre abastecimentos da mesma placa ─────────────
    df_s = df_t.sort_values(["placa", "data_hora"])
    df_s["_diff_h"] = (
        df_s.groupby("placa")["data_hora"]
        .diff()
        .dt.total_seconds()
        .div(3600)
    )
    less_1h = df_s[df_s["_diff_h"].notna() & (df_s["_diff_h"] < 1)].copy()

    # Colunas comuns de exibição
    _COLS = [
        c for c in [
            "data_hora", "placa", "condutor", "modelo",
            "secretaria", "combustivel", "litros", "km_por_litro", "valor",
        ] if c in df.columns
    ]
    _CFG = {
        "data_hora":    st.column_config.DatetimeColumn("Data/Hora",  format="DD/MM/YYYY HH:mm"),
        "litros":       st.column_config.NumberColumn("Litros",       format="%.2f L"),
        "km_por_litro": st.column_config.NumberColumn("KM/L",         format="%.2f"),
        "valor":        st.column_config.NumberColumn("Valor R$",      format="R$ %.2f"),
    }
    _cfg = {k: v for k, v in _CFG.items() if k in df.columns}

    col_a, col_b = st.columns(2)
    col_a.metric("Veículos — múlt. abast. mesmo dia", f"{same_day['placa'].nunique()}")
    col_b.metric("Abastecimentos com intervalo < 1h", f"{len(less_1h)}")

    st.divider()

    st.markdown("#### 📅 Múltiplos abastecimentos no mesmo dia")
    if same_day.empty:
        st.success("Nenhum abastecimento duplicado no mesmo dia encontrado.")
    else:
        st.dataframe(
            same_day[_COLS], use_container_width=True,
            column_config=_cfg, hide_index=True,
        )

    st.markdown("#### ⏱️ Abastecimentos com intervalo inferior a 1 hora")
    if less_1h.empty:
        st.success("Nenhum abastecimento com intervalo inferior a 1 hora.")
    else:
        _COLS_1H = [c for c in _COLS if c != "km_por_litro"] + ["_diff_h"]
        _COLS_1H = [c for c in _COLS_1H if c in less_1h.columns]
        _cfg_1h  = {**_cfg, "_diff_h": st.column_config.NumberColumn("Intervalo (h)", format="%.2f h")}
        st.dataframe(
            less_1h[_COLS_1H], use_container_width=True,
            column_config=_cfg_1h, hide_index=True,
        )


def _gerar_docx_notif(notif: dict) -> bytes:
    """
    Gera .docx seguindo o Modelo_relatorio.

    Estrutura do modelo (por índice de parágrafo):
      0  – Título (CENTER, 12pt)
      1  – vazio (CENTER)
      2  – Data (RIGHT)   → dinâmico
      3-4 – vazios (RIGHT)
      5  – "Prezado(a) NOME" (JUSTIFY, bold) → dinâmico
      6  – SECRETARIA (JUSTIFY, bold)        → dinâmico
      7  – vazio
      8  – Corpo / introdução (JUSTIFY)      → dinâmico
      9  – "OCORRENCIAS IDENTIFICADAS:" (JUSTIFY)
      10 – primeira ocorrência (JUSTIFY)     → dinâmico
      11 – segunda ocorrência (JUSTIFY)      → dinâmico / duplicar se necessário
      12 – "Diante do exposto…"
      13 – "Ressaltamos…"
      14 – vazio
      15 – "Atenciosamente,"
      16 – vazio
      17 – assinatura 1 (bold)
      18 – cargo 1 (CENTER, bold)
      19-22 – vazios
      23 – assinatura 2 (CENTER, bold)
      24 – cargo 2 (CENTER, bold)
    """
    import datetime as _dt
    from copy import deepcopy
    from docx import Document
    from docx.oxml.ns import qn

    _MESES_PT = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
    ]

    def _set_para(para, text: str, bold: bool | None = None) -> None:
        """Substitui o texto do parágrafo preservando a formatação do primeiro run."""
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
            if bold is not None:
                para.runs[0].bold = bold
        else:
            r = para.add_run(text)
            if bold is not None:
                r.bold = bold

    def _clone_para_after(ref_elem, parent, text: str) -> None:
        """Insere uma cópia do elemento ref_elem logo após si mesmo, com o texto dado."""
        new_elem = deepcopy(ref_elem)
        for r in new_elem.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                t.text = text
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            break
        idx = list(parent).index(ref_elem)
        parent.insert(idx + 1, new_elem)
    # ── Abre o modelo a partir dos bytes (não altera o arquivo original) ────
    template_bytes = _MODELO_RELATORIO.read_bytes()
    doc = Document(BytesIO(template_bytes))
    paras = doc.paragraphs

    # ── Data ────────────────────────────────────────────────────────────────
    today = _dt.date.today()
    date_str = (
        f"Santana de Parnaíba,  {today.day} de "
        f"{_MESES_PT[today.month - 1]} de {today.year}."
    )
    _set_para(paras[2], date_str)

    # ── Destinatário ─────────────────────────────────────────────────────────
    _set_para(paras[5], f"Prezado(a) {notif.get('condutor', '').upper()}", bold=True)
    _set_para(paras[6], notif.get("unidade", "").upper(), bold=True)

    # ── Corpo introdutório ───────────────────────────────────────────────────
    modelo_txt = f" ({notif['modelo']})" if notif.get("modelo") else ""
    intro = (
        f"Foram identificadas inconsistencias tecnicas nos registros de "
        f"abastecimento vinculados ao veiculo {notif.get('placa', '?')}{modelo_txt}, "
        f"realizado em {notif.get('data_abastecimento', '')} "
        f"no estabelecimento {notif.get('estabelecimento', '') or 'nao informado'}."
    )
    _set_para(paras[8], intro)

    # ── Ocorrências ──────────────────────────────────────────────────────────
    ocorrencias = notif.get("ocorrencias", [])

    def _oc_texto(oc: dict) -> str:
        return (
            f"  [{oc.get('codigo_regra', '?')}] "
            f"{oc.get('tipo_ocorrencia', '').title()}: "
            f"{oc.get('descricao_tecnica', '')}"
        )

    if not ocorrencias:
        # Remove as duas linhas de ocorrência do template
        for _ in range(2):
            p = doc.paragraphs[10]
            p._p.getparent().remove(p._p)
    elif len(ocorrencias) == 1:
        _set_para(doc.paragraphs[10], _oc_texto(ocorrencias[0]))
        # Remove segunda linha placeholder
        p11 = doc.paragraphs[11]
        p11._p.getparent().remove(p11._p)
    else:
        _set_para(doc.paragraphs[10], _oc_texto(ocorrencias[0]))
        _set_para(doc.paragraphs[11], _oc_texto(ocorrencias[1]))
        # Insere linhas extras após parágrafo 11 (índice dinâmico após edições)
        if len(ocorrencias) > 2:
            parent = doc.paragraphs[11]._p.getparent()
            # `ref` começa no para 11 e avança para o último inserido a cada iteração
            ref = doc.paragraphs[11]._p
            for oc in ocorrencias[2:]:
                _clone_para_after(ref, parent, _oc_texto(oc))
                # O novo elemento foi inserido em idx+1; ref avança para ele
                ref = list(parent)[list(parent).index(ref) + 1]

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _render_notificacoes(resultado: dict) -> None:
    notifs = resultado.get("notificacoes", [])
    if not notifs:
        st.info("Nenhuma notificação gerada.")
        return

    st.markdown(f"**{len(notifs)} notificação(ões) gerada(s)**")

    for i, notif in enumerate(notifs):
        placa     = notif.get("placa", f"#{i+1}")
        grav      = notif.get("gravidade_max", "")
        condutor  = notif.get("condutor", "")
        unidade   = notif.get("unidade", "")
        texto     = notif.get("texto_notificacao", notif.get("minuta", ""))
        label_extra = f" · {condutor}" if condutor else ""
        label_extra += f" · {unidade}" if unidade else ""
        label = f"📄 Placa **{placa}**{label_extra} — {_grav_badge(grav)}"

        with st.expander(label, expanded=(i == 0)):
            st.text_area(
                "Minuta",
                value=texto,
                height=280,
                key=f"aud_minuta_{i}",
                label_visibility="collapsed",
            )

            col_word, col_email = st.columns(2)

            # ── Word ─────────────────────────────────────────────────────────
            with col_word:
                docx_bytes = _gerar_docx_notif(notif)
                st.download_button(
                    "📝 Exportar Word",
                    data=docx_bytes,
                    file_name=f"notificacao_{placa.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"aud_word_{i}",
                )

            # ── E-mail ────────────────────────────────────────────────────────
            with col_email:
                with st.popover("📧 Enviar por e-mail", use_container_width=True):
                    _smtp = _load_smtp_cfg()
                    host     = st.text_input("Servidor SMTP", value=_smtp.get("host", "smtp.gmail.com"), key=f"smtp_host_{i}")
                    port     = st.number_input("Porta", value=int(_smtp.get("port", 465)), step=1, key=f"smtp_port_{i}")
                    rem      = st.text_input("Remetente", value=_smtp.get("remetente", ""), key=f"smtp_rem_{i}")
                    senha    = st.text_input("Senha / App Password", value=_smtp.get("senha", ""), type="password", key=f"smtp_pwd_{i}")
                    dest     = st.text_input(
                        "Destinatário",
                        value=_lookup_contato(notif.get("condutor", ""), notif.get("unidade", "")),
                        key=f"smtp_dest_{i}",
                    )
                    usar_ssl = st.checkbox("Usar SSL", value=bool(_smtp.get("usar_ssl", True)), key=f"smtp_ssl_{i}")
                    col_salvar, col_enviar = st.columns(2)
                    with col_salvar:
                        if st.button("💾 Salvar cfg", key=f"btn_save_smtp_{i}", use_container_width=True):
                            _save_smtp_cfg(host, int(port), rem, senha, usar_ssl)
                            st.success("Configuração salva!")
                    with col_enviar:
                        if st.button("Enviar", key=f"btn_email_{i}", type="primary", use_container_width=True):
                            ok, err = _enviar_email_smtp(notif, {
                                "host": host, "port": port, "remetente": rem,
                                "senha": senha, "destinatario": dest, "usar_ssl": usar_ssl,
                            })
                            st.success("E-mail enviado!") if ok else st.error(f"Falha: {err}")

            # ── Resolver ocorrências desta notificação ────────────────────────
            ocs_notif = notif.get("ocorrencias", [])
            if ocs_notif:
                resolucoes = get_resolucoes(_AUDIT_DB)
                st.divider()
                st.markdown("**✏️ Registrar resolução das ocorrências**")
                for j, oc in enumerate(ocs_notif):
                    oid = oc.get("id_ocorrencia")
                    if not oid:
                        continue
                    status_atual = resolucoes.get(oid, {}).get("status", "PENDENTE")
                    regra = oc.get("codigo_regra", "?")
                    dt    = oc.get("data_hora", "")
                    dt_str = dt.strftime("%d/%m/%Y %H:%M") if hasattr(dt, "strftime") else str(dt)

                    badge = {"JUSTIFICADA": "🟢 JUSTIFICADA", "DESCARTADA": "⚫ DESCARTADA"}.get(status_atual, "🔴 PENDENTE")
                    with st.expander(f"[{regra}] {dt_str} — {badge}", expanded=(status_atual == "PENDENTE")):
                        if status_atual != "PENDENTE":
                            obs_atual = resolucoes.get(oid, {}).get("observacao", "")
                            if obs_atual:
                                st.caption(f"Observação: {obs_atual}")
                            if st.button("↩️ Re-abrir como pendente",
                                         key=f"notif_reopen_{i}_{j}"):
                                remover_resolucao(_AUDIT_DB, oid)
                                st.rerun()
                        else:
                            c_s, c_o = st.columns([1, 3])
                            with c_s:
                                novo = st.radio("Resolução",
                                               ["JUSTIFICADA", "DESCARTADA"],
                                               key=f"notif_res_{i}_{j}")
                            with c_o:
                                obs = st.text_area("Observação",
                                                   height=68,
                                                   key=f"notif_obs_{i}_{j}",
                                                   placeholder="Justificativa ou motivo do descarte…")
                            if st.button("💾 Gravar",
                                         key=f"notif_salvar_{i}_{j}",
                                         type="primary"):
                                salvar_resolucao(_AUDIT_DB, oid, novo, obs)
                                st.success(f"Marcada como **{novo}**.")
                                st.rerun()


def _enviar_email_smtp(notif: dict, cfg: dict) -> tuple[bool, str]:
    try:
        remetente = cfg["remetente"].strip()
        senha     = cfg["senha"].replace(" ", "").strip()
        dest      = cfg["destinatario"].strip()
        msg = MIMEMultipart()
        msg["From"]    = remetente
        msg["To"]      = dest
        msg["Subject"] = (
            f"Notificação de Frota — Placa {notif.get('placa','?')} "
            f"[{notif.get('gravidade_max','?')}]"
        )
        msg.attach(MIMEText(notif.get("texto_notificacao", notif.get("minuta", "")), "plain", "utf-8"))
        ctx = ssl.create_default_context()
        if cfg.get("usar_ssl", True):
            with smtplib.SMTP_SSL(cfg["host"].strip(), int(cfg["port"]), context=ctx) as s:
                s.login(remetente, senha)
                s.sendmail(remetente, [dest], msg.as_string())
        else:
            with smtplib.SMTP(cfg["host"].strip(), int(cfg["port"])) as s:
                s.ehlo()
                s.starttls(context=ctx)
                s.login(remetente, senha)
                s.sendmail(remetente, [dest], msg.as_string())
        return True, ""
    except Exception as e:
        return False, str(e)


def _render_exportacao(resultado: dict, df_filt: pd.DataFrame) -> None:
    ocs = resultado.get("ocorrencias", resultado.get("alertas", []))
    if not ocs:
        st.info("Sem ocorrências para exportar.")
        return
    df_oc = pd.DataFrame(ocs) if isinstance(ocs, list) else ocs
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        df_filt.drop(
            columns=["data", "ano", "mes", "mes_nome", "ano_mes"], errors="ignore"
        ).to_excel(w, sheet_name="Dados Brutos", index=False)
        df_oc.to_excel(w, sheet_name="Ocorrências", index=False)
    buf.seek(0)
    st.download_button(
        "📥 Baixar relatório Excel",
        data=buf,
        file_name="auditoria_frota.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )


def _render_gestao_contatos() -> None:
    """Expander para cadastrar/editar e-mails de contato por secretaria ou condutor."""
    with st.expander("📋 Gerenciar contatos de e-mail", expanded=False):
        try:
            data = json.loads(_CONFIG_JSON.read_text(encoding="utf-8"))
        except Exception:
            data = {}
        contatos = data.get("contatos", {"secretarias": {}, "condutores": {}})

        tab_sec, tab_cond = st.tabs(["Por Secretaria", "Por Condutor"])

        with tab_sec:
            secs = contatos.get("secretarias", {})
            if secs:
                df_sec = pd.DataFrame(
                    [{"Secretaria": k, "E-mail": v} for k, v in secs.items()]
                )
                st.dataframe(df_sec, use_container_width=True, hide_index=True)
            else:
                st.caption("Nenhum e-mail por secretaria cadastrado.")

            st.markdown("**Adicionar / Atualizar**")
            c1, c2, c3 = st.columns([1, 2, 1])
            with c1:
                sec_chave = st.text_input("Sigla", key="cont_sec_chave",
                                          placeholder="ex: SEMUTTRANS")
            with c2:
                sec_email = st.text_input("E-mail", key="cont_sec_email",
                                          placeholder="gestor@prefeitura.sp.gov.br")
            with c3:
                st.write("")
                st.write("")
                if st.button("💾 Salvar", key="cont_sec_salvar"):
                    if sec_chave and sec_email:
                        _save_contato("secretaria", sec_chave, sec_email)
                        st.success(f"Salvo: {sec_chave}")
                        st.rerun()
                    else:
                        st.warning("Preencha a sigla e o e-mail.")

        with tab_cond:
            conds = contatos.get("condutores", {})
            if conds:
                df_cond = pd.DataFrame(
                    [{"Condutor": k, "E-mail": v} for k, v in conds.items()]
                )
                st.dataframe(df_cond, use_container_width=True, hide_index=True)
            else:
                st.caption("Nenhum e-mail por condutor cadastrado.")

            st.markdown("**Adicionar / Atualizar**")
            c1, c2, c3 = st.columns([2, 2, 1])
            with c1:
                cond_chave = st.text_input("Nome do condutor", key="cont_cond_chave",
                                           placeholder="ex: JOSE DA SILVA")
            with c2:
                cond_email = st.text_input("E-mail", key="cont_cond_email",
                                           placeholder="condutor@prefeitura.sp.gov.br")
            with c3:
                st.write("")
                st.write("")
                if st.button("💾 Salvar", key="cont_cond_salvar"):
                    if cond_chave and cond_email:
                        _save_contato("condutor", cond_chave, cond_email)
                        st.success(f"Salvo: {cond_chave}")
                        st.rerun()
                    else:
                        st.warning("Preencha o nome e o e-mail.")


def _render_manual() -> None:
    if _MANUAL_HTML.exists():
        st.components.v1.html(_MANUAL_HTML.read_text(encoding="utf-8"), height=900, scrolling=True)
    else:
        st.info("Manual não encontrado em `assets/manual_auditoria.html`.")


# ---------------------------------------------------------------------------
# Helpers internos de run_auditoria_page
# ---------------------------------------------------------------------------

def _render_sidebar(df_all: pd.DataFrame, db_min, db_max) -> dict:
    """Renderiza a sidebar e devolve um dict com todas as seleções do usuário."""
    _hoje    = date.today()
    _ini_mes = _hoje.replace(day=1)
    _default_ini = max(_ini_mes, db_min)
    _default_fim = min(_hoje,    db_max)

    _DASH_PATH = Path(__file__).resolve().parent.parent.parent / "pages" / "Dashboard.py"

    with st.sidebar:
        st.markdown("### 🔍 Auditoria de Frota")
        if _DASH_PATH.exists():
            st.page_link("pages/Dashboard.py", label="← Dashboard", icon="📊")
        st.divider()
        st.subheader("🗓️ Período")
        data_ini = st.date_input("De",  value=_default_ini, min_value=db_min, max_value=db_max, key="aud_di")
        data_fim = st.date_input("Até", value=_default_fim, min_value=db_min, max_value=db_max, key="aud_df")

        st.subheader("🔎 Filtros")
        _fuel_opts  = ["Todos"] + sorted(_FUEL_DISPLAY.keys())
        _und_opts   = ["Todas"] + sorted(df_all["secretaria"].dropna().astype(str).unique().tolist())
        _marca_opts = ["Todas"] + sorted(df_all["marca"].dropna().astype(str).unique().tolist())
        _cond_opts  = ["Todos"] + sorted(df_all["condutor"].dropna().astype(str).unique().tolist())
        _placa_opts = ["Todas"] + sorted(df_all["placa"].dropna().astype(str).unique().tolist())

        sel_fuel  = st.selectbox("Combustível", _fuel_opts,  key="aud_fuel")
        sel_und   = st.selectbox("Secretaria",  _und_opts,   key="aud_und")
        sel_marca = st.selectbox("Marca",       _marca_opts, key="aud_marca")

        # Modelos dependentes da marca
        _df_mod   = df_all if sel_marca == "Todas" else df_all[df_all["marca"] == sel_marca]
        _mod_opts = sorted(
            _df_mod["modelo"].dropna().astype(str).str.strip()
            .replace("", pd.NA).dropna().unique().tolist()
        )
        if st.button("Selecionar todos os modelos", use_container_width=True, key="aud_all_mod"):
            st.session_state["aud_modelos"] = _mod_opts
            st.rerun()
        sel_modelos = st.multiselect(
            "Modelos", options=_mod_opts, key="aud_modelos",
            help="Deixe vazio = todos os modelos da marca selecionada.",
        )

        sel_cond  = st.selectbox("Motorista", _cond_opts,  key="aud_cond")
        sel_placa = st.selectbox("Placa",     _placa_opts, key="aud_placa")

        st.divider()
        sigma_mult = st.slider(
            "Sensibilidade outlier (σ)", 0.5, 4.0,
            float(st.session_state.get("aud_sigma", 2.0)), 0.1,
            help="consumo < média − σ×desvio → outlier. Menor σ = mais sensível.",
            key="aud_sigma_slider",
        )

        st.divider()
        col_ap, col_lim = st.columns(2)
        with col_ap:
            aplicar = st.button("🚀 Aplicar", type="primary", use_container_width=True, key="aud_aplicar")
        with col_lim:
            limpar  = st.button("🔄 Limpar",  use_container_width=True, key="aud_limpar")

        if limpar:
            st.session_state["aud_resultado"] = None
            st.session_state["aud_df_filt"]   = None
            st.rerun()

    return dict(
        data_ini=data_ini, data_fim=data_fim,
        sel_fuel=sel_fuel, sel_und=sel_und, sel_marca=sel_marca,
        sel_modelos=sel_modelos, sel_cond=sel_cond, sel_placa=sel_placa,
        sigma_mult=sigma_mult, aplicar=aplicar,
    )


def _filter_df(df_all: pd.DataFrame, data_ini, data_fim,
               sel_fuel, sel_und, sel_marca, sel_modelos,
               sel_cond, sel_placa) -> pd.DataFrame:
    """Filtra df_all pelo período e pelas seleções de sidebar."""
    mask = (
        df_all["data_hora"].notna()
        & (df_all["data_hora"].dt.date >= data_ini)
        & (df_all["data_hora"].dt.date <= data_fim)
    )
    df = df_all[mask].copy()
    if sel_fuel != "Todos":
        df = df[df["combustivel"] == _FUEL_DISPLAY.get(sel_fuel, sel_fuel)]
    if sel_und != "Todas":
        df = df[df["secretaria"] == sel_und]
    if sel_marca != "Todas":
        df = df[df["marca"] == sel_marca]
    if sel_modelos:
        df = df[df["modelo"].isin(sel_modelos)]
    if sel_cond != "Todos":
        df = df[df["condutor"] == sel_cond]
    if sel_placa != "Todas":
        df = df[df["placa"] == sel_placa]
    return df


def _run_pipeline(df, df_all, sigma_mult) -> None:
    """Executa o pipeline multiagente e armazena resultado no session_state."""
    _BRIDGE = {
        "valor":        "valor_total",
        "km_por_litro": "km_l_informado",
        "posto":        "estabelecimento",
        "secretaria":   "unidade",
        "combustivel":  "produto",
    }
    with st.spinner("Executando pipeline multiagente de auditoria…"):
        try:
            df_para_pipeline = df.rename(columns=_BRIDGE)
            orq = OrchestradorAuditoria(metadata={
                "outlier_sigma_mult":    sigma_mult,
                "df_historico_completo": df_all,
            })
            resultado = orq.run_pipeline(df_para_pipeline)
            st.session_state["aud_resultado"] = resultado
            st.session_state["aud_df_filt"]   = df.copy()
            st.rerun()
        except Exception as exc:
            import traceback
            st.error(f"Erro no pipeline: {exc}")
            st.code(traceback.format_exc())


# ---------------------------------------------------------------------------
# Ponto de entrada da página
# ---------------------------------------------------------------------------

def run_auditoria_page() -> None:
    st.set_page_config(
        page_title="Auditoria de Frota",
        page_icon="🔍",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    _apply_style()

    st.session_state.setdefault("aud_resultado", None)
    st.session_state.setdefault("aud_df_filt",   None)
    st.session_state.setdefault("aud_sigma",      2.0)

    db_path = str(settings.db_path)
    try:
        df_all = _load_all(db_path)
    except Exception as exc:
        st.error(f"Não foi possível ler o banco de dados: {exc}")
        return

    if df_all.empty:
        st.warning("Banco de dados vazio ou sem a tabela 'abastecimentos'.")
        return

    db_min, db_max = _date_bounds(df_all)

    sel = _render_sidebar(df_all, db_min, db_max)
    data_ini    = sel["data_ini"]
    data_fim    = sel["data_fim"]
    sigma_mult  = sel["sigma_mult"]
    aplicar     = sel["aplicar"]

    if data_ini > data_fim:
        st.warning("A data inicial deve ser anterior ou igual à data final.")
        return

    df = _filter_df(
        df_all, data_ini, data_fim,
        sel["sel_fuel"], sel["sel_und"], sel["sel_marca"],
        sel["sel_modelos"], sel["sel_cond"], sel["sel_placa"],
    )

    _n_periodo = int(
        (df_all["data_hora"].notna()
         & (df_all["data_hora"].dt.date >= data_ini)
         & (df_all["data_hora"].dt.date <= data_fim)).sum()
    )
    st.caption(
        f"🗄️ **{Path(db_path).name}** · "
        f"Período selecionado: **{data_ini.strftime('%d/%m/%Y')} – {data_fim.strftime('%d/%m/%Y')}** "
        f"({_n_periodo:,} registros no período) · "
        f"**{len(df):,}** após todos os filtros"
    )

    if df.empty:
        st.warning("⚠️ Nenhum dado encontrado para os filtros selecionados.")
        return

    if aplicar:
        st.session_state["aud_sigma"] = sigma_mult
        _run_pipeline(df, df_all, sigma_mult)

    resultado = st.session_state.get("aud_resultado")
    df_filt   = st.session_state.get("aud_df_filt", df)

    # ── Abas ────────────────────────────────────────────────────────────────
    tab_vis, tab_alertas, tab_oc, tab_notif, tab_export, tab_man = st.tabs([
        "📊 Visão Geral",
        "⚠️ Alertas Op.",
        "🚨 Ocorrências",
        "📨 Notificações",
        "📥 Exportar",
        "📖 Manual",
    ])

    with tab_vis:
        _render_stats_and_scatter(df, sigma_mult, sel["sel_modelos"], period=(data_ini, data_fim))
        if resultado is not None:
            st.divider()
            _render_painel_executivo(resultado)
        elif aplicar is False:
            st.info("👆 Clique em **🚀 Aplicar** na sidebar para executar a auditoria por agentes.")

    with tab_alertas:
        _render_alertas_operacionais(df)

    with tab_oc:
        if resultado is None:
            st.info("👆 Clique em **🚀 Aplicar** na sidebar para gerar as ocorrências.")
        else:
            _render_ocorrencias(resultado)

    with tab_notif:
        if resultado is None:
            st.info("👆 Clique em **🚀 Aplicar** na sidebar para gerar as notificações.")
        else:
            _render_notificacoes(resultado)

    with tab_export:
        if resultado is None:
            st.info("👆 Clique em **🚀 Aplicar** na sidebar para habilitar a exportação.")
        else:
            _render_exportacao(resultado, df_filt)
        st.divider()
        _render_gestao_contatos()

    with tab_man:
        _render_manual()
