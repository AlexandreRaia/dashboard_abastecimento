"""
Aplicação Streamlit — Auditor de Abastecimento de Frota.

Ponto de entrada da interface web. Responsável por:
- Receber o upload da planilha Excel de abastecimentos.
- Acionar o pipeline multiagente (OrchestradorAuditoria).
- Exibir o Painel Executivo, Indicadores Gerenciais e Ocorrências.
- Permitir filtros por data, placa e condutor.
- Exportar resultados em Excel e exibir minutas de notificação.
- Apresentar o manual operacional integrado.

Executar com:
    streamlit run main.py
"""
import sqlite3
import streamlit as st
import streamlit.components.v1 as st_components
import pandas as pd
import plotly.express as px
import numpy as np
import unicodedata
import hashlib
import re
from datetime import datetime
from pathlib import Path
from io import BytesIO

from agents import OrchestradorAuditoria
from agents.config import canonicalizar_modelo

PROJECT_ROOT = Path(__file__).resolve().parent.parent  # raiz do projeto
RELATORIO_DB = PROJECT_ROOT / "relatorio.db"            # fonte somente-leitura
RESULTADOS_DB = PROJECT_ROOT / "auditoria_resultados.db"  # banco de resultados
SQLITE_DB_PATH = RESULTADOS_DB  # alias de compatibilidade
MANUAL_DOCX_PATH = PROJECT_ROOT / "manual e treinamento.docx"  # opcional
ROLLBACK_TAG_NAME = "backup-pre-layout-manual"
MAX_ANALISE_ROWS = 25000
MAX_RENDER_ROWS = 5000

# Centralize vehicle annotations/tank capacities
VEHICLE_TANK_CAPACITIES = {
    'Mobi': 47,
    'Spin': 53,
    'Master': 105,
    'Van': 105,
    'Caminhão': 120,
}

# Helper function for robust date processing
def process_date_column(df_column):
    dates = pd.to_datetime(df_column, dayfirst=True, errors='coerce')
    numeric_like_values = pd.to_numeric(df_column, errors='coerce')
    excel_serial_mask = dates.isna() & numeric_like_values.notna()
    dates.loc[excel_serial_mask] = pd.to_datetime(numeric_like_values[excel_serial_mask], unit='D', origin='1899-12-30', errors='coerce')
    return dates


def normalize_text(value):
    if pd.isna(value):
        return ""
    text = str(value).strip().lower()
    return "".join(
        ch for ch in unicodedata.normalize("NFKD", text)
        if not unicodedata.combining(ch)
    )


def first_existing_column(df, candidates):
    return next((col for col in candidates if col in df.columns), None)


def safe_index(options, value):
    if value in options:
        return options.index(value)
    return 0


def render_dataframe_limited(df: pd.DataFrame, max_rows: int = MAX_RENDER_ROWS):
    total_rows = len(df)
    if total_rows > max_rows:
        st.warning(
            f"Exibicao limitada a {max_rows:,} linhas para evitar travamento da interface. "
            f"Total disponivel: {total_rows:,} linhas.".replace(',', '.')
        )
        st.dataframe(df.head(max_rows), use_container_width=True)
    else:
        st.dataframe(df, use_container_width=True)


@st.cache_data(show_spinner=False)
def carregar_manual_pdf_bytes(caminho_pdf: str):
    caminho = Path(caminho_pdf)
    if not caminho.exists() or caminho.suffix.lower() != '.pdf':
        return None
    return caminho.read_bytes()


def _limpar_prefixo_manual(texto: str) -> str:
    texto_limpo = str(texto or "").strip()
    while texto_limpo and not (texto_limpo[0].isalnum() or texto_limpo[0] in "("):
        texto_limpo = texto_limpo[1:].lstrip()
    return texto_limpo.strip()


def _eh_titulo_manual(texto: str) -> bool:
    texto_limpo = _limpar_prefixo_manual(texto)
    if not texto_limpo:
        return False
    if re.match(r'^\d+\.\s+', texto_limpo):
        return True
    return (
        len(texto_limpo) <= 90
        and texto_limpo.upper() == texto_limpo
        and any(ch.isalpha() for ch in texto_limpo)
    )


@st.cache_data(show_spinner=False)
def carregar_manual_docx_estruturado(caminho_docx: str):
    caminho = Path(caminho_docx)
    if not caminho.exists() or caminho.suffix.lower() != '.docx':
        return None

    from docx import Document

    doc = Document(caminho)
    secoes = []
    secao_atual = {"titulo": "Visao Geral", "linhas": []}

    for paragrafo in doc.paragraphs:
        texto = str(paragrafo.text or "").strip()
        if not texto:
            continue

        if _eh_titulo_manual(texto):
            if secao_atual["titulo"] or secao_atual["linhas"]:
                secoes.append(secao_atual)
            secao_atual = {
                "titulo": _limpar_prefixo_manual(texto),
                "linhas": [],
            }
            continue

        secao_atual["linhas"].append(texto)

    if secao_atual["titulo"] or secao_atual["linhas"]:
        secoes.append(secao_atual)

    secoes = [secao for secao in secoes if secao.get("titulo") or secao.get("linhas")]
    if not secoes:
        return None

    titulo_base = next(
        (_limpar_prefixo_manual(par.text) for par in doc.paragraphs if str(par.text or "").strip()),
        caminho.stem,
    )

    return {
        "titulo": titulo_base,
        "arquivo": caminho.name,
        "bytes": caminho.read_bytes(),
        "secoes": secoes,
    }


def aplicar_estilo_relatorio():
    st.markdown(
        """
        <style>
        .painel-executivo {
            background: linear-gradient(135deg, #f8fafc 0%, #eef6ff 100%);
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: 22px;
            padding: 1.2rem 1.25rem;
            margin-bottom: 1rem;
        }
        .painel-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(190px, 1fr));
            gap: 0.85rem;
            margin-top: 0.85rem;
        }
        .painel-card {
            background: rgba(255, 255, 255, 0.92);
            border: 1px solid rgba(148, 163, 184, 0.22);
            border-radius: 18px;
            padding: 1rem;
            box-shadow: 0 12px 30px rgba(15, 23, 42, 0.06);
        }
        .painel-card.alerta {
            background: linear-gradient(180deg, #fff1f2 0%, #ffe4e6 100%);
            border-color: rgba(190, 24, 93, 0.18);
        }
        .painel-card.media {
            background: linear-gradient(180deg, #fff7ed 0%, #ffedd5 100%);
            border-color: rgba(194, 65, 12, 0.18);
        }
        .painel-card .rotulo {
            display: block;
            color: #475569;
            font-size: 0.82rem;
            margin-bottom: 0.45rem;
            text-transform: uppercase;
            letter-spacing: 0.04em;
        }
        .painel-card .valor {
            display: block;
            color: #0f172a;
            font-size: 1.9rem;
            font-weight: 700;
            line-height: 1.05;
        }
        .painel-card .apoio {
            display: block;
            color: #334155;
            font-size: 0.92rem;
            margin-top: 0.5rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def renderizar_painel_executivo(resultado_auditoria: dict, sqlite_info: dict):
    relatorio_agente = resultado_auditoria.get('relatorio', {})
    notificacoes_agente = resultado_auditoria.get('notificacoes', [])
    res_exec = relatorio_agente.get('resumo_executivo', {})
    ocorrencias = resultado_auditoria.get('ocorrencias', [])
    df_ocs = pd.DataFrame(ocorrencias) if ocorrencias else pd.DataFrame()

    n_alta = int(res_exec.get('ocorrencias_alta', 0))
    n_media = int(res_exec.get('ocorrencias_media', 0))
    total_ocorrencias = int(res_exec.get('total_ocorrencias', len(df_ocs)))
    total_registros = int(res_exec.get('total_registros', 0))
    total_placas = int(res_exec.get('total_placas', 0))
    placas_criticas = int(df_ocs[df_ocs.get('gravidade_final') == 'ALTA']['placa'].nunique()) if not df_ocs.empty and 'gravidade_final' in df_ocs.columns and 'placa' in df_ocs.columns else 0

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("Ocorrencias ALTA", f"{n_alta}", delta=f"{placas_criticas} placa(s) critica(s)")
    with c2:
        st.metric("Ocorrencias MEDIA", f"{n_media}")
        st.caption("Monitoramento preventivo antes de escalar")
    with c3:
        st.metric("Notificacoes", f"{len(notificacoes_agente)}")
        st.caption("Minutas administrativas prontas para uso")
    with c4:
        st.metric("Base auditada", f"{total_registros}", delta=f"{total_placas} placa(s)")
        st.caption(f"{total_ocorrencias} ocorrencia(s)")


def renderizar_manual_no_dashboard(manual_data: dict):
    HTML_MANUAL = """
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;font-size:14px;line-height:1.6;color:#1e293b;background:#f8fafc;padding:0}
.manual{max-width:860px;margin:0 auto;padding:20px 16px 40px}
.manual-header{background:linear-gradient(135deg,#1e3a8a 0%,#1e40af 60%,#2563eb 100%);color:white;border-radius:16px;padding:32px 32px 28px;margin-bottom:24px;text-align:center}
.manual-header h1{font-size:22px;font-weight:700;margin-bottom:6px}
.manual-header p{font-size:14px;opacity:.85}
.manual-header .badge{display:inline-block;background:rgba(255,255,255,.15);border:1px solid rgba(255,255,255,.3);border-radius:20px;padding:3px 12px;font-size:12px;margin-top:10px}
.toc{background:white;border:1px solid #e2e8f0;border-radius:12px;padding:20px 24px;margin-bottom:24px}
.toc h3{font-size:13px;text-transform:uppercase;letter-spacing:.06em;color:#64748b;margin-bottom:12px}
.toc ol{list-style:none;display:grid;grid-template-columns:1fr 1fr;gap:6px}
.toc ol li a{text-decoration:none;color:#1e40af;font-size:13px;font-weight:500}
.toc ol li a:hover{text-decoration:underline}
.section{background:white;border:1px solid #e2e8f0;border-radius:14px;padding:26px 28px;margin-bottom:20px;box-shadow:0 1px 4px rgba(0,0,0,.04)}
.section-title{display:flex;align-items:center;gap:12px;margin-bottom:18px}
.section-num{min-width:36px;height:36px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:14px;font-weight:700;color:white}
.section-title h2{font-size:17px;font-weight:700;color:#0f172a}
.section-title .icon{font-size:20px}
.c-blue{background:#1e40af}.c-sky{background:#0284c7}.c-purple{background:#7c3aed}.c-red{background:#dc2626}.c-indigo{background:#4f46e5}.c-orange{background:#d97706}.c-green{background:#16a34a}.c-yellow{background:#ca8a04}.c-teal{background:#0d9488}.c-slate{background:#475569}
.formula-box{background:#eff6ff;border:1px solid #bfdbfe;border-left:4px solid #1e40af;border-radius:8px;padding:14px 18px;font-family:'Courier New',Courier,monospace;font-size:14px;color:#1e3a8a;margin:12px 0;line-height:1.8}
.formula-label{font-size:11px;text-transform:uppercase;letter-spacing:.06em;color:#3b82f6;font-weight:600;margin-bottom:4px;font-family:-apple-system,sans-serif}
.example{border-radius:8px;padding:12px 16px;margin:10px 0;font-size:13px}
.example-normal{background:#f0fdf4;border-left:4px solid #16a34a}
.example-alert{background:#fff7ed;border-left:4px solid #d97706}
.example-critical{background:#fef2f2;border-left:4px solid #dc2626}
.example-info{background:#f0f9ff;border-left:4px solid #0284c7}
.example strong{display:block;font-size:12px;color:#6b7280;text-transform:uppercase;letter-spacing:.04em;margin-bottom:6px}
.rules-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(220px,1fr));gap:12px;margin-top:8px}
.rule-card{border:1px solid #e2e8f0;border-radius:10px;padding:14px;position:relative;overflow:hidden}
.rule-card::before{content:'';position:absolute;top:0;left:0;right:0;height:3px}
.rule-card.alta::before{background:#dc2626}.rule-card.media::before{background:#d97706}.rule-card.baixa::before{background:#16a34a}
.rule-code{font-family:'Courier New',monospace;font-size:12px;font-weight:700;color:#475569;background:#f1f5f9;border-radius:4px;padding:2px 6px;display:inline-block;margin-bottom:6px}
.rule-title{font-size:13px;font-weight:600;color:#0f172a;margin-bottom:4px}
.rule-desc{font-size:12px;color:#64748b;line-height:1.5}
.rule-example{font-size:11px;color:#94a3b8;margin-top:6px;padding-top:6px;border-top:1px solid #f1f5f9}
.badge-grav{font-size:10px;font-weight:600;padding:1px 6px;border-radius:10px;float:right}
.badge-alta{background:#fee2e2;color:#991b1b}.badge-media{background:#fef3c7;color:#92400e}.badge-baixa{background:#dcfce7;color:#166534}
.agent-pipeline{display:flex;flex-wrap:wrap;align-items:center;gap:4px;margin:16px 0}
.agent-step{background:#eff6ff;border:1px solid #bfdbfe;border-radius:8px;padding:6px 12px;font-size:12px;font-weight:600;color:#1e40af;white-space:nowrap}
.agent-arrow{color:#94a3b8;font-size:16px}
.agent-cards{display:grid;grid-template-columns:repeat(auto-fill,minmax(200px,1fr));gap:10px;margin-top:14px}
.agent-card{border:1px solid #e2e8f0;border-radius:10px;padding:14px;background:#fafafa}
.agent-card .agent-icon{font-size:22px;margin-bottom:6px}
.agent-card .agent-name{font-size:12px;font-weight:700;color:#1e40af;margin-bottom:4px;font-family:'Courier New',monospace}
.agent-card .agent-role{font-size:13px;font-weight:600;color:#0f172a;margin-bottom:6px}
.agent-card ul{list-style:none}
.agent-card ul li{font-size:12px;color:#64748b;padding:2px 0}
.agent-card ul li::before{content:'→ ';color:#3b82f6}
.grav-table{width:100%;border-collapse:collapse;margin-top:12px;font-size:13px}
.grav-table th{background:#f1f5f9;padding:8px 12px;text-align:left;font-size:11px;text-transform:uppercase;letter-spacing:.04em;color:#64748b;border-bottom:2px solid #e2e8f0}
.grav-table td{padding:9px 12px;border-bottom:1px solid #f1f5f9;vertical-align:top}
.grav-table tr:last-child td{border-bottom:none}
.chip-alta{background:#fee2e2;color:#991b1b;border-radius:4px;padding:2px 8px;font-weight:600;font-size:12px}
.chip-media{background:#fef3c7;color:#92400e;border-radius:4px;padding:2px 8px;font-weight:600;font-size:12px}
.chip-baixa{background:#dcfce7;color:#166534;border-radius:4px;padding:2px 8px;font-weight:600;font-size:12px}
.checklist{list-style:none}
.checklist li{padding:8px 0 8px 28px;position:relative;border-bottom:1px solid #f1f5f9;font-size:13px;color:#334155}
.checklist li:last-child{border-bottom:none}
.checklist li::before{content:'✓';position:absolute;left:0;color:#16a34a;font-weight:700}
.checklist.dont li::before{content:'✗';color:#dc2626}
.two-col{display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-top:8px}
@media(max-width:560px){.two-col{grid-template-columns:1fr}}
.flow-list{counter-reset:flow-counter;list-style:none;padding:0}
.flow-list li{counter-increment:flow-counter;display:flex;align-items:flex-start;gap:10px;padding:8px 0;border-bottom:1px solid #f1f5f9;font-size:13px;color:#334155}
.flow-list li:last-child{border-bottom:none}
.flow-list li::before{content:counter(flow-counter);background:#1e40af;color:white;min-width:22px;height:22px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:700;flex-shrink:0;margin-top:1px}
.exercise{border:1px solid #e2e8f0;border-radius:10px;padding:14px 16px;margin:10px 0}
.ex-header{display:flex;align-items:center;gap:8px;margin-bottom:10px}
.ex-num{background:#0d9488;color:white;border-radius:50%;width:28px;height:28px;display:flex;align-items:center;justify-content:center;font-size:12px;font-weight:700}
.ex-title{font-size:13px;font-weight:600;color:#0f172a}
.ex-data{display:flex;gap:16px;flex-wrap:wrap;margin:8px 0}
.ex-data-item .label{color:#94a3b8;font-size:11px;text-transform:uppercase;letter-spacing:.04em;display:block}
.ex-data-item .value{font-weight:600;color:#0f172a;font-size:15px}
.ex-result{border-radius:6px;padding:8px 12px;font-size:13px;font-weight:500;margin-top:8px}
.ex-ok{background:#f0fdf4;color:#166534;border:1px solid #bbf7d0}
.ex-nok{background:#fef2f2;color:#991b1b;border:1px solid #fecaca}
.quote{background:linear-gradient(135deg,#f0f9ff 0%,#eff6ff 100%);border-left:4px solid #1e40af;border-radius:0 10px 10px 0;padding:16px 20px;font-size:15px;font-style:italic;color:#1e3a8a;margin:14px 0;line-height:1.7}
.subsection{margin:16px 0}
.subsection h3{font-size:14px;font-weight:600;color:#334155;margin-bottom:8px;display:flex;align-items:center;gap:6px}
p{margin:8px 0;color:#334155;font-size:13px;line-height:1.65}
strong{color:#0f172a}
ul.default{padding-left:20px}
ul.default li{font-size:13px;color:#334155;padding:3px 0;line-height:1.5}
.info-box{background:#fafafa;border:1px solid #e2e8f0;border-radius:8px;padding:12px 16px;font-size:13px;color:#334155;margin:8px 0}
.compare{display:grid;grid-template-columns:1fr 1fr;gap:10px;margin:10px 0}
@media(max-width:480px){.compare{grid-template-columns:1fr}}
.compare-col{border-radius:8px;padding:14px}
.before{background:#fef2f2;border:1px solid #fecaca}
.after{background:#f0fdf4;border:1px solid #bbf7d0}
.compare-col h4{font-size:12px;text-transform:uppercase;letter-spacing:.05em;font-weight:700;margin-bottom:8px}
.before h4{color:#991b1b}.after h4{color:#166534}
.compare-col ul{list-style:none}
.compare-col ul li{font-size:12px;padding:3px 0;color:#334155}
.before ul li::before{content:'✗ ';color:#dc2626}
.after ul li::before{content:'✓ ';color:#16a34a}
</style></head><body><div class="manual">

<div class="manual-header">
    <div style="font-size:32px;margin-bottom:10px">📘</div>
    <h1>Manual Operacional e de Treinamento</h1>
    <p>Sistema de Auditoria de Abastecimento de Frota Municipal</p>
    <span class="badge">Versão consolidada · Santana de Parnaíba · 2026</span>
</div>

<div class="toc">
    <h3>Índice</h3>
    <ol>
        <li><a href="#s1">1. Objetivo do Sistema</a></li>
        <li><a href="#s2">2. Como o Sistema Funciona</a></li>
        <li><a href="#s3">3. Fundamentos Estatísticos</a></li>
        <li><a href="#s4">4. Regras de Auditoria (R01–R10)</a></li>
        <li><a href="#s5">5. Arquitetura de Agentes de IA</a></li>
        <li><a href="#s6">6. Classificação de Gravidade</a></li>
        <li><a href="#s7">7. Ação da Equipe</a></li>
        <li><a href="#s8">8. Limites do Sistema</a></li>
        <li><a href="#s9">9. Treinamento Prático</a></li>
        <li><a href="#s10">10. Mentalidade Final</a></li>
    </ol>
</div>

<!-- SEÇÃO 1 -->
<div class="section" id="s1">
    <div class="section-title">
        <div class="section-num c-blue">1</div>
        <span class="icon">🎯</span>
        <h2>Objetivo do Sistema</h2>
    </div>
    <p>O sistema realiza <strong>auditoria técnica automatizada</strong> dos registros de abastecimento da frota municipal, identificando inconsistências com base em evidências numéricas.</p>
    <p>A análise combina quatro pilares:</p>
    <ol class="flow-list" style="margin-top:10px">
        <li>Regras objetivas (R01 a R10) — critérios técnicos fixos e documentados</li>
        <li>Estatística (média e desvio padrão) — comportamento esperado por veículo</li>
        <li>Histórico individual por veículo — comparação com o padrão próprio de cada placa</li>
        <li>Evidência numérica — toda ocorrência tem base em dados verificáveis</li>
    </ol>
    <div class="quote" style="margin-top:16px">💡 &ldquo;Toda ocorrência precisa de justificativa baseada em números, não em opinião.&rdquo;</div>
</div>

<!-- SEÇÃO 2 -->
<div class="section" id="s2">
    <div class="section-title">
        <div class="section-num c-sky">2</div>
        <span class="icon">🔄</span>
        <h2>Como o Sistema Funciona</h2>
    </div>
    <p>O processamento segue um fluxo sequencial de agentes especializados, do dado bruto até a notificação administrativa:</p>
    <ol class="flow-list" style="margin-top:12px">
        <li>Leitura e padronização da planilha de abastecimento</li>
        <li>Limpeza e validação dos dados (remoção de inconsistências óbvias)</li>
        <li>Armazenamento no banco de dados local (SQLite)</li>
        <li>Aplicação das regras de auditoria (R01–R10)</li>
        <li>Comparação com o histórico individual de cada veículo</li>
        <li>Classificação de gravidade das ocorrências (ALTA / MÉDIA / BAIXA)</li>
        <li>Geração do relatório técnico consolidado</li>
        <li>Geração automática das minutas de notificação administrativa</li>
    </ol>
    <div class="example example-info" style="margin-top:14px">
        <strong>Tradução simples para a equipe</strong>
        👉 O sistema funciona como um <em>auditor automático baseado em dados</em>: não opina, não acusa — apenas aponta onde os números não fazem sentido.
    </div>
</div>

<!-- SEÇÃO 3 -->
<div class="section" id="s3">
    <div class="section-title">
        <div class="section-num c-purple">3</div>
        <span class="icon">📊</span>
        <h2>Fundamentos Estatísticos</h2>
    </div>

    <div class="subsection">
        <h3>🧠 3.1 Consumo (km/L)</h3>
        <p>O consumo é a métrica principal de eficiência do veículo. É calculado dividindo-se a distância percorrida pelo volume de combustível abastecido.</p>
        <div class="formula-label">Fórmula</div>
        <div class="formula-box">consumo = km_rodado ÷ litros</div>
        <div class="two-col">
            <div class="example example-normal">
                <strong>Exemplo normal</strong>
                KM rodado: 100 &nbsp;|&nbsp; Litros: 10<br>
                <span style="font-weight:700">Consumo = 10 km/L ✅</span>
            </div>
            <div class="example example-critical">
                <strong>Exemplo crítico</strong>
                KM rodado: 100 &nbsp;|&nbsp; Litros: 20<br>
                <span style="font-weight:700">Consumo = 5 km/L 🚨 Muito baixo</span>
            </div>
        </div>
    </div>

    <div class="subsection">
        <h3>📊 3.2 Média (μ)</h3>
        <p>A média aritmética representa o consumo <strong>típico</strong> de um veículo ao longo de seu histórico de abastecimentos. É usada como referência central para identificar desvios.</p>
        <div class="formula-label">Fórmula</div>
        <div class="formula-box">μ = (x₁ + x₂ + ... + xₙ) ÷ n</div>
        <div class="example example-info">
            <strong>Exemplo</strong>
            Consumos históricos: 10, 11, 9 km/L<br>
            μ = (10 + 11 + 9) ÷ 3 = <strong>10 km/L</strong><br>
            <span style="color:#64748b">→ Esse é o comportamento normal esperado para esse veículo.</span>
        </div>
    </div>

    <div class="subsection">
        <h3>📉 3.3 Desvio Percentual</h3>
        <p>Mede quanto o consumo atual se afastou da média histórica, expresso em porcentagem. É o indicador mais <strong>intuitivo</strong> para a equipe de fiscalização.</p>
        <div class="formula-label">Fórmula</div>
        <div class="formula-box">desvio% = ((x − μ) ÷ μ) × 100</div>
        <div class="example example-critical">
            <strong>Exemplo</strong>
            Média (μ) = 10 km/L &nbsp;|&nbsp; Consumo atual (x) = 6 km/L<br>
            desvio% = ((6 − 10) ÷ 10) × 100 = <strong>−40% 🚨</strong><br>
            <span style="color:#64748b">→ Consumo 40% abaixo do esperado — fora do padrão.</span>
        </div>
    </div>

    <div class="subsection">
        <h3>📊 3.4 Desvio Padrão (σ)</h3>
        <p>Mede a <strong>variabilidade</strong> do consumo do veículo. Um desvio padrão baixo indica comportamento consistente; um desvio padrão alto indica irregularidade no histórico.</p>
        <div class="formula-label">Fórmula</div>
        <div class="formula-box">σ = √[ Σ(xᵢ − μ)² ÷ (n − 1) ]</div>
        <div class="example example-info">
            <strong>Exemplo passo a passo</strong>
            Dados: 10, 11, 9 km/L &nbsp;|&nbsp; μ = 10<br>
            Diferenças: (10−10)²=0 &nbsp; (11−10)²=1 &nbsp; (9−10)²=1<br>
            σ = √[(0 + 1 + 1) ÷ (3 − 1)] = √1 = <strong>σ = 1</strong><br>
            <span style="color:#64748b">→ Variação pequena = veículo com consumo consistente.</span>
        </div>
    </div>

    <div class="subsection">
        <h3>🚨 3.5 Regra de Outlier (utilizada no sistema)</h3>
        <p>Define o <strong>limiar de alerta estatístico</strong>. Se o consumo cair abaixo desse limite, o sistema registra uma ocorrência automaticamente. O valor de N (multiplicador do desvio, geralmente entre 1,5 e 2,5) é configurável.</p>
        <div class="formula-label">Condição de alerta</div>
        <div class="formula-box">consumo &lt; μ − N × σ</div>
        <div class="example example-alert">
            <strong>Exemplo</strong>
            μ = 10 km/L &nbsp;|&nbsp; σ = 1 &nbsp;|&nbsp; N = 2<br>
            Limite inferior = 10 − (2 × 1) = <strong>8 km/L</strong><br>
            Consumo atual = 6 km/L → 6 &lt; 8 → <span style="font-weight:700">🚨 Alerta disparado</span>
        </div>
    </div>

    <div class="subsection">
        <h3>📊 3.6 KM Esperado (baseado no histórico)</h3>
        <p>Calcula quantos quilômetros o veículo <em>deveria</em> ter percorrido com o volume abastecido, usando sua média histórica de consumo. Grandes divergências entre o esperado e o realizado indicam irregularidades.</p>
        <div class="formula-label">Fórmula</div>
        <div class="formula-box">km_esperado = litros × μ_consumo</div>
        <div class="two-col">
            <div class="example example-normal">
                <strong>Situação normal</strong>
                Litros: 40 &nbsp;|&nbsp; μ = 10 km/L<br>
                KM esperado = <strong>400 km</strong><br>
                KM realizado = 380 km → ✅ coerente
            </div>
            <div class="example example-critical">
                <strong>Situação suspeita</strong>
                Litros: 40 &nbsp;|&nbsp; μ = 10 km/L<br>
                KM esperado = <strong>400 km</strong><br>
                KM realizado = 150 km → 🚨 diferença de 250 km!
            </div>
        </div>
    </div>
</div>

<!-- SEÇÃO 4 -->
<div class="section" id="s4">
    <div class="section-title">
        <div class="section-num c-red">4</div>
        <span class="icon">🔴</span>
        <h2>Regras de Auditoria (R01–R10)</h2>
    </div>
    <p>O sistema aplica 10 regras objetivas a cada registro de abastecimento. Um registro pode acionar mais de uma regra simultaneamente, o que eleva automaticamente a gravidade final.</p>
    <div class="rules-grid" style="margin-top:14px">

        <div class="rule-card alta">
            <span class="badge-grav badge-alta">ALTA</span>
            <div class="rule-code">R01</div>
            <div class="rule-title">Capacidade do tanque excedida</div>
            <div class="rule-desc">Volume abastecido superior à capacidade máxima do tanque do veículo — fisicamente impossível sem adulteração.</div>
            <div class="rule-example">Tanque: 50 L → Abastecido: 58 L 🚨</div>
        </div>

        <div class="rule-card media">
            <span class="badge-grav badge-media">MÉDIA</span>
            <div class="rule-code">R02</div>
            <div class="rule-title">Consumo fora da faixa configurada</div>
            <div class="rule-desc">KM/L abaixo do mínimo ou acima do máximo definido para o veículo ou modelo.</div>
            <div class="rule-example">KM/L fora do intervalo [Min, Máx] configurados</div>
        </div>

        <div class="rule-card alta">
            <span class="badge-grav badge-alta">ALTA</span>
            <div class="rule-code">R03</div>
            <div class="rule-title">Consumo estatisticamente crítico</div>
            <div class="rule-desc">Consumo muito abaixo da média histórica, ultrapassando o limiar estatístico: consumo &lt; μ − N × σ.</div>
            <div class="rule-example">Limite = 8 km/L &amp; consumo = 5 km/L → 🚨</div>
        </div>

        <div class="rule-card alta">
            <span class="badge-grav badge-alta">ALTA</span>
            <div class="rule-code">R04</div>
            <div class="rule-title">Hodômetro inconsistente</div>
            <div class="rule-desc">O KM registrado no abastecimento atual é menor que o KM do abastecimento anterior — hodômetro regressivo.</div>
            <div class="rule-example">KM anterior: 10.000 → KM atual: 9.800 🚨</div>
        </div>

        <div class="rule-card alta">
            <span class="badge-grav badge-alta">ALTA</span>
            <div class="rule-code">R05</div>
            <div class="rule-title">Lógica incoerente (pouco KM + muito combustível)</div>
            <div class="rule-desc">Distância percorrida muito curta com volume abastecido muito alto — matematicamente impossível para o veículo.</div>
            <div class="rule-example">50 km rodados com 40 L abastecidos 🚨</div>
        </div>

        <div class="rule-card media">
            <span class="badge-grav badge-media">MÉDIA</span>
            <div class="rule-code">R06</div>
            <div class="rule-title">Abastecimentos muito próximos</div>
            <div class="rule-desc">Dois abastecimentos do mesmo veículo com intervalo de tempo ou quilometragem suspeito, sugerindo duplicidade.</div>
            <div class="rule-example">2 abastecimentos no mesmo dia → 🚨 suspeito</div>
        </div>

        <div class="rule-card media">
            <span class="badge-grav badge-media">MÉDIA</span>
            <div class="rule-code">R07</div>
            <div class="rule-title">Preço acima do contratado</div>
            <div class="rule-desc">Valor por litro praticado pelo fornecedor supera o preço estabelecido em contrato.</div>
            <div class="rule-example">Contrato: R$ 5,80/L → Cobrado: R$ 6,40/L 🚨</div>
        </div>

        <div class="rule-card media">
            <span class="badge-grav badge-media">MÉDIA</span>
            <div class="rule-code">R08</div>
            <div class="rule-title">Valor total inconsistente</div>
            <div class="rule-desc">O valor total cobrado não corresponde ao produto entre o volume e o preço unitário.</div>
            <div class="formula-box" style="margin:6px 0;padding:8px 12px;font-size:12px">valor ≠ litros × preço_unitário</div>
        </div>

        <div class="rule-card media">
            <span class="badge-grav badge-media">MÉDIA</span>
            <div class="rule-code">R09</div>
            <div class="rule-title">Volume abastecido fora do histórico</div>
            <div class="rule-desc">Litros abastecidos em um único evento muito acima da média histórica do veículo para um abastecimento.</div>
            <div class="rule-example">Média por evento: 30 L → Abastecido: 90 L</div>
        </div>

        <div class="rule-card baixa">
            <span class="badge-grav badge-baixa">BAIXA</span>
            <div class="rule-code">R10</div>
            <div class="rule-title">KM realizado abaixo do esperado</div>
            <div class="rule-desc">Distância percorrida significativamente inferior ao KM esperado calculado pelo histórico de consumo (divergência &gt; 25%).</div>
            <div class="rule-example">km_esperado = litros × μ → desvio &gt; 25%</div>
        </div>

    </div>
    <div class="example example-info" style="margin-top:14px">
        <strong>Combinação de regras</strong>
        As regras podem se acumular em um mesmo registro. Quanto mais regras acionadas, maior a gravidade final. Exemplo: R02 + R05 acionadas juntas → classificação ALTA automática.
    </div>
</div>

<!-- SEÇÃO 5 -->
<div class="section" id="s5">
    <div class="section-title">
        <div class="section-num c-indigo">5</div>
        <span class="icon">🤖</span>
        <h2>Arquitetura de Agentes de IA</h2>
    </div>
    <p>O sistema não é um programa único — é uma <strong>equipe de agentes especializados</strong> que trabalham em sequência, cada um com responsabilidade exclusiva sobre uma etapa da auditoria:</p>
    <div class="agent-pipeline">
        <div class="agent-step">📥 Ingestion</div><div class="agent-arrow">→</div>
        <div class="agent-step">🔍 Validação</div><div class="agent-arrow">→</div>
        <div class="agent-step">💾 Storage</div><div class="agent-arrow">→</div>
        <div class="agent-step">📏 Regras</div><div class="agent-arrow">→</div>
        <div class="agent-step">📈 Histórico</div><div class="agent-arrow">→</div>
        <div class="agent-step">⚖️ Classificação</div><div class="agent-arrow">→</div>
        <div class="agent-step">📋 Relatório</div><div class="agent-arrow">→</div>
        <div class="agent-step">📬 Notificação</div>
    </div>
    <div class="agent-cards">
        <div class="agent-card">
            <div class="agent-icon">📥</div>
            <div class="agent-name">AgentIngestion</div>
            <div class="agent-role">Receptor de Dados</div>
            <ul><li>Lê a planilha Excel/CSV</li><li>Padroniza nomes de colunas</li><li>Converte datas e números</li><li>Calcula consumo quando ausente</li></ul>
        </div>
        <div class="agent-card">
            <div class="agent-icon">🔍</div>
            <div class="agent-name">AgentValidacao</div>
            <div class="agent-role">Filtro de Qualidade</div>
            <ul><li>Remove registros inválidos</li><li>Identifica duplicatas</li><li>Garante coerência dos números</li></ul>
        </div>
        <div class="agent-card">
            <div class="agent-icon">💾</div>
            <div class="agent-name">AgentStorageSQLite</div>
            <div class="agent-role">Memória do Sistema</div>
            <ul><li>Grava dados no banco local</li><li>Mantém histórico por veículo</li><li>Permite análise retroativa</li></ul>
        </div>
        <div class="agent-card">
            <div class="agent-icon">📏</div>
            <div class="agent-name">AgentRegras</div>
            <div class="agent-role">Auditor Técnico</div>
            <ul><li>Aplica regras R01 a R10</li><li>Gera evidência por ocorrência</li><li>Responde: "esse registro faz sentido?"</li></ul>
        </div>
        <div class="agent-card">
            <div class="agent-icon">📈</div>
            <div class="agent-name">AgentHistorico</div>
            <div class="agent-role">Analista Estatístico</div>
            <ul><li>Calcula μ e σ por veículo</li><li>Detecta outliers estatísticos</li><li>Responde: "é normal para esse carro?"</li></ul>
        </div>
        <div class="agent-card">
            <div class="agent-icon">⚖️</div>
            <div class="agent-name">AgentClassificacao</div>
            <div class="agent-role">Supervisor de Gravidade</div>
            <ul><li>Agrega todas as regras acionadas</li><li>Define gravidade final</li><li>Combina evidências múltiplas</li></ul>
        </div>
        <div class="agent-card">
            <div class="agent-icon">📋</div>
            <div class="agent-name">AgentRelatorio</div>
            <div class="agent-role">Redator Técnico</div>
            <ul><li>Gera relatório consolidado</li><li>Organiza dados por gravidade</li><li>Produz resumo executivo</li></ul>
        </div>
        <div class="agent-card">
            <div class="agent-icon">📬</div>
            <div class="agent-name">AgentNotificacao</div>
            <div class="agent-role">Redator Administrativo</div>
            <ul><li>Gera minutas de notificação</li><li>Agrupa ocorrências por condutor</li><li>Prazo padrão: 48 horas</li></ul>
        </div>
    </div>
    <div class="example example-info" style="margin-top:14px">
        <strong>Comparação com uma repartição pública</strong>
        👉 Um <em>recebe</em> → um <em>confere</em> → um <em>armazena</em> → um <em>analisa</em> → um <em>compara</em> → um <em>decide</em> → um <em>documenta</em> → um <em>notifica</em>.
    </div>
</div>

<!-- SEÇÃO 6 -->
<div class="section" id="s6">
    <div class="section-title">
        <div class="section-num c-orange">6</div>
        <span class="icon">🚨</span>
        <h2>Classificação de Gravidade</h2>
    </div>
    <p>Cada ocorrência recebe uma classificação com base nas regras acionadas e na intensidade do desvio detectado.</p>
    <table class="grav-table">
        <thead><tr><th>Nível</th><th>Critério</th><th>Ação esperada</th></tr></thead>
        <tbody>
            <tr>
                <td><span class="chip-alta">🔴 ALTA</span></td>
                <td>Regra crítica acionada (R01, R03, R04 ou R05), ou combinação de múltiplas regras, ou desvio estatístico severo.</td>
                <td>Notificação administrativa imediata. Prazo de 48 horas para justificativa formal.</td>
            </tr>
            <tr>
                <td><span class="chip-media">🟡 MÉDIA</span></td>
                <td>Regras secundárias (R02, R06, R07, R08, R09) ou desvio moderado sem combinação crítica.</td>
                <td>Monitoramento preventivo ou orientação ao condutor/unidade responsável.</td>
            </tr>
            <tr>
                <td><span class="chip-baixa">🟢 BAIXA</span></td>
                <td>Desvio pequeno dentro da margem aceitável. Geralmente R10 isolado com divergência entre 25–35%.</td>
                <td>Registro para histórico. Sem intervenção imediata necessária.</td>
            </tr>
        </tbody>
    </table>
    <div class="example example-alert" style="margin-top:12px">
        <strong>Princípio da combinação de regras</strong>
        Múltiplos alertas em um mesmo registro <strong>elevam a gravidade final automaticamente</strong>. Exemplo: R02 (consumo fora da faixa) + R05 (lógica incoerente) → resultado final ALTA.
    </div>
</div>

<!-- SEÇÃO 7 -->
<div class="section" id="s7">
    <div class="section-title">
        <div class="section-num c-green">7</div>
        <span class="icon">🧾</span>
        <h2>Ação da Equipe</h2>
    </div>
    <div class="two-col">
        <div>
            <p style="font-weight:600;color:#16a34a;margin-bottom:8px">✔️ O que fazer</p>
            <ul class="checklist">
                <li>Identificar a ocorrência na lista de alertas</li>
                <li>Ler a evidência técnica gerada pelo sistema</li>
                <li>Classificar o nível de prioridade</li>
                <li>Enviar a notificação ao condutor ou unidade</li>
                <li>Aguardar a justificativa no prazo de 48h</li>
                <li>Tomar decisão com base na defesa apresentada</li>
                <li>Registrar o resultado no sistema</li>
            </ul>
        </div>
        <div>
            <p style="font-weight:600;color:#dc2626;margin-bottom:8px">✗ O que evitar</p>
            <ul class="checklist dont">
                <li>Acusar o condutor sem análise prévia</li>
                <li>Ignorar o contexto operacional do veículo</li>
                <li>Decidir com base em opinião, sem evidência</li>
                <li>Descartar ocorrências sem justificativa formal</li>
            </ul>
        </div>
    </div>
    <div class="quote" style="margin-top:14px">💬 &ldquo;Isso faz sentido na vida real? Sempre pergunte isso antes de agir.&rdquo;</div>
    <div class="example example-info" style="margin-top:12px">
        <strong>Fluxo administrativo resumido</strong>
        Identificar → Validar → Classificar → Notificar (48h) → Analisar justificativa → Decidir → Registrar
    </div>
</div>

<!-- SEÇÃO 8 -->
<div class="section" id="s8">
    <div class="section-title">
        <div class="section-num c-yellow">8</div>
        <span class="icon">⚠️</span>
        <h2>Limites do Sistema</h2>
    </div>
    <div class="compare">
        <div class="compare-col before">
            <h4>O sistema NÃO faz</h4>
            <ul>
                <li>Acusa condutores diretamente</li>
                <li>Substitui a análise e o julgamento humano</li>
                <li>Considera fatores externos automaticamente</li>
                <li>Garante certeza absoluta — gera hipóteses técnicas</li>
                <li>Funciona como prova jurídica isolada</li>
            </ul>
        </div>
        <div class="compare-col after">
            <h4>O sistema FAZ</h4>
            <ul>
                <li>Detecta padrões estatisticamente anômalos</li>
                <li>Aplica regras objetivas e padronizadas</li>
                <li>Gera evidência numérica rastreável</li>
                <li>Produz notificações formais padronizadas</li>
                <li>Mantém histórico auditável por veículo</li>
            </ul>
        </div>
    </div>
    <div class="example example-alert" style="margin-top:14px">
        <strong>📌 Fatores externos a considerar antes de concluir</strong>
        Condições de trânsito intenso &nbsp;|&nbsp; Carga pesada ou transporte especial &nbsp;|&nbsp; Manutenção em andamento &nbsp;|&nbsp; Uso operacional atípico do veículo &nbsp;|&nbsp; Falha de hodômetro
    </div>
</div>

<!-- SEÇÃO 9 -->
<div class="section" id="s9">
    <div class="section-title">
        <div class="section-num c-teal">9</div>
        <span class="icon">🎓</span>
        <h2>Treinamento Prático</h2>
    </div>
    <p>Use os exercícios abaixo para treinar a equipe na leitura e interpretação das ocorrências geradas pelo sistema.</p>

    <div class="exercise">
        <div class="ex-header"><div class="ex-num">1</div><div class="ex-title">Verificação de consumo básico</div></div>
        <div class="ex-data">
            <div class="ex-data-item"><span class="label">KM Rodado</span><span class="value">100</span></div>
            <div class="ex-data-item"><span class="label">Litros</span><span class="value">10</span></div>
            <div class="ex-data-item"><span class="label">Consumo</span><span class="value">10 km/L</span></div>
        </div>
        <div class="ex-result ex-ok">✅ Dentro do padrão — nenhuma regra acionada.</div>
    </div>

    <div class="exercise">
        <div class="ex-header"><div class="ex-num">2</div><div class="ex-title">Consumo muito baixo</div></div>
        <div class="ex-data">
            <div class="ex-data-item"><span class="label">KM Rodado</span><span class="value">100</span></div>
            <div class="ex-data-item"><span class="label">Litros</span><span class="value">20</span></div>
            <div class="ex-data-item"><span class="label">Consumo</span><span class="value">5 km/L</span></div>
        </div>
        <div class="ex-result ex-nok">🚨 Consumo abaixo do limite mínimo — aciona R02 e possivelmente R03 (outlier estatístico).</div>
    </div>

    <div class="exercise">
        <div class="ex-header"><div class="ex-num">3</div><div class="ex-title">Hodômetro regressivo</div></div>
        <div class="ex-data">
            <div class="ex-data-item"><span class="label">KM Anterior</span><span class="value">10.000</span></div>
            <div class="ex-data-item"><span class="label">KM Atual</span><span class="value">9.900</span></div>
        </div>
        <div class="ex-result ex-nok">🚨 KM atual menor que o anterior — aciona R04 (hodômetro inconsistente). Gravidade ALTA.</div>
    </div>

    <div class="exercise">
        <div class="ex-header"><div class="ex-num">4</div><div class="ex-title">KM muito abaixo do esperado</div></div>
        <div class="ex-data">
            <div class="ex-data-item"><span class="label">Litros</span><span class="value">40</span></div>
            <div class="ex-data-item"><span class="label">Média (μ)</span><span class="value">10 km/L</span></div>
            <div class="ex-data-item"><span class="label">KM Esperado</span><span class="value">400 km</span></div>
            <div class="ex-data-item"><span class="label">KM Realizado</span><span class="value">150 km</span></div>
        </div>
        <div class="ex-result ex-nok">🚨 Diferença de 250 km (62,5%) — aciona R05 (lógica incoerente) e R10 (KM abaixo do esperado). Gravidade ALTA.</div>
    </div>

    <div class="exercise">
        <div class="ex-header"><div class="ex-num">5</div><div class="ex-title">Dois abastecimentos no mesmo dia</div></div>
        <div class="ex-data">
            <div class="ex-data-item"><span class="label">Placa</span><span class="value">ABC-1234</span></div>
            <div class="ex-data-item"><span class="label">1º Abastecimento</span><span class="value">14/04 08:15</span></div>
            <div class="ex-data-item"><span class="label">2º Abastecimento</span><span class="value">14/04 11:20</span></div>
        </div>
        <div class="ex-result ex-nok">🚨 Intervalo muito curto — aciona R06 (abastecimentos próximos). Verificar se o tanque comportava os dois volumes e se o KM rodado é coerente.</div>
    </div>
</div>

<!-- SEÇÃO 10 -->
<div class="section" id="s10">
    <div class="section-title">
        <div class="section-num c-slate">10</div>
        <span class="icon">🎯</span>
        <h2>Mentalidade Final</h2>
    </div>
    <div class="quote">💬 &ldquo;O sistema não diz quem errou. Ele mostra onde os números não fazem sentido.&rdquo;</div>
    <p style="margin-top:16px">Com esse sistema, a gestão de frota passa a contar com:</p>
    <div class="compare" style="margin-top:12px">
        <div class="compare-col before">
            <h4>❌ Antes</h4>
            <ul>
                <li>Análise manual e demorada</li>
                <li>Critérios subjetivos</li>
                <li>Sem rastreabilidade por veículo</li>
                <li>Difícil de padronizar e replicar</li>
            </ul>
        </div>
        <div class="compare-col after">
            <h4>✅ Depois</h4>
            <ul>
                <li>Auditoria automática e sistemática</li>
                <li>Critérios objetivos e documentados</li>
                <li>Histórico rastreável por veículo</li>
                <li>Notificações padronizadas e auditáveis</li>
            </ul>
        </div>
    </div>
    <div class="quote" style="margin-top:16px;background:linear-gradient(135deg,#f0fdf4,#ecfdf5);border-left-color:#16a34a;color:#166534">
        🚀 &ldquo;Cada agente resolve uma parte do problema — juntos, eles garantem a auditoria completa.&rdquo;
    </div>
    <div class="info-box" style="margin-top:16px;background:#eff6ff;border-color:#bfdbfe">
        <strong>💡 Princípio fundamental:</strong><br>
        Toda inconsistência deve ser comprovada por números. A defesa do condutor também deve ser baseada em fatos documentados. O sistema garante tratamento igual e rastreável para todos os registros analisados.
    </div>
</div>

</div>"""

    st_components.html(HTML_MANUAL, height=4800, scrolling=True)

    if manual_data and manual_data.get('bytes'):
        st.download_button(
            label=f"Baixar original ({manual_data.get('arquivo', 'manual.docx')})",
            data=manual_data.get('bytes', b''),
            file_name=manual_data.get('arquivo', 'manual.docx'),
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            use_container_width=True,
            key='download_manual_docx_dashboard',
        )


def apply_filters(df, filtros, fuel_map):
    # Evita copia completa da base em toda iteracao dos filtros.
    base_df = df

    ano_filter = filtros.get('ano', 'Todos')
    if ano_filter and ano_filter != 'Todos' and 'ano' in base_df.columns:
        base_df = base_df[base_df['ano'].astype(str) == str(ano_filter)]

    mes_filter = filtros.get('mes', 'Todos')
    if mes_filter and mes_filter != 'Todos' and 'mes' in base_df.columns:
        _MONTHS_LOCAL = {
            1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
            5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
            9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
        }
        _num = {v: k for k, v in _MONTHS_LOCAL.items()}.get(mes_filter)
        if _num is not None:
            base_df = base_df[base_df['mes'].astype(int) == _num]

    fuel_filter = filtros.get('fuel', 'Todos')
    if 'combustivel_norm' in base_df.columns and fuel_filter in fuel_map:
        base_df = base_df[base_df['combustivel_norm'] == fuel_map[fuel_filter]]

    marca_filter = filtros.get('marca', 'Todos')
    if marca_filter != 'Todos' and 'Marca' in base_df.columns:
        base_df = base_df[base_df['Marca'].astype(str) == marca_filter]

    modelo_filter = filtros.get('modelo', 'Todos')
    if isinstance(modelo_filter, list):
        modelo_filter = modelo_filter[0] if modelo_filter else 'Todos'
    if modelo_filter != 'Todos':
        alvo_modelo = canonicalizar_modelo(modelo_filter)
        if 'modelo_norm' in base_df.columns:
            base_df = base_df[base_df['modelo_norm'] == alvo_modelo]
        else:
            base_df = base_df[
                base_df['Modelo'].astype(str).apply(canonicalizar_modelo) == alvo_modelo
            ]

    condutor_filter = filtros.get('condutor', 'Todos')
    if condutor_filter != 'Todos':
        base_df = base_df[base_df['Condutor'] == condutor_filter]

    unidade_filter = filtros.get('unidade', 'Todas')
    if unidade_filter and unidade_filter != 'Todas' and 'unidade_alerta' in base_df.columns:
        base_df = base_df[base_df['unidade_alerta'].astype(str) == unidade_filter]

    placa_filter = filtros.get('placa', 'Todos')
    if placa_filter != 'Todos':
        base_df = base_df[base_df['Placa'] == placa_filter]

    modelo_filter_list = [] if modelo_filter == 'Todos' else [modelo_filter]
    return base_df, modelo_filter_list


@st.cache_data(show_spinner=False, ttl=300, max_entries=128)
def apply_filters_cached(_df, filtros_tuple, fuel_map_tuple, data_version):
    filtros = dict(filtros_tuple)
    fuel_map = dict(fuel_map_tuple)
    return apply_filters(_df, filtros, fuel_map)


def filtrar_resultado_auditoria_por_recorte(filtered_df: pd.DataFrame, resultado_auditoria: dict) -> dict:
    """Filtra ocorrências/notificações para manter estrita coerência com o recorte atual."""
    if not resultado_auditoria or filtered_df is None or filtered_df.empty:
        return resultado_auditoria

    def _to_dt(value):
        try:
            return pd.to_datetime(value, errors='coerce')
        except Exception:
            return pd.NaT

    base = filtered_df.copy()
    base['__placa_key__'] = base.get('Placa', '').astype(str).str.strip()
    base['__data_key__'] = pd.to_datetime(base.get('data', pd.NaT), errors='coerce').dt.strftime('%Y-%m-%d %H:%M')
    chaves_validas = set(
        (base['__placa_key__'] + '|' + base['__data_key__'])
        .dropna()
        .astype(str)
        .tolist()
    )

    def _chave_evento(placa, data_hora):
        placa_s = '' if placa is None else str(placa).strip()
        dt = _to_dt(data_hora)
        if pd.isna(dt):
            return None
        return f"{placa_s}|{dt.strftime('%Y-%m-%d %H:%M')}"

    ocs = resultado_auditoria.get('ocorrencias', []) or []
    ocs_filtradas = []
    for oc in ocs:
        chave = _chave_evento(oc.get('placa'), oc.get('data_hora'))
        if chave and chave in chaves_validas:
            ocs_filtradas.append(oc)

    notifs = resultado_auditoria.get('notificacoes', []) or []
    notifs_filtradas = []
    for notif in notifs:
        ocorrencias_notif = notif.get('ocorrencias', []) or []
        ocorrencias_validas = []
        for oc in ocorrencias_notif:
            chave_oc = _chave_evento(oc.get('placa'), oc.get('data_hora'))
            if chave_oc and chave_oc in chaves_validas:
                ocorrencias_validas.append(oc)

        chave_notif = _chave_evento(notif.get('placa'), notif.get('data_hora'))
        notif_valida = bool(chave_notif and chave_notif in chaves_validas)

        if ocorrencias_validas or notif_valida:
            notif_copy = dict(notif)
            if ocorrencias_notif:
                notif_copy['ocorrencias'] = ocorrencias_validas
            notifs_filtradas.append(notif_copy)

    resultado_filtrado = dict(resultado_auditoria)
    resultado_filtrado['ocorrencias'] = ocs_filtradas
    resultado_filtrado['notificacoes'] = notifs_filtradas
    return resultado_filtrado


def build_excel_report(filtered_df, freq_suspeita, sequenciais):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        resumo = pd.DataFrame([
            {'Indicador': 'Total de abastecimentos', 'Valor': len(filtered_df)},
            {'Indicador': 'Total de placas', 'Valor': filtered_df['Placa'].nunique()},
            {'Indicador': 'Total de condutores', 'Valor': filtered_df['Condutor'].nunique()},
            {'Indicador': 'Litros totais', 'Valor': float(filtered_df['litros'].sum())},
            {'Indicador': 'KM total', 'Valor': float(filtered_df['km'].sum())},
            {'Indicador': 'Consumo medio KM/L', 'Valor': float(filtered_df['consumo'].mean()) if len(filtered_df) else 0},
            {'Indicador': 'Multiplos no mesmo dia', 'Valor': len(freq_suspeita)},
            {'Indicador': 'Sequenciais < 1h', 'Valor': len(sequenciais)},
        ])
        resumo.to_excel(writer, sheet_name='Resumo', index=False)
        filtered_df.sort_values('data').to_excel(writer, sheet_name='Base Filtrada', index=False)
        if not sequenciais.empty:
            sequenciais.sort_values('data').to_excel(writer, sheet_name='Seq < 1h', index=False)

    output.seek(0)
    return output.getvalue()


def build_auditoria_analitica_excel(filtered_df, resultado_auditoria, filtros_ui, freq_suspeita, sequenciais, sigma_mult=2.0):
    output = BytesIO()

    ocs = resultado_auditoria.get('ocorrencias', []) if resultado_auditoria else []
    ocs_df = pd.DataFrame(ocs)

    consumo_base = pd.to_numeric(filtered_df.get('consumo', np.nan), errors='coerce')
    media_consumo = float(consumo_base.mean()) if len(filtered_df) else 0.0
    desvio_consumo = float(consumo_base.std()) if len(filtered_df) else 0.0
    limiar_outlier = media_consumo - (float(sigma_mult) * desvio_consumo)

    resumo_exec = (resultado_auditoria or {}).get('relatorio', {}).get('resumo_executivo', {})
    filtro_placa = (filtros_ui or {}).get('placa', 'Todos')
    todos_veiculos = str(filtro_placa) == 'Todos'

    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        criterios = pd.DataFrame([
            {'Parâmetro': 'Combustível', 'Valor': (filtros_ui or {}).get('fuel', 'Todos')},
            {'Parâmetro': 'Marca', 'Valor': (filtros_ui or {}).get('marca', 'Todos')},
            {'Parâmetro': 'Modelo', 'Valor': (filtros_ui or {}).get('modelo', 'Todos')},
            {'Parâmetro': 'Condutor', 'Valor': (filtros_ui or {}).get('condutor', 'Todos')},
            {'Parâmetro': 'Placa', 'Valor': filtro_placa},
            {'Parâmetro': 'Período', 'Valor': str((filtros_ui or {}).get('periodo', 'Todos'))},
            {'Parâmetro': 'Sigma Outlier', 'Valor': float(sigma_mult)},
        ])
        criterios.to_excel(writer, sheet_name='Criterios_Filtro', index=False)

        resumo = pd.DataFrame([
            {'Indicador': 'Registros auditados', 'Valor': len(filtered_df)},
            {'Indicador': 'Placas únicas', 'Valor': filtered_df['Placa'].nunique() if 'Placa' in filtered_df.columns else 0},
            {'Indicador': 'Modelos únicos', 'Valor': filtered_df['Modelo'].nunique() if 'Modelo' in filtered_df.columns else 0},
            {'Indicador': 'Condutores únicos', 'Valor': filtered_df['Condutor'].nunique() if 'Condutor' in filtered_df.columns else 0},
            {'Indicador': 'Litros totais', 'Valor': float(pd.to_numeric(filtered_df.get('litros', 0), errors='coerce').sum())},
            {'Indicador': 'KM rodados totais', 'Valor': float(pd.to_numeric(filtered_df.get('km', 0), errors='coerce').sum())},
            {'Indicador': 'Consumo médio (KM/L)', 'Valor': media_consumo},
            {'Indicador': 'Desvio padrão consumo', 'Valor': desvio_consumo},
            {'Indicador': 'Limiar outlier (média - Nσ)', 'Valor': limiar_outlier},
            {'Indicador': 'Múltiplos no mesmo dia', 'Valor': int(len(freq_suspeita))},
            {'Indicador': 'Sequenciais < 1h', 'Valor': int(len(sequenciais))},
            {'Indicador': 'Ocorrências ALTA', 'Valor': int(resumo_exec.get('ocorrencias_alta', 0))},
            {'Indicador': 'Ocorrências MÉDIA', 'Valor': int(resumo_exec.get('ocorrencias_media', 0))},
            {'Indicador': 'Ocorrências BAIXA', 'Valor': int(resumo_exec.get('ocorrencias_baixa', 0))},
        ])
        resumo.to_excel(writer, sheet_name='Resumo_Auditoria', index=False)

        if todos_veiculos and 'Modelo' in filtered_df.columns:
            base_modelo = filtered_df.copy()
            base_modelo['litros'] = pd.to_numeric(base_modelo.get('litros', np.nan), errors='coerce')
            base_modelo['km'] = pd.to_numeric(base_modelo.get('km', np.nan), errors='coerce')
            base_modelo['consumo'] = pd.to_numeric(base_modelo.get('consumo', np.nan), errors='coerce')
            base_modelo['outlier_consumo'] = base_modelo['consumo'] < limiar_outlier

            analise_modelo = (
                base_modelo.groupby('Modelo', dropna=False)
                .agg(
                    Abastecimentos=('Modelo', 'size'),
                    Placas=('Placa', 'nunique'),
                    Condutores=('Condutor', 'nunique'),
                    Litros_Totais=('litros', 'sum'),
                    KM_Rodados_Totais=('km', 'sum'),
                    Consumo_Medio_KML=('consumo', 'mean'),
                    Consumo_Min_KML=('consumo', 'min'),
                    Consumo_Max_KML=('consumo', 'max'),
                    Outliers_Consumo=('outlier_consumo', 'sum'),
                )
                .reset_index()
                .sort_values('Abastecimentos', ascending=False)
            )
            analise_modelo.to_excel(writer, sheet_name='Analise_por_Modelo', index=False)

        if not ocs_df.empty:
            cols_pref = [c for c in [
                'placa', 'condutor', 'modelo', 'unidade', 'data_hora',
                'codigo_regra', 'tipo_ocorrencia', 'gravidade_final', 'evidencia',
                'km_anterior', 'km_atual', 'km_rodados', 'km_esperado',
                'km_l', 'litros', 'media', 'min', 'max', 'desvio',
                'estabelecimento', 'descricao_tecnica', 'qtd_evidencias_evento', 'recomendacao',
            ] if c in ocs_df.columns]
            ocs_export = ocs_df[cols_pref].copy() if cols_pref else ocs_df.copy()
            ocs_export.to_excel(writer, sheet_name='Ocorrencias_Detalhadas', index=False)

        base_cols = [c for c in [
            'data', 'Placa', 'Condutor', 'Marca', 'Modelo', 'posto', 'Produto',
            'ult_km_alerta', 'km_atual_alerta', 'km', 'litros', 'consumo',
        ] if c in filtered_df.columns]
        filtered_df.sort_values('data')[base_cols].to_excel(writer, sheet_name='Base_Filtrada', index=False)

        if not sequenciais.empty:
            sequenciais.sort_values('data')[
                [c for c in ['Placa', 'data', 'posto', 'litros', 'km', 'consumo', 'delta_min'] if c in sequenciais.columns]
            ].to_excel(writer, sheet_name='Alertas_Sequenciais', index=False)

        if not freq_suspeita.empty:
            freq_suspeita.sort_values(['Placa', 'data_dia']).to_excel(writer, sheet_name='Multiplos_Mesmo_Dia', index=False)

    output.seek(0)
    return output.getvalue()


def build_auditoria_analitica_pdf(filtered_df, resultado_auditoria, filtros_ui, freq_suspeita, sequenciais, sigma_mult=2.0):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=18,
        rightMargin=18,
        topMargin=18,
        bottomMargin=18,
    )
    styles = getSampleStyleSheet()
    story = []

    header_style = ParagraphStyle(
        name='audit_pdf_header',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=8,
        textColor=colors.white,
    )
    body_style = ParagraphStyle(
        name='audit_pdf_body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=6.2,
        leading=7.2,
        wordWrap='CJK',
    )
    small_style = ParagraphStyle(
        name='audit_pdf_small',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10,
    )
    section_title_style = ParagraphStyle(
        name='audit_pdf_section_title',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=14,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=6,
    )
    model_title_style = ParagraphStyle(
        name='audit_pdf_model_title',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=15,
        textColor=colors.HexColor('#0f766e'),
        spaceAfter=6,
    )
    note_style = ParagraphStyle(
        name='audit_pdf_note',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7.5,
        leading=9,
        textColor=colors.HexColor('#334155'),
    )

    def _safe_paragraph(value, style):
        text = str(value if pd.notna(value) else '-')
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        return Paragraph(text, style)

    def _format_num(value, decimals=2):
        if pd.isna(value):
            return '-'
        try:
            return f"{float(value):.{decimals}f}"
        except Exception:
            return str(value)

    ocs = resultado_auditoria.get('ocorrencias', []) if resultado_auditoria else []
    ocs_df = pd.DataFrame(ocs)
    resumo_exec = (resultado_auditoria or {}).get('relatorio', {}).get('resumo_executivo', {})
    filtro_placa = (filtros_ui or {}).get('placa', 'Todos')
    todos_veiculos = str(filtro_placa) == 'Todos'

    consumo_base = pd.to_numeric(filtered_df.get('consumo', np.nan), errors='coerce')
    media_consumo = float(consumo_base.mean()) if len(filtered_df) else 0.0
    desvio_consumo = float(consumo_base.std()) if len(filtered_df) else 0.0
    limiar_outlier = media_consumo - (float(sigma_mult) * desvio_consumo)

    story.append(Paragraph("<b>Relatório Analítico de Auditoria</b>", styles['Title']))
    story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", small_style))
    story.append(Spacer(1, 8))

    criterios = [
        f"Combustível: {(filtros_ui or {}).get('fuel', 'Todos')}",
        f"Marca: {(filtros_ui or {}).get('marca', 'Todos')}",
        f"Modelo: {(filtros_ui or {}).get('modelo', 'Todos')}",
        f"Condutor: {(filtros_ui or {}).get('condutor', 'Todos')}",
        f"Placa: {filtro_placa}",
        f"Período: {(filtros_ui or {}).get('periodo', 'Todos')}",
        f"Sigma Outlier: {float(sigma_mult):.1f}",
    ]
    story.append(Paragraph("Critérios do Filtro Aplicado", section_title_style))
    for item in criterios:
        story.append(Paragraph(item, small_style))
    story.append(Spacer(1, 8))

    resumo_data = [
        ['Indicador', 'Valor'],
        ['Registros auditados', str(len(filtered_df))],
        ['Placas únicas', str(filtered_df['Placa'].nunique() if 'Placa' in filtered_df.columns else 0)],
        ['Modelos únicos', str(filtered_df['Modelo'].nunique() if 'Modelo' in filtered_df.columns else 0)],
        ['Condutores únicos', str(filtered_df['Condutor'].nunique() if 'Condutor' in filtered_df.columns else 0)],
        ['Litros totais', f"{float(pd.to_numeric(filtered_df.get('litros', 0), errors='coerce').sum()):.2f}"],
        ['KM rodados totais', f"{float(pd.to_numeric(filtered_df.get('km', 0), errors='coerce').sum()):.2f}"],
        ['Consumo médio (KM/L)', f"{media_consumo:.2f}"],
        ['Desvio padrão consumo', f"{desvio_consumo:.2f}"],
        ['Limiar outlier', f"{limiar_outlier:.2f}"],
        ['Múltiplos no mesmo dia', str(int(len(freq_suspeita)))],
        ['Sequenciais < 1h', str(int(len(sequenciais)))],
        ['Ocorrências ALTA', str(int(resumo_exec.get('ocorrencias_alta', 0)))],
        ['Ocorrências MÉDIA', str(int(resumo_exec.get('ocorrencias_media', 0)))],
        ['Ocorrências BAIXA', str(int(resumo_exec.get('ocorrencias_baixa', 0)))],
    ]
    resumo_table = Table(resumo_data, repeatRows=1, colWidths=[170, 110])
    resumo_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('LEADING', (0, 0), (-1, -1), 10),
    ]))
    story.append(Paragraph("Resumo da Auditoria", section_title_style))
    story.append(resumo_table)
    story.append(Spacer(1, 10))

    base_modelo = filtered_df.copy()
    if 'Modelo' not in base_modelo.columns:
        base_modelo['Modelo'] = 'Não informado'
    base_modelo['Modelo'] = base_modelo['Modelo'].fillna('Não informado').astype(str)
    base_modelo['litros'] = pd.to_numeric(base_modelo.get('litros', np.nan), errors='coerce')
    base_modelo['km'] = pd.to_numeric(base_modelo.get('km', np.nan), errors='coerce')
    base_modelo['consumo'] = pd.to_numeric(base_modelo.get('consumo', np.nan), errors='coerce')
    base_modelo['outlier_consumo'] = base_modelo['consumo'] < limiar_outlier

    if ocs_df.empty:
        story.append(Paragraph("Nenhuma ocorrência detalhada encontrada para os filtros atuais.", small_style))
    else:
        if 'modelo' not in ocs_df.columns:
            ocs_df['modelo'] = 'Não informado'
        ocs_df['modelo'] = ocs_df['modelo'].fillna('Não informado').astype(str)

        if 'placa' not in ocs_df.columns:
            ocs_df['placa'] = '-'
        if 'data_hora' not in ocs_df.columns:
            ocs_df['data_hora'] = '-'

        # Painel executivo por modelo para priorização.
        ranking_modelo = (
            ocs_df.groupby('modelo', dropna=False)
            .agg(
                Ocorrencias=('modelo', 'size'),
                Alta=('gravidade_final', lambda s: int((s == 'ALTA').sum())),
                Media=('gravidade_final', lambda s: int((s == 'MEDIA').sum())),
                Baixa=('gravidade_final', lambda s: int((s == 'BAIXA').sum())),
                Placas=('placa', 'nunique'),
            )
            .reset_index()
            .rename(columns={'modelo': 'Modelo'})
            .sort_values(['Alta', 'Ocorrencias'], ascending=[False, False])
        )

        if todos_veiculos and not ranking_modelo.empty:
            story.append(Paragraph("Ranking de Modelos Prioritários", section_title_style))
            ranking_rows = [[
                _safe_paragraph('Modelo', header_style),
                _safe_paragraph('Ocorrências', header_style),
                _safe_paragraph('ALTA', header_style),
                _safe_paragraph('MÉDIA', header_style),
                _safe_paragraph('BAIXA', header_style),
                _safe_paragraph('Placas', header_style),
            ]]
            for _, row in ranking_modelo.head(10).iterrows():
                ranking_rows.append([
                    _safe_paragraph(row['Modelo'], body_style),
                    _safe_paragraph(row['Ocorrencias'], body_style),
                    _safe_paragraph(row['Alta'], body_style),
                    _safe_paragraph(row['Media'], body_style),
                    _safe_paragraph(row['Baixa'], body_style),
                    _safe_paragraph(row['Placas'], body_style),
                ])
            ranking_table = Table(ranking_rows, repeatRows=1, colWidths=[190, 72, 50, 55, 55, 55])
            ranking_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#334155')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.whitesmoke]),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(ranking_table)
            story.append(Spacer(1, 10))

        modelos_ordenados = (
            base_modelo.groupby('Modelo', dropna=False)
            .size()
            .sort_values(ascending=False)
            .index.tolist()
        )
        modelos_com_ocorrencia = [m for m in modelos_ordenados if m in set(ocs_df['modelo'])]
        if not modelos_com_ocorrencia:
            modelos_com_ocorrencia = sorted(ocs_df['modelo'].unique().tolist())

        story.append(Paragraph("Ocorrências Auditadas por Modelo", section_title_style))
        story.append(Spacer(1, 4))

        for idx, modelo in enumerate(modelos_com_ocorrencia):
            base_m = base_modelo[base_modelo['Modelo'] == modelo].copy()
            ocs_m = ocs_df[ocs_df['modelo'] == modelo].copy()
            if ocs_m.empty:
                continue

            altas = int((ocs_m.get('gravidade_final', pd.Series(dtype='object')) == 'ALTA').sum())
            medias = int((ocs_m.get('gravidade_final', pd.Series(dtype='object')) == 'MEDIA').sum())
            baixas = int((ocs_m.get('gravidade_final', pd.Series(dtype='object')) == 'BAIXA').sum())

            story.append(Paragraph(f"Modelo: {modelo}", model_title_style))

            header_modelo = [
                ['Indicador', 'Valor', 'Indicador', 'Valor'],
                ['Abastecimentos', str(len(base_m)), 'Placas', str(base_m['Placa'].nunique() if 'Placa' in base_m.columns else 0)],
                ['Condutores', str(base_m['Condutor'].nunique() if 'Condutor' in base_m.columns else 0), 'Litros totais', _format_num(base_m['litros'].sum())],
                ['KM rodados', _format_num(base_m['km'].sum()), 'Consumo médio (KM/L)', _format_num(base_m['consumo'].mean())],
                ['Outliers de consumo', str(int(base_m['outlier_consumo'].sum())), 'Ocorrências', str(len(ocs_m))],
                ['Gravidade ALTA', str(altas), 'Gravidade MÉDIA/BAIXA', f"{medias}/{baixas}"],
            ]
            table_header_modelo = Table(header_modelo, repeatRows=1, colWidths=[120, 70, 120, 70])
            table_header_modelo.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f766e')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('LEADING', (0, 0), (-1, -1), 10),
            ]))
            story.append(table_header_modelo)
            story.append(Spacer(1, 6))

            story.append(Paragraph(
                "Tabela-resumo das ocorrências auditadas. Evidência e recomendação completas seguem logo abaixo.",
                note_style,
            ))
            story.append(Spacer(1, 4))

            cols_m = [c for c in [
                'placa', 'data_hora', 'codigo_regra', 'tipo_ocorrencia', 'gravidade_final',
            ] if c in ocs_m.columns]
            ocs_view = ocs_m[cols_m].copy().fillna('-')

            col_labels = {
                'placa': 'Placa',
                'data_hora': 'Data/Hora',
                'codigo_regra': 'Código',
                'tipo_ocorrencia': 'Tipo',
                'gravidade_final': 'Gravidade',
            }
            table_rows = [[_safe_paragraph(col_labels.get(c, c), header_style) for c in ocs_view.columns]]
            for _, row in ocs_view.iterrows():
                table_rows.append([_safe_paragraph(row[c], body_style) for c in ocs_view.columns])

            col_widths = []
            for c in ocs_view.columns:
                if c == 'placa':
                    col_widths.append(48)
                elif c == 'data_hora':
                    col_widths.append(62)
                elif c == 'codigo_regra':
                    col_widths.append(42)
                elif c == 'tipo_ocorrencia':
                    col_widths.append(240)
                elif c == 'gravidade_final':
                    col_widths.append(62)
                else:
                    col_widths.append(60)

            table_modelo_ocs = Table(table_rows, repeatRows=1, colWidths=col_widths)
            table_modelo_ocs.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c2d12')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.whitesmoke]),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))
            story.append(table_modelo_ocs)
            story.append(Spacer(1, 6))

            story.append(Paragraph("Evidências e Recomendações", styles['Heading3']))
            for row_idx, (_, row) in enumerate(ocs_m.reset_index(drop=True).iterrows(), start=1):
                placa = row.get('placa', '-')
                data_hora = row.get('data_hora', '-')
                codigo = row.get('codigo_regra', '-')
                tipo = row.get('tipo_ocorrencia', '-')
                gravidade = row.get('gravidade_final', '-')
                evidencia = row.get('evidencia', '-')
                recomendacao = row.get('recomendacao', '-')

                story.append(Paragraph(
                    f"<b>Ocorrência {row_idx}</b> | Placa: {placa} | Data/Hora: {data_hora} | Código: {codigo} | Gravidade: {gravidade}",
                    small_style,
                ))
                story.append(Paragraph(f"Tipo: {tipo}", note_style))
                story.append(Paragraph(f"Evidência: {evidencia}", note_style))
                story.append(Paragraph(f"Recomendação: {recomendacao}", note_style))
                story.append(Spacer(1, 4))

            if idx < len(modelos_com_ocorrencia) - 1:
                story.append(PageBreak())

    doc.build(story)
    output.seek(0)
    return output.getvalue()


def build_pdf_report(filtered_df, freq_suspeita, sequenciais):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    output = BytesIO()
    c = canvas.Canvas(output, pagesize=A4)
    w, h = A4
    y = h - 40

    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, "Relatorio de Risco - Frota Inteligente")
    y -= 20
    c.setFont("Helvetica", 9)
    c.drawString(40, y, f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    y -= 25

    c.setFont("Helvetica-Bold", 11)
    c.drawString(40, y, "Resumo Executivo")
    y -= 18
    c.setFont("Helvetica", 9)

    linhas = [
        f"- Total de abastecimentos: {len(filtered_df)}",
        f"- Total de placas: {filtered_df['Placa'].nunique()}",
        f"- Total de condutores: {filtered_df['Condutor'].nunique()}",
        f"- Litros totais: {filtered_df['litros'].sum():.2f}",
        f"- KM total: {filtered_df['km'].sum():.2f}",
        f"- Consumo medio KM/L: {filtered_df['consumo'].mean():.2f}" if len(filtered_df) else "- Consumo medio KM/L: 0.00",
        f"- Multiplos no mesmo dia: {len(freq_suspeita)}",
        f"- Sequenciais < 1h: {len(sequenciais)}",
    ]

    for linha in linhas:
        if y < 50:
            c.showPage()
            y = h - 40
            c.setFont("Helvetica", 9)
        c.drawString(40, y, linha)
        y -= 15

    c.showPage()
    c.save()
    output.seek(0)
    return output.getvalue()


def build_outliers_pdf_report(tabela_outliers, sigma_mult, limiar_outlier):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=18,
        rightMargin=18,
        topMargin=18,
        bottomMargin=18,
    )
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("<b>Relatorio de Outliers - Frota Inteligente</b>", styles['Title']))
    story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
    story.append(
        Paragraph(
            f"Regra: consumo &lt; media - {sigma_mult:.1f}σ | limiar: {limiar_outlier:.2f} KM/L",
            styles['Normal'],
        )
    )
    story.append(Spacer(1, 8))

    if tabela_outliers is None or tabela_outliers.empty:
        story.append(Paragraph("Nenhum outlier encontrado para os filtros atuais.", styles['Normal']))
        doc.build(story)
        output.seek(0)
        return output.getvalue()

    story.append(Paragraph(f"Total de outliers: {len(tabela_outliers)}", styles['Normal']))
    story.append(Spacer(1, 6))

    df_pdf = tabela_outliers.copy()
    cols_2c = [
        'KM Anterior', 'KM Atual', 'KM Rodados', 'KM Esperado',
        'KM/L', 'Litros', 'Média', 'Mín', 'Máx', 'Desvio',
    ]
    for col in cols_2c:
        if col in df_pdf.columns:
            df_pdf[col] = pd.to_numeric(df_pdf[col], errors='coerce').map(
                lambda x: '-' if pd.isna(x) else f"{x:.2f}"
            )

    for col in df_pdf.columns:
        df_pdf[col] = df_pdf[col].astype(str).replace({'nan': '-', 'None': '-'})

    def _safe_text(value):
        text = str(value)
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    header_style = ParagraphStyle(
        name='pdf_header',
        fontName='Helvetica-Bold',
        fontSize=6,
        leading=7,
        textColor=colors.white,
    )
    body_style = ParagraphStyle(
        name='pdf_body',
        fontName='Helvetica',
        fontSize=5.5,
        leading=6.5,
    )

    table_data = [[Paragraph(_safe_text(col), header_style) for col in df_pdf.columns.tolist()]]
    for row in df_pdf.values.tolist():
        table_data.append([Paragraph(_safe_text(cell), body_style) for cell in row])

    col_widths = [
        52, 45, 45, 70, 45, 45, 45, 45, 45,
        40, 40, 38, 38, 38, 38, 85, 130,
    ]
    if len(col_widths) == len(df_pdf.columns):
        available_width = landscape(A4)[0] - doc.leftMargin - doc.rightMargin
        total_width = sum(col_widths)
        if total_width > available_width:
            scale = available_width / total_width
            col_widths = [max(24, w * scale) for w in col_widths]
    else:
        col_widths = None

    table = Table(table_data, repeatRows=1, colWidths=col_widths)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
        ('LEFTPADDING', (0, 0), (-1, -1), 2),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))

    story.append(table)
    doc.build(story)
    output.seek(0)
    return output.getvalue()


def build_notificacao_docx(notificacoes, metadata=None):
    """
    Gera documento Word padronizado usando o Modelo_relatorio como template,
    preservando cabeçalho timbrado, margens e estilos originais.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    MESES_PT = {
        1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
        5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
        9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro',
    }

    # Mapeamento sigla → nome completo das secretarias
    _SIGLA_NOME = {
        'SEMUTTRANS': 'SECRETARIA MUNICIPAL DE MOBILIDADE URBANA E TRÂNSITO',
        'SMCC':       'SECRETARIA MUNICIPAL DE COMUNICAÇÃO E CULTURA',
        'SMDS':       'SECRETARIA MUNICIPAL DE DESENVOLVIMENTO SOCIAL',
        'SMSU':       'SECRETARIA MUNICIPAL DE SERVIÇOS URBANOS',
        'SMAFEL':     'SECRETARIA MUNICIPAL DE ADMINISTRAÇÃO, FINANÇAS, ESPORTE E LAZER',
        'SMS':        'SECRETARIA MUNICIPAL DE SAÚDE',
        'SMMAP':      'SECRETARIA MUNICIPAL DE MEIO AMBIENTE E PROTEÇÃO ANIMAL',
        'SMCS':       'SECRETARIA MUNICIPAL DE CULTURA E SHOWS',
        'SMO':        'SECRETARIA MUNICIPAL DE OBRAS',
        'SMGAED':     'SECRETARIA MUNICIPAL DE GESTÃO ADMINISTRATIVA E DESENVOLVIMENTO',
        'SMF':        'SECRETARIA MUNICIPAL DE FAZENDA',
        'SMOU':       'SECRETARIA MUNICIPAL DE OBRAS E URBANISMO',
        'SME':        'SECRETARIA MUNICIPAL DE EDUCAÇÃO',
        'SMA':        'SECRETARIA MUNICIPAL DE ADMINISTRAÇÃO',
        'SMTI':       'SECRETARIA MUNICIPAL DE TECNOLOGIA DA INFORMAÇÃO',
        'SMSM':       'SECRETARIA MUNICIPAL DE SERVIÇOS E MANUTENÇÃO',
        'SEMEDES':    'SECRETARIA MUNICIPAL DE DESENVOLVIMENTO ECONÔMICO',
        'SMCT':       'SECRETARIA MUNICIPAL DE CIÊNCIA E TECNOLOGIA',
        'SMMF':       'SECRETARIA MUNICIPAL DE MEIO AMBIENTE E FAZENDAS',
    }

    def _expandir_unidade(sigla: str) -> str:
        """Retorna o nome completo da secretaria pela sigla, ou a própria string se não encontrada."""
        key = sigla.strip().upper()
        return _SIGLA_NOME.get(key, sigla.strip().upper())

    metadata = metadata or {}
    cidade = metadata.get('municipio', '') or 'Santana de Parnaíba'
    responsavel = metadata.get('responsavel', 'Departamento de Transportes')

    agora = datetime.now()
    data_longa = f"{agora.day} de {MESES_PT[agora.month]} de {agora.year}"

    # Abre o Modelo_relatorio como template (preserva cabeçalho timbrado, margens e estilos)
    modelo_path = PROJECT_ROOT / 'Modelo_relatorio'
    if modelo_path.exists():
        doc = Document(str(modelo_path))
        # Limpa o corpo mantendo apenas o sectPr (configurações de seção)
        body = doc.element.body
        for child in list(body):
            if child.tag != qn('w:sectPr'):
                body.remove(child)
    else:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(4.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    def _p(text='', bold=False, italic=False, size=12,
           align=WD_ALIGN_PARAGRAPH.JUSTIFY,
           first_indent_cm=0.0, space_before=240, space_after=240):
        """Adiciona parágrafo com formatação padrão do modelo (12pt, justify, spacing 240)."""
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # Usa twips diretamente para replicar w:spacing w:before="240" w:after="240"
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(space_before))
        spacing.set(qn('w:after'), str(space_after))
        pPr.append(spacing)
        if first_indent_cm:
            p.paragraph_format.first_line_indent = Cm(first_indent_cm)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
        return p

    for i, notif in enumerate(notificacoes):
        if i > 0:
            doc.add_page_break()

        condutor = str(notif.get('condutor', '') or '').strip().upper()
        _unidade_raw = str(notif.get('unidade', '') or '').strip()
        # Ignora placeholders e expande sigla para nome completo
        _unidade_valida = _unidade_raw not in {'', '-', 'nan', 'None', 'NAN', 'NONE'}
        unidade = _expandir_unidade(_unidade_raw) if _unidade_valida else ''
        placa = str(notif.get('placa', ''))
        modelo_veiculo = str(notif.get('modelo', '') or '')
        data_abast = str(notif.get('data_abastecimento', ''))
        estabelecimento = str(notif.get('estabelecimento') or 'não informado')

        # ── Título ─────────────────────────────────────────────────────
        _p('NOTIFICAÇÃO ADMINISTRATIVA — CONTROLE DE ABASTECIMENTO',
           bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

        # ── Data ───────────────────────────────────────────────────────
        _p(f'{cidade}, {data_longa}.', size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)

        _p()  # linha em branco

        # ── Destinatário ───────────────────────────────────────────────
        p_dest = doc.add_paragraph()
        p_dest.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p_dest.add_run(f'Prezado(a) {condutor}')
        r1.bold = True
        r1.font.size = Pt(12)

        if unidade:
            p_uni = doc.add_paragraph()
            p_uni.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_uni = p_uni.add_run(unidade)
            r_uni.bold = True
            r_uni.font.size = Pt(12)

        # ── Corpo ──────────────────────────────────────────────────────
        _p(
            f'Foram identificadas inconsistências técnicas nos registros de '
            f'abastecimento vinculados ao veículo {placa}'
            f'{(" (" + modelo_veiculo + ")") if modelo_veiculo else ""}, '
            f'realizado em {data_abast} no estabelecimento {estabelecimento}.',
            size=12,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_indent_cm=1.25,
        )

        # ── Ocorrências ────────────────────────────────────────────────
        _p('OCORRÊNCIAS IDENTIFICADAS:', size=12,
           align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent_cm=1.25)

        ocorrencias = notif.get('ocorrencias', [])
        for oc in ocorrencias:
            codigo = oc.get('codigo_regra', '')
            tipo = oc.get('tipo_ocorrencia', '').title()
            descricao = oc.get('descricao_tecnica', '')
            linha = f'[{codigo}] {tipo}: {descricao}' if codigo else f'{tipo}: {descricao}'
            _p(linha, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── Solicitação ────────────────────────────────────────────────
        _p(
            'Diante do exposto, solicitamos manifestação formal com apresentação '
            'de justificativas e documentos comprobatórios no prazo de '
            '48 (quarenta e oito) horas a contar do recebimento desta notificação.',
            size=12,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_indent_cm=1.25,
        )

        # ── Ressalva jurídica ──────────────────────────────────────────
        _p(
            'Ressaltamos que a presente análise possui caráter técnico e preliminar, '
            'sendo assegurados ao notificado os princípios do contraditório e da '
            'ampla defesa, nos termos da legislação vigente.',
            size=12,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_indent_cm=1.25,
        )

        _p()  # linha em branco

        # ── Encerramento ───────────────────────────────────────────────
        _p('Atenciosamente,', size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        _p()
        _p()

        # ── Assinatura ─────────────────────────────────────────────────
        p_sign = doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_sign = p_sign.add_run(f'                                       {responsavel}')
        r_sign.bold = True
        r_sign.font.size = Pt(12)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def build_ocorrencias_pdf_report(df_ocs):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=18,
        rightMargin=18,
        topMargin=18,
        bottomMargin=18,
    )
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("<b>Relatorio de Ocorrencias - Auditoria Tecnica</b>", styles['Title']))
    story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
    story.append(Spacer(1, 8))

    if df_ocs is None or df_ocs.empty:
        story.append(Paragraph("Nenhuma ocorrencia identificada para os filtros atuais.", styles['Normal']))
        doc.build(story)
        output.seek(0)
        return output.getvalue()

    colunas = [
        'placa', 'condutor', 'modelo', 'unidade', 'data_hora',
        'codigo_regra', 'tipo_ocorrencia', 'gravidade_final', 'evidencia',
    ]
    colunas_presentes = [c for c in colunas if c in df_ocs.columns]
    df_pdf = df_ocs[colunas_presentes].copy() if colunas_presentes else df_ocs.copy()

    for col in df_pdf.columns:
        df_pdf[col] = df_pdf[col].astype(str).replace({'nan': '-', 'None': '-'})

    def _safe_text(value):
        text = str(value)
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    header_style = ParagraphStyle(
        name='pdf_header_ocs',
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=8,
        textColor=colors.white,
    )
    body_style = ParagraphStyle(
        name='pdf_body_ocs',
        fontName='Helvetica',
        fontSize=6,
        leading=7,
    )

    table_data = [[Paragraph(_safe_text(col), header_style) for col in df_pdf.columns.tolist()]]
    for row in df_pdf.values.tolist():
        table_data.append([Paragraph(_safe_text(cell), body_style) for cell in row])

    col_widths_map = {
        'placa': 45,
        'condutor': 85,
        'modelo': 60,
        'unidade': 45,
        'data_hora': 65,
        'codigo_regra': 42,
        'tipo_ocorrencia': 110,
        'gravidade_final': 55,
        'evidencia': 190,
    }
    col_widths = [col_widths_map.get(c, 70) for c in df_pdf.columns.tolist()]
    available_width = landscape(A4)[0] - doc.leftMargin - doc.rightMargin
    total_width = sum(col_widths)
    if total_width > available_width:
        scale = available_width / total_width
        col_widths = [max(28, w * scale) for w in col_widths]

    table = Table(table_data, repeatRows=1, colWidths=col_widths)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
        ('LEFTPADDING', (0, 0), (-1, -1), 2),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))

    story.append(Paragraph(f"Total de ocorrencias: {len(df_pdf)}", styles['Normal']))
    story.append(Spacer(1, 6))
    story.append(table)

    doc.build(story)
    output.seek(0)
    return output.getvalue()


def load_and_process_data(uploaded_file):
    df = pd.read_excel(BytesIO(uploaded_file))

    # Canoniza variacoes de modelo para melhorar consistencia dos filtros.
    if 'Modelo' in df.columns:
        df['Modelo'] = df['Modelo'].apply(canonicalizar_modelo)
    
    # Converter para datetime usando dayfirst=True
    date_col_candidates = ['Data', 'Data/Hora', 'data', 'data_hora', 'Data Abastecimento']
    date_col = first_existing_column(df, date_col_candidates)
    if date_col:
        df['data'] = process_date_column(df[date_col])
    else:
        return pd.DataFrame()

    # Extract year and month
    df['ano'] = df['data'].dt.year
    df['mes'] = df['data'].dt.month

    # Combustível
    fuel_col_candidates = ['Produto', 'Combustível', 'Combustivel', 'Tipo Combustível', 'Tipo Combustivel']
    fuel_col = first_existing_column(df, fuel_col_candidates)
    if fuel_col:
        df['combustivel'] = df[fuel_col].astype(str).str.strip()
        df['combustivel_norm'] = df['combustivel'].apply(normalize_text)

    # Posto
    posto_col_candidates = ['Estabelecimento', 'Posto', 'Nome Posto', 'Posto de Abastecimento']
    posto_col = first_existing_column(df, posto_col_candidates)
    if posto_col:
        df['posto'] = df[posto_col].astype(str).str.strip()
    else:
        df['posto'] = 'Nao informado'

    # Colunas adicionais
    unidade_col = first_existing_column(df, ['Unidade'])
    ult_km_col = first_existing_column(df, ['Ult. km', 'Ult km', 'Ult.km'])
    km_atual_col = first_existing_column(df, ['km Atual', 'Km Atual', 'KM Atual'])
    produto_col = first_existing_column(df, ['Produto'])

    df['unidade_alerta'] = df[unidade_col].fillna('-').astype(str) if unidade_col else '-'
    df['ult_km_alerta'] = df[ult_km_col] if ult_km_col else np.nan
    df['km_atual_alerta'] = df[km_atual_col] if km_atual_col else np.nan
    df['produto_alerta'] = df[produto_col] if produto_col else 'Nao informado'
    df['modelo_alerta'] = df['Modelo'] if 'Modelo' in df.columns else '-'
    if 'Modelo' not in df.columns:
        df['Modelo'] = '-'

    # Converter numéricas
    df['km'] = pd.to_numeric(df['Km Rodado'], errors='coerce').fillna(0)
    df['litros'] = pd.to_numeric(df['Qtde (L)'], errors='coerce').fillna(0)
    df['ult_km_num'] = pd.to_numeric(df['ult_km_alerta'], errors='coerce')
    df['km_atual_num'] = pd.to_numeric(df['km_atual_alerta'], errors='coerce')

    # Consumo
    df['consumo_calculado'] = df['km'] / df['litros'].replace(0, np.nan)
    if 'km/L' in df.columns:
        df['consumo'] = pd.to_numeric(df['km/L'], errors='coerce').fillna(df['consumo_calculado'])
    else:
        df['consumo'] = df['consumo_calculado']

    df.dropna(subset=['data', 'consumo'], inplace=True)
    return df


def show_auto_dismiss_success(message, duration=4):
    """
    Exibe mensagem de sucesso sem bloquear o rerun.
    duration: mantido apenas por compatibilidade de assinatura
    """
    _ = duration
    st.toast(message, icon="✅")


def show_auto_dismiss_warning(message, duration=5):
    """
    Exibe mensagem de aviso sem bloquear o rerun.
    duration: mantido apenas por compatibilidade de assinatura
    """
    _ = duration
    st.toast(message, icon="⚠️")


def show_data_source_panel(info_base=None):
    """
    Exibe painel da fonte de dados ativa na sidebar.
    """
    with st.sidebar:
        st.divider()
        with st.expander("📊 **Fonte de Dados Ativa**", expanded=True):
            info_base = info_base or {}
            sqlite_carga = (st.session_state.ultima_carga or {}).get('sqlite', {})
            db_path = (
                sqlite_carga.get('db_path')
                or info_base.get('db_path')
                or st.session_state.get('active_db_path', str(SQLITE_DB_PATH))
            )
            db_name = Path(str(db_path)).name if db_path else "-"
            total_historico = sqlite_carga.get('rows_loaded', info_base.get('rows_loaded', 0))
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("📂 Banco", db_name, delta=None)
            with col2:
                st.metric("📍 Modo", "Local")
            
            st.divider()
            
            if info_base.get('status') == 'OK' or st.session_state.ultima_carga:
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("✅ Inseridos", sqlite_carga.get('rows_inserted', 0))
                with col2:
                    st.metric("📦 Total Histórico", total_historico)

                if info_base.get('status') == 'OK':
                    st.caption(f"Fonte ativa: {db_path}")
                else:
                    st.caption("Última carga: sessão atual.")
            else:
                st.info("Nenhuma carga realizada nesta sessão.")
        
        st.divider()


@st.cache_data(show_spinner=False)
def preparar_df_dashboard(_df_internal):
    """
    Adapta base interna (agentes/SQLite) para colunas usadas na interface,
    mantendo as colunas originais para auditoria pre_validated.
    """
    if _df_internal is None or _df_internal.empty:
        return pd.DataFrame()

    df = _df_internal.copy()

    if 'data_hora' in df.columns:
        df['data'] = pd.to_datetime(df['data_hora'], errors='coerce')
    elif 'data' in df.columns:
        df['data'] = pd.to_datetime(df['data'], errors='coerce')

    if 'placa' in df.columns:
        df['Placa'] = df['placa'].astype(str)
    if 'condutor' in df.columns:
        df['Condutor'] = df['condutor'].astype(str)
    if 'modelo' in df.columns:
        df['Modelo'] = df['modelo'].astype(str)
        df['modelo_norm'] = df['Modelo'].apply(canonicalizar_modelo)
    if 'marca' in df.columns:
        df['Marca'] = df['marca'].astype(str)

    if 'produto' in df.columns:
        df['combustivel'] = df['produto'].astype(str)
        df['combustivel_norm'] = df['combustivel'].apply(normalize_text)

    # Garante coluna Produto para exibicao, com fallback entre fontes disponiveis.
    if 'Produto' not in df.columns:
        if 'produto' in df.columns:
            df['Produto'] = df['produto']
        elif 'produto_alerta' in df.columns:
            df['Produto'] = df['produto_alerta']
        elif 'combustivel' in df.columns:
            df['Produto'] = df['combustivel']
        else:
            df['Produto'] = pd.NA

    if 'Produto' in df.columns:
        df['Produto'] = (
            df['Produto']
            .replace(['None', 'none', 'nan', 'NaN', ''], pd.NA)
            .fillna('-')
            .astype(str)
            .str.strip()
        )

    if 'estabelecimento' in df.columns:
        df['posto'] = df['estabelecimento'].astype(str)
    elif 'posto' not in df.columns:
        df['posto'] = 'Nao informado'

    if 'unidade' in df.columns:
        df['unidade_alerta'] = df['unidade'].fillna('-').astype(str)
    elif 'unidade_alerta' not in df.columns:
        df['unidade_alerta'] = '-'
    if 'ult_km' in df.columns:
        df['ult_km_alerta'] = df['ult_km']
    if 'km_atual' in df.columns:
        df['km_atual_alerta'] = df['km_atual']
    if 'produto' in df.columns:
        df['produto_alerta'] = df['produto']
    if 'modelo' in df.columns:
        df['modelo_alerta'] = df['modelo']
        if 'Modelo' not in df.columns:
            df['Modelo'] = df['modelo'].astype(str)
    if 'Modelo' not in df.columns:
        df['Modelo'] = '-'

    if 'km_rodado' in df.columns:
        df['km'] = pd.to_numeric(df['km_rodado'], errors='coerce')
    if 'litros' in df.columns:
        df['litros'] = pd.to_numeric(df['litros'], errors='coerce')

    if 'consumo' not in df.columns and 'km' in df.columns and 'litros' in df.columns:
        df['consumo'] = df['km'] / df['litros'].replace(0, np.nan)
    else:
        df['consumo'] = pd.to_numeric(df.get('consumo', np.nan), errors='coerce')

    if 'data' in df.columns:
        df['ano'] = df['data'].dt.year
        df['mes'] = df['data'].dt.month

    df.dropna(subset=['data', 'consumo'], inplace=True)
    return df


# ═════════════════════════════════════════════════════════════════
# CONFIGURAÇÃO E SETUP
# ═════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="Auditoria de Frota",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Bloqueia tradutores automáticos do browser (causa do NotFoundError React)
st_components.html(
    """
    <script>
    (function() {
        var r = window.parent.document.documentElement;
        r.setAttribute('translate', 'no');
        r.setAttribute('lang', 'pt-BR');
        var m = window.parent.document.createElement('meta');
        m.name = 'google'; m.content = 'notranslate';
        window.parent.document.head.appendChild(m);
    })();
    </script>
    """,
    height=0,
)

# ── Injeta o mesmo CSS do Financeiro para identidade visual consistente ──
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@500;700&family=Space+Grotesk:wght@400;500;700&display=swap');

    #MainMenu, footer { display: none !important; }

    header[data-testid="stHeader"] {
        background: transparent !important;
        border-bottom: 0 !important;
    }

    div[data-testid="stDecoration"] { height: 0 !important; }

    .stApp {
        background:
            radial-gradient(circle at 20% -10%, rgba(56,189,248,0.20), transparent 35%),
            radial-gradient(circle at 90% 0%, rgba(45,212,191,0.16), transparent 30%),
            #0a121b;
        color: #e7eef8;
        font-family: 'Space Grotesk', sans-serif;
    }

    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f1826 0%, #0a121b 100%);
        border-right: 1px solid rgba(142,163,190,0.15);
    }

    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3,
    section[data-testid="stSidebar"] h4,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] .stMarkdown,
    section[data-testid="stSidebar"] [data-testid="stWidgetLabel"] {
        color: #dbe7f5 !important;
        opacity: 1 !important;
    }

    section[data-testid="stSidebar"] [data-testid="stWidgetLabel"] p {
        color: #c9d8ea !important;
        font-weight: 600 !important;
    }

    .block-container {
        padding-top: 0.35rem;
        padding-bottom: 1.2rem;
        max-width: 1600px;
    }

    .section-title {
        font-family: 'Rajdhani', 'Space Grotesk', sans-serif;
        font-size: 1.5rem;
        font-weight: 700;
        letter-spacing: 0.02em;
        margin: 1.1rem 0 0.15rem 0;
        padding-bottom: 0.3rem;
        border-bottom: 1px solid rgba(56,189,248,0.18);
        color: #e7eef8;
    }

    div[data-testid="stDataFrame"] {
        border: 1px solid rgba(142,163,190,0.20);
        border-radius: 10px;
        overflow: hidden;
    }

    .sidebar-brand {
        display: flex;
        align-items: center;
        gap: 14px;
        padding: 8px 0 14px 0;
        border-bottom: 1px solid rgba(142,163,190,0.18);
        margin-bottom: 6px;
    }
    .sidebar-brand-icon { font-size: 2rem; line-height: 1; }
    .sidebar-brand-title {
        font-family: 'Rajdhani', sans-serif;
        font-size: 1.15rem;
        font-weight: 700;
        color: #e7eef8;
        line-height: 1.1;
        letter-spacing: 0.04em;
    }
    .sidebar-brand-sub {
        font-size: 0.7rem;
        color: #8ea3be;
        text-transform: uppercase;
        letter-spacing: 0.1em;
    }
    .sidebar-section {
        font-size: 0.68rem;
        font-weight: 700;
        letter-spacing: 0.14em;
        color: #8ea3be;
        text-transform: uppercase;
        margin: 14px 0 8px 0;
        padding-bottom: 4px;
        border-bottom: 1px solid rgba(142,163,190,0.18);
    }
    .sidebar-footer {
        font-size: 0.68rem;
        color: #8ea3be;
        margin-top: 10px;
        padding-top: 8px;
        border-top: 1px solid rgba(142,163,190,0.15);
    }
    .alert-card {
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 9px 12px;
        border-radius: 8px;
        margin-bottom: 7px;
        background: rgba(10, 18, 27, 0.85);
        border: 1px solid rgba(142,163,190,0.12);
    }
    .alert-card.alert-red    { border-left: 4px solid #ef4444; }
    .alert-card.alert-yellow { border-left: 4px solid #eab308; }
    .alert-card.alert-green  { border-left: 4px solid #22c55e; }
    .alert-icon { font-size: 1.4rem; line-height: 1; }
    .alert-body { display: flex; flex-direction: column; gap: 1px; }
    .alert-count-red    { font-family:'Rajdhani',sans-serif; font-size:1.3rem; font-weight:700; color:#ef4444; line-height:1; }
    .alert-count-yellow { font-family:'Rajdhani',sans-serif; font-size:1.3rem; font-weight:700; color:#eab308; line-height:1; }
    .alert-count-green  { font-family:'Rajdhani',sans-serif; font-size:1.3rem; font-weight:700; color:#22c55e; line-height:1; }
    .alert-label { font-size: 0.75rem; color: #c9d8ea; }
    .alert-detail { font-size: 0.65rem; color: #8ea3be; line-height: 1.5; }
    </style>
    """,
    unsafe_allow_html=True,
)

# Inicializar session_state
if 'df_filtrado' not in st.session_state:
    st.session_state.df_filtrado = None
if 'filtros_aplicados' not in st.session_state:
    st.session_state.filtros_aplicados = {}
if 'resultado_processado' not in st.session_state:
    st.session_state.resultado_processado = None
if 'upload_hash' not in st.session_state:
    st.session_state.upload_hash = None
if 'ultima_carga' not in st.session_state:
    st.session_state.ultima_carga = None
if 'active_db_path' not in st.session_state:
    st.session_state.active_db_path = str(RESULTADOS_DB)
if 'filtros_ui' not in st.session_state:
    st.session_state.filtros_ui = {
        'ano': 'Todos',
        'mes': 'Todos',
        'marca': 'Todos',
        'modelo': 'Todos',
        'condutor': 'Todos',
        'placa': 'Todos',
    }
if 'outlier_sigma_mult' not in st.session_state:
    st.session_state.outlier_sigma_mult = 2.0
if 'show_carga_success' not in st.session_state:
    st.session_state.show_carga_success = False
if 'show_carga_warning' not in st.session_state:
    st.session_state.show_carga_warning = False
if 'df_hist_cache' not in st.session_state:
    st.session_state.df_hist_cache = None
if 'info_hist_cache' not in st.session_state:
    st.session_state.info_hist_cache = {}
if 'modo_consulta' not in st.session_state:
    st.session_state.modo_consulta = 'Completo (base inteira)'

# ═════════════════════════════════════════════════════════════════
# CARREGAMENTO DO BANCO RELATORIO.DB (somente leitura)
# ═════════════════════════════════════════════════════════════════


@st.cache_data(show_spinner="Carregando dados de abastecimento...")
def carregar_dados_relatorio() -> pd.DataFrame:
    """Lê todos os registros de relatorio.db em modo somente leitura."""
    if not RELATORIO_DB.exists():
        return pd.DataFrame()
    uri = f"file:{RELATORIO_DB.as_posix()}?mode=ro"
    conn = sqlite3.connect(uri, uri=True)
    try:
        df = pd.read_sql("SELECT * FROM abastecimentos", conn)
    finally:
        conn.close()
    return df


if not RELATORIO_DB.exists():
    st.error(
        "Banco de dados não encontrado. "
        "Importe os dados na página **Financeiro** e retorne aqui."
    )
    st.stop()

# Carrega a base completa (cacheada)
df_raw_relatorio = carregar_dados_relatorio()

if df_raw_relatorio.empty:
    st.warning("Banco relatorio.db está vazio ou sem a tabela 'abastecimentos'.")
    st.stop()

# Processa pelo pipeline de agentes (ingestão + validação → SQLite de resultados)
_fonte_hash = str(len(df_raw_relatorio))
if st.session_state.df_hist_cache is None or st.session_state.upload_hash != _fonte_hash:
    with st.spinner("Processando pipeline de auditoria..."):
        orq_carga = OrchestradorAuditoria(metadata={
            'db_path': str(RESULTADOS_DB),
            'db_table_historico': 'abastecimentos_historico',
            'db_table': 'abastecimentos_validados',
        })
        carga = orq_carga.atualizar_base_diaria(df_raw_relatorio)

    st.session_state.ultima_carga = carga
    st.session_state.active_db_path = str(RESULTADOS_DB)

    orq_refresh = OrchestradorAuditoria(metadata={
        'db_path': str(RESULTADOS_DB),
        'db_table_historico': 'abastecimentos_historico',
        'db_table': 'abastecimentos_validados',
    })
    df_hist_new, info_hist_new = orq_refresh.carregar_base_historica()
    st.session_state.df_hist_cache = df_hist_new
    st.session_state.info_hist_cache = info_hist_new
    st.session_state.upload_hash = _fonte_hash
    st.session_state.resultado_processado = None

# Exibe informações da carga mais recente
if st.session_state.ultima_carga:
    sqlite_carga = st.session_state.ultima_carga.get('sqlite', {})
    if sqlite_carga.get('status') == 'ERRO_ESCRITA':
        st.error(
            "Falha ao gravar no banco de resultados. "
            f"Caminho: {sqlite_carga.get('db_path', RESULTADOS_DB)} | "
            f"Motivo: {sqlite_carga.get('fallback_reason', 'erro operacional')}"
        )

df_hist = st.session_state.df_hist_cache
info_hist = st.session_state.info_hist_cache

if info_hist.get('db_path'):
    st.session_state.active_db_path = info_hist.get('db_path')

# ── Cabeçalho principal no estilo do Financeiro ──
_ano_ctx = st.session_state.filtros_ui.get('ano', 'Todos')
_filtro_ctx = _ano_ctx if _ano_ctx != 'Todos' else 'Todos os anos'
st.markdown(
    f'<div style="display:flex;align-items:baseline;gap:16px;margin-bottom:0.5rem;padding-bottom:0.4rem;border-bottom:1px solid rgba(142,163,190,0.18);">'
    f'<span style="font-family:\'Rajdhani\',sans-serif;font-size:2rem;font-weight:700;color:#e7eef8;letter-spacing:0.02em;">AUDITORIA DE ABASTECIMENTO</span>'
    f'<span style="font-size:0.88rem;color:#8ea3be;font-family:\'Space Grotesk\',sans-serif;">Análise inteligente da frota &nbsp;•&nbsp; {_filtro_ctx}</span>'
    f'</div>',
    unsafe_allow_html=True,
)

if df_hist is None or df_hist.empty:
    st.warning("Nenhum dado processado ainda. Tente recarregar a página.")
    st.stop()

df = preparar_df_dashboard(df_hist)
if df.empty:
    st.warning("Base historica encontrada, mas sem dados validos para analise.")
    st.stop()

# ═════════════════════════════════════════════════════════════════
# SEÇÃO DE FILTROS
# ═════════════════════════════════════════════════════════════════

_MONTHS = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}
_MONTH_NAME_TO_NUMBER = {v: k for k, v in _MONTHS.items()}

fuel_map = {
    'Gasolina': 'gasolina',
    'Alcool': 'alcool',
    'DIESEL S10': 'diesel s10'
}
filtros_ui = st.session_state.filtros_ui

# Opções dos filtros
anos_disponiveis = ["Todos"] + sorted(
    df['ano'].dropna().unique().astype(int).astype(str).tolist(), reverse=True
) if 'ano' in df.columns else ["Todos"]
fuel_options = ["Todos"] + sorted(df['combustivel'].dropna().unique().tolist()) if 'combustivel' in df.columns else ["Todos", "Gasolina", "Alcool", "DIESEL S10"]
unidade_options = ["Todas"] + sorted(df['unidade_alerta'].dropna().astype(str).unique().tolist()) if 'unidade_alerta' in df.columns else ["Todas"]
meses_disponiveis = ["Todos"] + [_MONTHS[m] for m in sorted(_MONTHS.keys())]
marca_options = ["Todos"] + sorted(df['Marca'].dropna().astype(str).unique().tolist()) if 'Marca' in df.columns else ["Todos"]
cond_options = ["Todos"] + sorted(df['Condutor'].dropna().astype(str).unique().tolist())
placa_options = ["Todos"] + sorted(df['Placa'].dropna().astype(str).unique().tolist())

# Gera modelo_options com base na marca selecionada
_marca_sel = filtros_ui.get('marca', 'Todos')
df_modelos_filtro = df[df['Marca'].astype(str) == _marca_sel] if (_marca_sel != 'Todos' and 'Marca' in df.columns) else df
modelos_norm = (
    df_modelos_filtro['Modelo'].dropna().astype(str).str.strip()
    .replace('', pd.NA).dropna().apply(canonicalizar_modelo)
) if 'Modelo' in df_modelos_filtro.columns else pd.Series(dtype=str)
modelo_options = ["Todos"] + sorted(modelos_norm.dropna().unique().tolist())

# Índices padrão
_ano_default = filtros_ui.get('ano', str(datetime.now().year))
_ano_idx = anos_disponiveis.index(_ano_default) if _ano_default in anos_disponiveis else 0

with st.sidebar:
    # ── Brand ──
    st.markdown(
        """
        <div class="sidebar-brand">
            <span class="sidebar-brand-icon">🔍</span>
            <div>
                <div class="sidebar-brand-title">AUDITORIA DE FROTA</div>
                <div class="sidebar-brand-sub">Sistema Multiagente</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander("🔍 Filtros", expanded=False):
        selected_ano = st.selectbox("Ano", anos_disponiveis, index=_ano_idx, key="aud_sel_ano")
        selected_mes = st.selectbox("Mês", meses_disponiveis, index=safe_index(meses_disponiveis, filtros_ui.get('mes', 'Todos')), key="aud_sel_mes")
        selected_marca = st.selectbox("Marca", marca_options, index=safe_index(marca_options, filtros_ui.get('marca', 'Todos')), key="aud_sel_marca")
        selected_modelo = st.selectbox("Modelo", modelo_options, index=safe_index(modelo_options, filtros_ui.get('modelo', 'Todos')), key="aud_sel_modelo")
        selected_condutor = st.selectbox("Motorista", cond_options, index=safe_index(cond_options, filtros_ui.get('condutor', 'Todos')), key="aud_sel_condutor")
        selected_placa = st.selectbox("Placa", placa_options, index=safe_index(placa_options, filtros_ui.get('placa', 'Todos')), key="aud_sel_placa")
        sigma_mult = st.slider(
            "Sensibilidade de Outlier (σ)",
            min_value=0.5, max_value=3.0,
            value=float(st.session_state.outlier_sigma_mult),
            step=0.1,
            help="Regra usada: consumo < media - σ*desvio. Quanto menor o σ, mais sensivel.",
        )
        col_ap, col_lim = st.columns(2)
        with col_ap:
            aplicar_filtros = st.button("✅ Aplicar", use_container_width=True)
        with col_lim:
            if st.button("🗑️ Limpar", use_container_width=True):
                st.session_state.filtros_ui = {
                    'ano': 'Todos', 'mes': 'Todos', 'marca': 'Todos',
                    'modelo': 'Todos', 'condutor': 'Todos', 'placa': 'Todos',
                }
                st.session_state.resultado_processado = None
                st.rerun()

    # ── Rodapé da sidebar ──
    _data_max_aud = pd.to_datetime(df['data'], errors='coerce').max() if 'data' in df.columns and not df.empty else None
    _ultima_aud = _data_max_aud.strftime('%d/%m/%Y') if _data_max_aud is not None and not pd.isna(_data_max_aud) else '—'
    st.markdown(
        f'<div class="sidebar-footer">🕐 Última atualização: <b>{_ultima_aud}</b><br>Base: {RELATORIO_DB.name}</div>',
        unsafe_allow_html=True,
    )

if aplicar_filtros:
    st.session_state.outlier_sigma_mult = float(sigma_mult)
    st.session_state.filtros_ui = {
        'ano': selected_ano,
        'mes': selected_mes,
        'marca': selected_marca,
        'modelo': selected_modelo,
        'condutor': selected_condutor,
        'placa': selected_placa,
    }

    data_version = (
        st.session_state.get('upload_hash', ''),
        st.session_state.get('active_db_path', ''),
        int(len(df)),
        str(df['ano'].max()) if 'ano' in df.columns and not df.empty else '',
    )
    filtered_df_submit, modelo_filter_list_submit = apply_filters_cached(
        df,
        tuple(st.session_state.filtros_ui.items()),
        tuple(fuel_map.items()),
        data_version,
    )
    st.session_state.df_filtrado = filtered_df_submit

    if filtered_df_submit.empty:
        st.session_state.resultado_processado = None
        st.warning("⚠️ Nenhum dado encontrado com os filtros aplicados.")
    else:
        with st.spinner("Processando dados da planilha filtrada..."):
            resultado_processado = {}

            total_registros_filtrados = len(filtered_df_submit)
            analise_limitada = False
            filtered_df_analise = filtered_df_submit

            # Estatísticas apenas sobre consumo no range válido (0 < consumo <= 100 km/L)
            # Exclui km zerado (consumo <= 0) e hodômetros absurdos (consumo > 100)
            _CONSUMO_MAX_VALIDO = 100.0
            _consumo_serie = pd.to_numeric(filtered_df_analise['consumo'], errors='coerce')
            _consumo_ok = _consumo_serie[(_consumo_serie > 0) & (_consumo_serie <= _CONSUMO_MAX_VALIDO)].dropna()
            media = _consumo_ok.mean()
            desvio = _consumo_ok.std()
            consumo_max = _consumo_ok.max()
            consumo_min = _consumo_ok.min()
            sigma_mult = float(st.session_state.get('outlier_sigma_mult', 2.0))
            limiar_outlier = media - (sigma_mult * desvio)

            resultado_processado['stats'] = {
                'media': media,
                'desvio': desvio,
                'max': consumo_max,
                'min': consumo_min,
                'sigma_mult': sigma_mult,
                'limiar_outlier': limiar_outlier,
            }

            df_outliers = filtered_df_analise.assign(
                outlier=(pd.to_numeric(filtered_df_analise.get('consumo', np.nan), errors='coerce') < limiar_outlier)
            )
            resultado_processado['df_outliers'] = df_outliers

            df_outliers_only = df_outliers[df_outliers['outlier']].copy()
            if not df_outliers_only.empty:
                df_outliers_only['Data/Hora'] = pd.to_datetime(df_outliers_only['data'], errors='coerce').dt.strftime('%d/%m/%Y %H:%M')
                df_outliers_only['Motorista'] = df_outliers_only.get('Condutor', np.nan)
                df_outliers_only['Unidade'] = df_outliers_only.get('unidade_alerta', np.nan)
                df_outliers_only['KM Anterior'] = pd.to_numeric(df_outliers_only.get('ult_km_alerta', np.nan), errors='coerce')
                df_outliers_only['KM Atual'] = pd.to_numeric(df_outliers_only.get('km_atual_alerta', np.nan), errors='coerce')
                df_outliers_only['KM Rodados'] = pd.to_numeric(df_outliers_only.get('km', np.nan), errors='coerce')
                df_outliers_only['KM Esperado'] = (media * pd.to_numeric(df_outliers_only.get('litros', np.nan), errors='coerce')).round(2)
                df_outliers_only['KM/L'] = pd.to_numeric(df_outliers_only.get('consumo', np.nan), errors='coerce')
                df_outliers_only['Litros'] = pd.to_numeric(df_outliers_only.get('litros', np.nan), errors='coerce')
                df_outliers_only['Média'] = float(media)
                df_outliers_only['Mín'] = float(consumo_min)
                df_outliers_only['Máx'] = float(consumo_max)
                df_outliers_only['Desvio'] = float(desvio)
                df_outliers_only['Tipo de Anomalia'] = 'Rendimento abaixo do limiar estatistico'
                df_outliers_only['Evidência'] = (
                    'KM/L '
                    + df_outliers_only['KM/L'].round(2).astype(str)
                    + f' abaixo de {limiar_outlier:.2f} (média - {sigma_mult:.1f}σ)'
                )

                colunas_outlier_ordem = [
                    'Data/Hora',
                    'Placa',
                    'Modelo',
                    'Motorista',
                    'Unidade',
                    'KM Anterior',
                    'KM Atual',
                    'KM Rodados',
                    'KM Esperado',
                    'KM/L',
                    'Litros',
                    'Média',
                    'Mín',
                    'Máx',
                    'Desvio',
                    'Tipo de Anomalia',
                    'Evidência',
                ]
                resultado_processado['tabela_outliers'] = df_outliers_only[colunas_outlier_ordem].copy()
            else:
                resultado_processado['tabela_outliers'] = pd.DataFrame()

            risco_placa = df_outliers[df_outliers['outlier']].groupby('Placa').size().sort_values(ascending=False)
            risco_condutor = df_outliers[df_outliers['outlier']].groupby('Condutor').size().sort_values(ascending=False)
            resultado_processado['risco_placa'] = risco_placa
            resultado_processado['risco_condutor'] = risco_condutor

            orq = OrchestradorAuditoria(metadata={
                'municipio': st.session_state.get('municipio', ''),
                'responsavel': st.session_state.get('responsavel', 'Departamento de Transportes'),
                'db_path': str(RESULTADOS_DB),
                'db_table_historico': 'abastecimentos_historico',
                'db_table': 'abastecimentos_validados',
                'outlier_sigma_mult': float(st.session_state.get('outlier_sigma_mult', 2.0)),
            })
            resultado_auditoria = orq.run_pipeline(filtered_df_analise, pre_validated=True)
            resultado_auditoria = filtrar_resultado_auditoria_por_recorte(filtered_df_analise, resultado_auditoria)
            resultado_processado['auditoria'] = resultado_auditoria
            resultado_processado['orq'] = orq
            resultado_processado['sqlite'] = resultado_auditoria.get('sqlite', {})

            alert_df = filtered_df_analise.sort_values(['Placa', 'data']).copy()
            alert_df['data_dia'] = alert_df['data'].dt.date
            alert_df['data_hora'] = alert_df['data'].dt.strftime('%d/%m/%Y %H:%M')
            alert_df['ordem_abastecimento'] = alert_df.groupby('Placa').cumcount() + 1

            freq = alert_df.groupby(['Placa', 'data_dia']).size().reset_index(name='qtd_abastecimentos')
            freq_suspeita = freq[freq['qtd_abastecimentos'] >= 2]

            alert_df['delta_min'] = (
                alert_df.groupby('Placa')['data']
                .diff().dt.total_seconds().div(60)
            )
            alert_df['data_hora_anterior'] = alert_df.groupby('Placa')['data_hora'].shift(1)
            alert_df['posto_anterior'] = alert_df.groupby('Placa')['posto'].shift(1)
            alert_df['litros_anterior'] = alert_df.groupby('Placa')['litros'].shift(1)
            alert_df['km_atual_evento'] = pd.to_numeric(alert_df.get('km_atual_alerta', np.nan), errors='coerce')
            alert_df['km_atual_evento'] = alert_df['km_atual_evento'].combine_first(pd.to_numeric(alert_df.get('km', np.nan), errors='coerce'))
            alert_df['km_anterior'] = alert_df.groupby('Placa')['km_atual_evento'].shift(1)
            alert_df['consumo_anterior'] = alert_df.groupby('Placa')['consumo'].shift(1)
            alert_df['sequencia'] = alert_df['ordem_abastecimento'].astype(str).radd('#')
            sequenciais = alert_df[(alert_df['delta_min'].notna()) & (alert_df['delta_min'] < 60)].copy()

            detalhes_freq = alert_df.merge(
                freq_suspeita[['Placa', 'data_dia']],
                on=['Placa', 'data_dia'], how='inner'
            )

            resultado_processado['freq_suspeita'] = freq_suspeita
            resultado_processado['detalhes_freq'] = detalhes_freq
            resultado_processado['sequenciais'] = sequenciais
            resultado_processado['alert_df'] = alert_df
            resultado_processado['filtered_df'] = filtered_df_analise
            resultado_processado['modelo_filter'] = modelo_filter_list_submit
            resultado_processado['total_registros_filtrados'] = total_registros_filtrados
            resultado_processado['registros_em_analise'] = len(filtered_df_analise)

            st.session_state.resultado_processado = resultado_processado

data_version = (
    st.session_state.get('upload_hash', ''),
    st.session_state.get('active_db_path', ''),
    int(len(df)),
    str(df['data'].max()) if 'data' in df.columns and not df.empty else '',
)
filtered_df, modelo_filter_list = apply_filters_cached(
    df,
    tuple(st.session_state.filtros_ui.items()),
    tuple(fuel_map.items()),
    data_version,
)

st.session_state.filtros_aplicados = {
    'fuel': st.session_state.filtros_ui.get('fuel', 'Todos'),
    'marca': st.session_state.filtros_ui.get('marca', 'Todos'),
    'modelo': modelo_filter_list,
    'condutor': st.session_state.filtros_ui.get('condutor', 'Todos'),
    'placa': st.session_state.filtros_ui.get('placa', 'Todos'),
}
st.session_state.df_filtrado = filtered_df

st.sidebar.caption(f"Registros no recorte: {len(filtered_df):,}".replace(",", "."))

# ═════════════════════════════════════════════════════════════════
# ABAS DE RESULTADOS
# ═════════════════════════════════════════════════════════════════

if st.session_state.resultado_processado is not None:
    res = st.session_state.resultado_processado
    filtered_df = res['filtered_df']
    media = res['stats']['media']
    desvio = res['stats']['desvio']
    consumo_max = res['stats']['max']
    consumo_min = res['stats']['min']
    sigma_mult = res['stats'].get('sigma_mult', 2.0)
    limiar_outlier = res['stats'].get('limiar_outlier', np.nan)
    resultado_auditoria = res['auditoria']
    sqlite_info = res.get('sqlite', {})

    st.info(
            f"ℹ️ Analise executada sobre {res.get('total_registros_filtrados', 0):,} registros.".replace(',', '.')
        )

    aplicar_estilo_relatorio()
    st.caption("")

    tab1, tab2, tab3, tab4 = st.tabs([
        "📊 Visão Geral",
        "🔍 Ocorrências",
        "⚠️ Alertas & Timeline",
        "📥 Dados & Relatórios",
    ])

    # ══════════════════════════════════════════════════════════════
    # TAB 1 — Visão Geral: KPIs + dispersão de consumo + outliers
    # ══════════════════════════════════════════════════════════════
    with tab1:
        # ── Painel executivo dos agentes ─────────────────────────
        renderizar_painel_executivo(resultado_auditoria, sqlite_info)

        st.divider()

        # ── Estatística + gráfico de dispersão ───────────────────
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Máx (KM/L)", f"{consumo_max:.2f}")
        c2.metric("Mín (KM/L)", f"{consumo_min:.2f}")
        c3.metric("Média (KM/L)", f"{media:.2f}")
        c4.metric("Desvio padrão", f"{desvio:.2f}")
        st.caption(f"Limiar de outlier: consumo < {limiar_outlier:.2f} km/L  (média − {sigma_mult:.1f}σ)")

        # Filtra registros com consumo absurdo antes de plotar
        _CONSUMO_MAX_VALIDO = 100.0
        _df_plot = res['df_outliers'].copy()
        _df_plot['_consumo_num'] = pd.to_numeric(_df_plot['consumo'], errors='coerce')
        _excluidos_plot = int((_df_plot['_consumo_num'] > _CONSUMO_MAX_VALIDO).sum())
        _df_plot = _df_plot[(_df_plot['_consumo_num'] > 0) & (_df_plot['_consumo_num'] <= _CONSUMO_MAX_VALIDO)].copy()
        if _excluidos_plot:
            st.caption(f"⚠️ {_excluidos_plot} registro(s) com consumo > {_CONSUMO_MAX_VALIDO:.0f} km/L excluídos do gráfico (hodômetro inválido).")

        fig_disp = px.scatter(
            _df_plot,
            x='consumo', y='litros',
            color='outlier',
            color_discrete_map={True: '#ef4444', False: '#38bdf8'},
            hover_data=['Placa', 'Modelo', 'Condutor', 'consumo', 'litros', 'km', 'data'],
            title='Dispersão KM/L × Litros abastecidos',
            template='plotly_dark',
        )
        for vtype, cap in VEHICLE_TANK_CAPACITIES.items():
            if not res['modelo_filter'] or any(vtype.lower() in m.lower() for m in res['modelo_filter']):
                fig_disp.add_hline(y=cap, line_dash='dot',
                                   annotation_text=f'{vtype} {cap}L',
                                   annotation_position='bottom right')
        fig_disp.update_layout(
            xaxis_title='KM/L', yaxis_title='Litros',
            paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,24,38,0.6)',
            margin=dict(l=8, r=8, t=40, b=8),
        )
        st.plotly_chart(fig_disp, use_container_width=True)

        # ── Outliers detalhados (colapsado por padrão) ────────────
        tabela_outliers = res.get('tabela_outliers', pd.DataFrame())
        if not tabela_outliers.empty:
            with st.expander(f"🔎 Outliers detalhados — {len(tabela_outliers)} registro(s)", expanded=False):
                cols_2c = ['KM Anterior','KM Atual','KM Rodados','KM Esperado','KM/L','Litros','Média','Mín','Máx','Desvio']
                tov = tabela_outliers.copy()
                for col in cols_2c:
                    if col in tov.columns:
                        tov[col] = pd.to_numeric(tov[col], errors='coerce').round(2)
                render_dataframe_limited(tov)

                d1, d2 = st.columns(2)
                csv_ot = tov.to_csv(index=False, sep=';', encoding='utf-8-sig')
                d1.download_button("📥 CSV", data=csv_ot, file_name="outliers.csv", mime="text/csv", use_container_width=True)
                xlsx_ot = BytesIO()
                with pd.ExcelWriter(xlsx_ot, engine='openpyxl') as w:
                    tov.to_excel(w, index=False, sheet_name='Outliers')
                xlsx_ot.seek(0)
                d2.download_button("📥 Excel", data=xlsx_ot.getvalue(), file_name="outliers.xlsx",
                                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                   use_container_width=True)

    # ══════════════════════════════════════════════════════════════
    # TAB 2 — Ocorrências: lista dos agentes + ranking de placas
    # ══════════════════════════════════════════════════════════════
    with tab2:
        ocs_class = resultado_auditoria.get('ocorrencias', [])
        relatorio_agente = resultado_auditoria.get('relatorio', {})
        notificacoes_agente = resultado_auditoria.get('notificacoes', [])
        res_exec = relatorio_agente.get('resumo_executivo', {})

        r1, r2, r3 = st.columns(3)
        r1.metric("Ocorrências totais", res_exec.get('total_ocorrencias', 0))
        r2.metric("Placas com risco", len(relatorio_agente.get('ranking_placa', [])))
        r3.metric("Notificações prontas", len(notificacoes_agente))

        # Log colapsado
        with st.expander("📋 Log do pipeline", expanded=False):
            for linha in resultado_auditoria.get('log', []):
                st.code(linha, language=None)

        if ocs_class:
            df_ocs = pd.DataFrame(ocs_class)

            gravidade_sel = st.radio(
                "Filtrar gravidade",
                options=["ALTA", "MÉDIA", "TODAS"],
                horizontal=True,
                key="visao_ocorrencias",
            )
            if gravidade_sel == "ALTA":
                df_ocs_view = df_ocs[df_ocs.get('gravidade_final', df_ocs.get('gravidade_inicial', '')).isin(['ALTA', 'CRITICA'])]
            elif gravidade_sel == "MÉDIA":
                df_ocs_view = df_ocs[df_ocs.get('gravidade_final', df_ocs.get('gravidade_inicial', '')).isin(['ALTA', 'CRITICA', 'MEDIA'])]
            else:
                df_ocs_view = df_ocs

            if df_ocs_view.empty:
                st.success("✅ Nenhuma ocorrência nessa faixa de gravidade.")
            else:
                _ocs_rename = {
                    'data_hora': 'Data/Hora', 'placa': 'Placa', 'modelo': 'Modelo',
                    'condutor': 'Condutor', 'unidade': 'Unidade', 'litros': 'Litros',
                    'km_l': 'KM/L', 'valor_observado': 'Valor Observado',
                    'valor_referencia': 'Valor Referência', 'gravidade_final': 'Gravidade',
                    'gravidade_inicial': 'Gravidade Inicial', 'codigo_regra': 'Regra',
                    'tipo_ocorrencia': 'Tipo de Anomalia', 'descricao_tecnica': 'Evidência',
                    'recomendacao': 'Recomendação',
                }
                _src_cols = [c for c in [
                    'data_hora','placa','modelo','condutor','unidade',
                    'gravidade_final','gravidade_inicial','codigo_regra','tipo_ocorrencia',
                    'litros','km_l','valor_observado','valor_referencia','descricao_tecnica','recomendacao'
                ] if c in df_ocs_view.columns]
                _ocs_display = df_ocs_view[_src_cols].rename(columns=_ocs_rename).reset_index(drop=True)
                if 'Data/Hora' in _ocs_display.columns:
                    _ocs_display['Data/Hora'] = pd.to_datetime(_ocs_display['Data/Hora'], errors='coerce').dt.strftime('%d/%m/%Y %H:%M').fillna('-')
                for _nc in ['Litros','KM/L','Valor Observado','Valor Referência']:
                    if _nc in _ocs_display.columns:
                        _ocs_display[_nc] = pd.to_numeric(_ocs_display[_nc], errors='coerce').round(2)
                render_dataframe_limited(_ocs_display)

                # Notificações — uma por acordeon
                if notificacoes_agente:
                    st.markdown(f"**📬 Minutas administrativas — {len(notificacoes_agente)} notificação(ões)**")
                    for i, notif in enumerate(notificacoes_agente, 1):
                        label = f"{i}. Placa {notif.get('placa','?')} · {notif.get('condutor','?')} — {notif.get('gravidade_max','?')}"
                        with st.expander(label, expanded=False):
                            st.text_area("", value=notif.get('texto_notificacao', notif.get('minuta', '')), height=260, key=f"notif_{i}", label_visibility="collapsed")
                            try:
                                _word_bytes = build_notificacao_docx(
                                    [notif],
                                    metadata={
                                        'municipio': st.session_state.get('municipio', ''),
                                        'responsavel': st.session_state.get('responsavel', 'Departamento de Transportes'),
                                    }
                                )
                                _placa_slug = str(notif.get('placa', 'notificacao')).replace(' ', '_')
                                st.download_button(
                                    "📄 Baixar Notificação (Word)",
                                    data=_word_bytes,
                                    file_name=f"notificacao_{_placa_slug}_{i}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_notif_word_{i}",
                                    use_container_width=True,
                                )
                            except Exception as _e:
                                st.warning(f"⚠️ Erro ao gerar Word: {_e}")
        else:
            st.success("✅ Nenhuma ocorrência identificada com os dados filtrados.")

    # ══════════════════════════════════════════════════════════════
    # TAB 3 — Alertas & Timeline
    # ══════════════════════════════════════════════════════════════
    with tab3:
        # ── Alertas de frequência ─────────────────────────────────
        alert_df = res.get('alert_df')
        freq_suspeita = res.get('freq_suspeita')
        detalhes_freq = res.get('detalhes_freq')

        if alert_df is None or alert_df.empty:
            alert_df = filtered_df.sort_values(['Placa', 'data']).copy()
            alert_df['data_dia'] = pd.to_datetime(alert_df['data'], errors='coerce').dt.date
            alert_df['data_hora'] = pd.to_datetime(alert_df['data'], errors='coerce').dt.strftime('%d/%m/%Y %H:%M')
            alert_df['ordem_abastecimento'] = alert_df.groupby('Placa').cumcount() + 1
            freq = alert_df.groupby(['Placa', 'data_dia']).size().reset_index(name='qtd_abastecimentos')
            freq_suspeita = freq[freq['qtd_abastecimentos'] >= 2].copy()
            alert_df['delta_min'] = alert_df.groupby('Placa')['data'].diff().dt.total_seconds().div(60)
            alert_df['data_hora_anterior'] = alert_df.groupby('Placa')['data_hora'].shift(1)
            alert_df['posto_anterior'] = alert_df.groupby('Placa')['posto'].shift(1)
            alert_df['litros_anterior'] = alert_df.groupby('Placa')['litros'].shift(1)
            alert_df['km_rodados'] = pd.to_numeric(alert_df.get('km', np.nan), errors='coerce')
            alert_df['km_atual_evento'] = pd.to_numeric(alert_df.get('km_atual_alerta', np.nan), errors='coerce')
            alert_df['km_anterior_base'] = pd.to_numeric(alert_df.get('ult_km_alerta', np.nan), errors='coerce')
            alert_df['km_atual_evento'] = alert_df['km_atual_evento'].combine_first(alert_df['km_rodados'])
            alert_df['km_anterior'] = alert_df.groupby('Placa')['km_atual_evento'].shift(1)
            alert_df['consumo_anterior'] = alert_df.groupby('Placa')['consumo'].shift(1)
            alert_df['sequencia'] = alert_df['ordem_abastecimento'].astype(str).radd('#')
            detalhes_freq = alert_df.merge(freq_suspeita[['Placa', 'data_dia']], on=['Placa', 'data_dia'], how='inner')
        else:
            alert_df = alert_df.copy()
            freq_suspeita = freq_suspeita if freq_suspeita is not None else pd.DataFrame()
            # Garantir colunas derivadas para exibição (podem estar ausentes em cache antigo)
            if 'data_hora' not in alert_df.columns:
                alert_df['data_hora'] = pd.to_datetime(alert_df['data'], errors='coerce').dt.strftime('%d/%m/%Y %H:%M')
            if 'data_dia' not in alert_df.columns:
                alert_df['data_dia'] = pd.to_datetime(alert_df['data'], errors='coerce').dt.date
            if 'ordem_abastecimento' not in alert_df.columns:
                alert_df['ordem_abastecimento'] = alert_df.groupby('Placa').cumcount() + 1
            if 'sequencia' not in alert_df.columns:
                alert_df['sequencia'] = alert_df['ordem_abastecimento'].astype(str).radd('#')
            if 'delta_min' not in alert_df.columns:
                alert_df['delta_min'] = alert_df.groupby('Placa')['data'].diff().dt.total_seconds().div(60)
            if 'km_atual_evento' not in alert_df.columns:
                alert_df['km_atual_evento'] = pd.to_numeric(alert_df.get('km_atual_alerta', np.nan), errors='coerce')
                alert_df['km_atual_evento'] = alert_df['km_atual_evento'].combine_first(
                    pd.to_numeric(alert_df.get('km', np.nan), errors='coerce'))
            if 'km_anterior' not in alert_df.columns:
                alert_df['km_anterior'] = alert_df.groupby('Placa')['km_atual_evento'].shift(1)
            if 'consumo_anterior' not in alert_df.columns:
                alert_df['consumo_anterior'] = alert_df.groupby('Placa')['consumo'].shift(1)
            if 'data_hora_anterior' not in alert_df.columns:
                alert_df['data_hora_anterior'] = alert_df.groupby('Placa')['data_hora'].shift(1)
            if 'posto_anterior' not in alert_df.columns:
                alert_df['posto_anterior'] = alert_df.groupby('Placa')['posto'].shift(1)
            if 'litros_anterior' not in alert_df.columns:
                alert_df['litros_anterior'] = alert_df.groupby('Placa')['litros'].shift(1)
            # Reconstruir detalhes_freq com todas as colunas atuais
            if not freq_suspeita.empty and 'data_dia' in alert_df.columns:
                detalhes_freq = alert_df.merge(freq_suspeita[['Placa', 'data_dia']], on=['Placa', 'data_dia'], how='inner')
            else:
                detalhes_freq = pd.DataFrame()

        sequenciais_alerta = alert_df[(alert_df['delta_min'].notna()) & (alert_df['delta_min'] < 60)].copy()
        proximo_delta = alert_df.groupby('Placa')['delta_min'].shift(-1)
        mascara_bloco = (
            ((alert_df['delta_min'].notna()) & (alert_df['delta_min'] < 60))
            | ((proximo_delta.notna()) & (proximo_delta < 60))
        )
        sequenciais = alert_df[mascara_bloco].copy()

        a1, a2 = st.columns(2)
        a1.metric("Múltiplos no mesmo dia", int(len(freq_suspeita)))
        a2.metric("Sequenciais < 1h", int(len(sequenciais_alerta)))

        consumo_media_alertas = pd.to_numeric(alert_df.get('consumo', np.nan), errors='coerce').mean()
        consumo_min_alertas = pd.to_numeric(alert_df.get('consumo', np.nan), errors='coerce').min()
        consumo_max_alertas = pd.to_numeric(alert_df.get('consumo', np.nan), errors='coerce').max()
        consumo_desvio_alertas = pd.to_numeric(alert_df.get('consumo', np.nan), errors='coerce').std()

        def _padronizar_alertas(df_src, colunas_alvo):
            if df_src is None or df_src.empty:
                return pd.DataFrame(columns=colunas_alvo)
            # Remover colunas que serão geradas pelo rename para evitar duplicatas
            _to_drop = []
            _rename_map = {
                'sequencia': 'Sequência', 'data_hora_anterior': 'Data/Hora Anterior',
                'data_hora': 'Data/Hora Atual', 'posto_anterior': 'Posto Anterior',
                'posto': 'Posto Atual', 'delta_min': 'Intervalo (min)',
                'litros_anterior': 'Litros Anterior', 'litros': 'Litros',
                'km_anterior': 'KM Anterior', 'km_atual_evento': 'KM Atual',
                'km': 'KM Rodados', 'consumo_anterior': 'KM/L Anterior', 'consumo': 'KM/L Atual',
                'unidade_alerta': 'Unidade', 'modelo_alerta': 'Modelo',
            }
            # Se já existe a coluna destino E existe a coluna origem, dropar a coluna destino antes
            src = df_src.copy()
            for orig, dest in _rename_map.items():
                if orig in src.columns and dest in src.columns and orig != dest:
                    src = src.drop(columns=[dest])
            view = src.rename(columns=_rename_map)
            # Fallbacks para colunas que podem ter nomes alternativos
            if 'Modelo' not in view.columns and 'modelo' in view.columns:
                view['Modelo'] = view['modelo']
            if 'Unidade' not in view.columns and 'unidade' in view.columns:
                view['Unidade'] = view['unidade']
            # Remover quaisquer duplicatas remanescentes (mantém a primeira)
            view = view.loc[:, ~view.columns.duplicated()]
            if 'KM Esperado' in colunas_alvo:
                view['KM Esperado'] = pd.to_numeric(view.get('Litros', np.nan), errors='coerce') * float(consumo_media_alertas)
            if 'Média' in colunas_alvo:
                view['Média'] = float(consumo_media_alertas)
            for c in colunas_alvo:
                if c not in view.columns:
                    view[c] = pd.NA
            cols_num = [c for c in ['Intervalo (min)','Litros','Litros Anterior','KM Anterior','KM Atual','KM Rodados','KM Esperado','Média','KM/L Anterior','KM/L Atual'] if c in colunas_alvo]
            for c in cols_num:
                view[c] = pd.to_numeric(view[c], errors='coerce').round(2)
            # Ordena e substitui NaN por '-' sem desalinhar índice
            _sorted = view[colunas_alvo].sort_values(['Placa', 'Data/Hora Atual'])
            return _sorted.where(pd.notna(_sorted), '-')

        col_mult = ['Placa','Modelo','Unidade','Sequência','Data/Hora Atual','Posto Atual','Condutor','KM Anterior','KM Atual','KM Rodados','Litros','KM/L Atual','Média']
        col_seq  = ['Placa','Modelo','Unidade','Sequência','Data/Hora Anterior','Data/Hora Atual','Posto Anterior','Posto Atual','Condutor','Intervalo (min)','Litros Anterior','Litros','KM Rodados','KM/L Anterior','KM/L Atual']

        with st.expander("📌 Múltiplos abastecimentos no mesmo dia", expanded=True):
            if detalhes_freq.empty:
                st.info("Nenhum caso encontrado.")
            else:
                render_dataframe_limited(_padronizar_alertas(detalhes_freq, col_mult))

        with st.expander("⏱️ Abastecimentos sequenciais < 1h", expanded=True):
            if sequenciais.empty:
                st.info("Nenhum caso encontrado.")
            else:
                render_dataframe_limited(_padronizar_alertas(sequenciais, col_seq))

        st.divider()

        # ── Timeline por placa (fragment: só re-executa este bloco ao mudar a placa) ──
        @st.fragment
        def _timeline_fragment():
            _res = st.session_state.get('resultado_processado')
            if _res is None:
                return
            _fdf = _res['filtered_df']
            _limiar = _res['stats'].get('limiar_outlier', float('nan'))
            _lim_ok = _limiar == _limiar  # False se NaN

            st.markdown("#### Timeline por Placa")
            placas_disp = sorted(_fdf['Placa'].dropna().astype(str).unique().tolist())
            placa_sel = st.selectbox("Placa", placas_disp, key="sel_placa_tab3")
            _tl_base = _fdf[_fdf['Placa'].astype(str) == str(placa_sel)]
            timeline_df = _tl_base.sort_values('data').copy()
            if timeline_df.empty:
                st.warning("Sem dados para esta placa nos filtros aplicados.")
                return
            timeline_df['consumo_num'] = pd.to_numeric(timeline_df['consumo'], errors='coerce')
            _litros_raw = pd.to_numeric(timeline_df['litros'], errors='coerce').fillna(0)
            timeline_df['litros_size'] = _litros_raw.clip(lower=1)
            timeline_df['outlier'] = timeline_df['consumo_num'] < _limiar if _lim_ok else False
            timeline_df['_cor'] = timeline_df['outlier'].map({True: 'Outlier', False: 'Normal'})

            fig_tl = px.scatter(
                timeline_df, x='data', y='consumo_num',
                size='litros_size', color='_cor',
                color_discrete_map={'Normal': '#38bdf8', 'Outlier': '#ef4444'},
                hover_data={'Placa': True, 'Condutor': True, 'posto': True,
                            'km': True, 'litros': True, 'consumo_num': True,
                            'litros_size': False, '_cor': False, 'outlier': False},
                title=f"Placa {placa_sel} — consumo ao longo do tempo",
                template='plotly_dark',
            )
            fig_tl.update_layout(
                xaxis_title='Data', yaxis_title='KM/L',
                paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(15,24,38,0.6)',
                margin=dict(l=8, r=8, t=40, b=8),
                showlegend=True, legend_title_text='',
            )
            _y_vals = timeline_df['consumo_num'].dropna()
            if not _y_vals.empty:
                _y_min, _y_max = _y_vals.min(), _y_vals.max()
                _y_pad = max((_y_max - _y_min) * 0.15, 1)
                _rmin, _rmax = _y_min - _y_pad, _y_max + _y_pad
                fig_tl.update_yaxes(range=[_rmin, _rmax])
                if _lim_ok and _rmin <= _limiar <= _rmax:
                    fig_tl.add_hline(y=_limiar, line_dash='dash', line_color='#ef4444',
                                     annotation_text=f'Limiar {_limiar:.2f}',
                                     annotation_position='bottom right')
            st.plotly_chart(fig_tl, use_container_width=True)

            cols_tl = [c for c in ['data','Placa','Modelo','unidade_alerta','Condutor','posto',
                                    'ult_km_alerta','km_atual_alerta','km','litros','consumo_num']
                       if c in timeline_df.columns]
            tl_view = timeline_df[cols_tl].rename(columns={
                'data': 'Data/Hora', 'posto': 'Posto', 'unidade_alerta': 'Unidade',
                'ult_km_alerta': 'KM Anterior', 'km_atual_alerta': 'KM Atual',
                'km': 'KM Rodados', 'litros': 'Litros', 'consumo_num': 'KM/L',
            })
            for col in ['KM Anterior', 'KM Atual', 'KM Rodados', 'Litros', 'KM/L']:
                if col in tl_view.columns:
                    tl_view[col] = pd.to_numeric(tl_view[col], errors='coerce').round(2)
            if 'Data/Hora' in tl_view.columns:
                tl_view['Data/Hora'] = pd.to_datetime(tl_view['Data/Hora'], errors='coerce').dt.strftime('%d/%m/%Y %H:%M')
            st.dataframe(tl_view.reset_index(drop=True), use_container_width=True)

        _timeline_fragment()

    # ══════════════════════════════════════════════════════════════
    # TAB 4 — Dados & Relatórios
    # ══════════════════════════════════════════════════════════════
    with tab4:
        dt1, dt2 = st.tabs(["🗂️ Tabela de Dados", "📥 Exportações & Manual"])

        with dt1:
            tabela_base = filtered_df.copy()
            if 'Produto' not in tabela_base.columns:
                for src in ['produto', 'produto_alerta', 'combustivel']:
                    if src in tabela_base.columns:
                        tabela_base['Produto'] = tabela_base[src]; break
                else:
                    tabela_base['Produto'] = pd.NA
            tabela_base['Produto'] = tabela_base['Produto'].replace(['None','none','nan','NaN',''], pd.NA).fillna('-').astype(str).str.strip()

            colunas_base = ['data','Placa','Modelo','unidade_alerta','Condutor','Marca','posto','Produto','ult_km_alerta','km_atual_alerta','km','litros','consumo']
            for c in colunas_base:
                if c not in tabela_base.columns:
                    tabela_base[c] = pd.NA
            tabela_exibicao = tabela_base[[c for c in colunas_base if c in tabela_base.columns]].copy()
            tabela_exibicao = tabela_exibicao.rename(columns={
                'data':'Data/Hora','unidade_alerta':'Unidade','posto':'Posto',
                'ult_km_alerta':'KM Anterior','km_atual_alerta':'KM Atual',
                'km':'KM Rodados','litros':'Litros','consumo':'KM/L',
            })
            for col in ['KM Anterior','KM Atual','KM Rodados','Litros','KM/L']:
                if col in tabela_exibicao.columns:
                    tabela_exibicao[col] = pd.to_numeric(tabela_exibicao[col], errors='coerce').round(2)
            if 'Data/Hora' in tabela_exibicao.columns:
                tabela_exibicao['Data/Hora'] = pd.to_datetime(tabela_exibicao['Data/Hora'], errors='coerce').dt.strftime('%d/%m/%Y %H:%M')
            tabela_exibicao = tabela_exibicao.where(pd.notna(tabela_exibicao), '-').replace({'None':'-','nan':'-','NaT':'-'})

            st.caption(f"{len(tabela_exibicao):,} registros no recorte atual".replace(',', '.'))
            render_dataframe_limited(tabela_exibicao)

            d1, d2 = st.columns(2)
            csv_tab = tabela_exibicao.to_csv(index=False, sep=';', encoding='utf-8-sig')
            d1.download_button("📥 CSV", data=csv_tab, file_name="dados_auditoria.csv", mime="text/csv", use_container_width=True)
            xlsx_tab = BytesIO()
            with pd.ExcelWriter(xlsx_tab, engine='openpyxl') as w:
                tabela_exibicao.to_excel(w, index=False, sheet_name='Dados')
            xlsx_tab.seek(0)
            d2.download_button("📥 Excel", data=xlsx_tab.getvalue(), file_name="dados_auditoria.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               use_container_width=True)

        with dt2:
            st.markdown("#### Manual")
            manual_data = carregar_manual_docx_estruturado(str(MANUAL_DOCX_PATH))
            renderizar_manual_no_dashboard(manual_data)

else:
    st.info("👈 Configure os filtros e clique em **'Processar Dados'** para ver os resultados.")
