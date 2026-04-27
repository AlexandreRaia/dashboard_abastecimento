"""
Lab — Leitura de Cupom Fiscal / NFCe SP
Experimento: OCR + QR Code → extração de chave → consulta SEFAZ SP → dossiê PDF
"""

from __future__ import annotations

import re
import warnings
from io import BytesIO

import numpy as np
import streamlit as st

# ── imports opcionais (com fallback amigável) ────────────────────────────────
try:
    import cv2
    _CV2_OK = True
except ImportError:
    _CV2_OK = False

try:
    import zxingcpp as _zxingcpp
    _ZXING_OK = True
except (ImportError, Exception):
    _ZXING_OK = False

try:
    from pyzbar import pyzbar as _pyzbar
    _PYZBAR_OK = True
except (ImportError, Exception):
    _PYZBAR_OK = False

try:
    from rapidocr_onnxruntime import RapidOCR as _RapidOCR
    from PIL import Image, ImageEnhance
    _RAPIDOCR_OK = True
except ImportError:
    _RAPIDOCR_OK = False



try:
    import pandas as pd
    _PANDAS_OK = True
except ImportError:
    _PANDAS_OK = False

try:
    from playwright.sync_api import sync_playwright as _sync_playwright
    _PLAYWRIGHT_OK = True
except (ImportError, Exception):
    _PLAYWRIGHT_OK = False

# ── configuração da página ───────────────────────────────────────────────────
st.set_page_config(page_title="Lab — NFCe", page_icon="🧪", layout="wide")
st.title("🧪 Lab — Cupom Fiscal / NFCe SP")
st.caption("Experimento: leitura de imagem → QR Code / OCR → consulta SEFAZ SP → dossiê PDF")

# Alerta de dependências ausentes
_missing = []
if not _CV2_OK:
    _missing.append("`opencv-python-headless`")
if not _RAPIDOCR_OK:
    _missing.append("`rapidocr-onnxruntime` / `Pillow`")

if _missing:
    st.warning(
        f"Pacotes não instalados: {', '.join(_missing)}. "
        "Execute `pip install -r requirements.txt` e reinicie o Streamlit."
    )

# ── session state ────────────────────────────────────────────────────────────
_STATE_KEYS = ("lab_chave", "lab_qr_url", "lab_ocr_text", "lab_ocr_data", "lab_nfe_data", "lab_nfce_screenshot", "lab_nfce_xml")
for _k in _STATE_KEYS:
    if _k not in st.session_state:
        st.session_state[_k] = None
if "lab_batch_items" not in st.session_state:
    st.session_state.lab_batch_items = []  # list[dict] — um item por cupom


# ══════════════════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════════════════

@st.cache_resource
def _get_rapid_ocr():
    """Carrega o motor RapidOCR (ONNX) uma única vez por sessão do servidor."""
    return _RapidOCR()


def _preprocess_image(img: "Image.Image") -> "Image.Image":
    """Converte para escala de cinza e melhora contraste para OCR."""
    img = img.convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    return img


def _make_variants(arr_gray: "np.ndarray") -> list:
    """Gera variantes de imagem para tentar decodificação de QR Code."""
    h, w = arr_gray.shape[:2]
    variants = [arr_gray]
    # Upscale 2x
    variants.append(cv2.resize(arr_gray, (w * 2, h * 2), interpolation=cv2.INTER_CUBIC))
    # Threshold Otsu
    _, otsu = cv2.threshold(arr_gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(otsu)
    # Upscale do Otsu
    variants.append(cv2.resize(otsu, (w * 2, h * 2), interpolation=cv2.INTER_NEAREST))
    # Threshold adaptativo
    adapt = cv2.adaptiveThreshold(
        arr_gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    variants.append(adapt)
    # Upscale do adaptativo
    variants.append(cv2.resize(adapt, (w * 2, h * 2), interpolation=cv2.INTER_NEAREST))
    # Inversão (QR dark-on-light vs light-on-dark)
    variants.append(cv2.bitwise_not(otsu))
    # Desfoque suave antes do threshold (reduz ruído de foto)
    blur = cv2.GaussianBlur(arr_gray, (3, 3), 0)
    _, otsu_blur = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(otsu_blur)
    variants.append(cv2.resize(otsu_blur, (w * 2, h * 2), interpolation=cv2.INTER_NEAREST))
    return variants


def _detect_qr(img: "Image.Image") -> str | None:
    """Detecta QR Code — tenta zxingcpp (melhor), pyzbar e OpenCV como fallback.
    Short-circuit: retorna imediatamente ao achar na imagem original ou grayscale,
    evitando gerar variantes pesadas desnecessariamente.
    """
    arr_rgb = np.array(img.convert("RGB"))
    arr_gray = cv2.cvtColor(arr_rgb, cv2.COLOR_RGB2GRAY) if _CV2_OK else None

    # ── 1. zxingcpp — melhor para fotos reais, sem DLLs extras ────────────────
    if _ZXING_OK:
        # Tenta primeiro nas versões simples (rápido, sem variantes)
        for pil in [img, img.convert("L")]:
            try:
                results = _zxingcpp.read_barcodes(np.array(pil))
                for r in results:
                    if "QR" in str(r.format).upper():
                        return r.text
                for r in results:
                    if "nfce" in r.text.lower() or "fazenda" in r.text.lower():
                        return r.text
            except Exception:
                pass
        # Só gera variantes pesadas se as simples falharam
        if _CV2_OK and arr_gray is not None:
            for v in _make_variants(arr_gray):
                try:
                    results = _zxingcpp.read_barcodes(np.array(Image.fromarray(v)))
                    for r in results:
                        if "QR" in str(r.format).upper():
                            return r.text
                    for r in results:
                        if "nfce" in r.text.lower() or "fazenda" in r.text.lower():
                            return r.text
                except Exception:
                    pass

    # ── 2. pyzbar fallback (se DLLs estiverem disponíveis) ──────────────────
    if _PYZBAR_OK:
        for pil in [img, img.convert("L")]:
            try:
                decoded = _pyzbar.decode(pil)
                for d in decoded:
                    if d.type in ("QRCODE", "QR_CODE"):
                        return d.data.decode("utf-8", errors="replace")
            except Exception:
                pass

    # ── 3. OpenCV fallback ───────────────────────────────────────────
    if _CV2_OK and arr_gray is not None:
        detector = cv2.QRCodeDetector()
        for v in _make_variants(arr_gray):
            bgr = cv2.cvtColor(v, cv2.COLOR_GRAY2BGR)
            try:
                data, _, _ = detector.detectAndDecode(bgr)
                if data:
                    return data
            except Exception:
                pass

    return None


def _chave_from_qr(qr_url: str) -> str | None:
    """Extrai os 44 dígitos da chave de acesso da URL do QR Code NFCe."""
    # Parâmetro p=CHAVE44|outros
    m = re.search(r'[?&]p=(\d{44})', qr_url)
    if m:
        return m.group(1)
    # Fallback: primeiros 44 dígitos contíguos na URL
    m = re.search(r'(\d{44})', qr_url)
    return m.group(1) if m else None


def _chave_from_text(text: str) -> str | None:
    """
    Extrai chave de acesso NFC-e/NFe (44 dígitos) do texto OCR.
    Suporta: 11 grupos de 4 dígitos (padrão SEFAZ), sequência contínua
    ou linha com exatamente 44 dígitos.
    """
    # 1. Padrão SEFAZ impresso: 11 grupos de 4 dígitos separados por espaços
    #    Ex: "3526 0420 8385 8300 0131 6500 4000 0892 5910 0977 0379"
    #    Usa [ \t]+ (não \s) para NÃO cruzar linhas
    m = re.search(
        r'(\d{4})[ \t]+(\d{4})[ \t]+(\d{4})[ \t]+(\d{4})[ \t]+(\d{4})[ \t]+'
        r'(\d{4})[ \t]+(\d{4})[ \t]+(\d{4})[ \t]+(\d{4})[ \t]+(\d{4})[ \t]+(\d{4})',
        text,
    )
    if m:
        return ''.join(m.groups())

    # 2. Linha com exatamente 44 dígitos (com ou sem espaços internos)
    for linha in text.splitlines():
        digits = re.sub(r'\D', '', linha)
        if len(digits) == 44:
            return digits
        if len(digits) > 44:
            m2 = re.findall(r'\d{44}', digits)
            if m2:
                return m2[0]

    # 3. Sequência contínua no texto todo (fallback)
    clean = re.sub(r'[\s\-\.]', '', text)
    found = re.findall(r'(?<![\d])\d{44}(?![\d])', clean)
    if found:
        return found[0]

    return None


def _parse_ocr_fields(text: str) -> dict:
    """Extrai campos estruturados do texto bruto do OCR."""
    data: dict = {}

    # CNPJ
    m = re.search(r'\d{2}[.\s]\d{3}[.\s]\d{3}[/\s]\d{4}[-\s]\d{2}', text)
    if m:
        data["cnpj"] = re.sub(r'\s', '', m.group())

    # Data DD/MM/AAAA
    m = re.search(r'\d{2}/\d{2}/\d{4}', text)
    if m:
        data["data_emissao"] = m.group()

    # Hora HH:MM
    m = re.search(r'\d{2}:\d{2}(?::\d{2})?', text)
    if m:
        data["hora"] = m.group()

    # Valor total
    m = re.search(r'(?:total|valor\s*total)[^\d\n]*?([\d.,]+)', text, re.IGNORECASE)
    if m:
        data["valor_total"] = m.group(1)
    else:
        vals = re.findall(r'R?\$?\s*(\d{1,3}(?:[.,]\d{3})*[.,]\d{2})', text)
        if vals:
            data["valor_total"] = vals[-1]

    # Itens / produtos
    # Linhas que parecem item: texto + quantidade + valor
    # Ex: "GASOLINA COMUM        1 UN   R$ 5,99"
    # Tenta detectar padrão: descrição seguida de qtd e valor monetário
    itens = []
    _STOP = re.compile(
        r'^(?:cnpj|cpf|ie\b|insc|data|hora|total|subtotal|desconto|troco|'
        r'valor|forma|pagamento|dinheiro|cartao|pix|documento|nota|chave|'
        r'consumidor|endere|fone|tel|fab|lote|val|sat|cfe|nfc|nfce|coo)'
        r'|^\s*$',
        re.IGNORECASE,
    )
    _VALOR = re.compile(r'[\d.,]{3,}')
    linhas = text.splitlines()
    i = 0
    while i < len(linhas):
        linha = linhas[i].strip()
        # Ignora linhas de cabeçalho/rodapé
        if _STOP.search(linha):
            i += 1
            continue
        # Linha deve ter pelo menos 4 chars de texto e conter um valor numérico
        letras = re.sub(r'[\d\s.,/*()\-]', '', linha)
        if len(letras) >= 3 and _VALOR.search(linha):
            # Tenta extrair: descrição | qtd | un | valor
            m_item = re.match(
                r'^(.+?)\s+(\d+[.,]?\d*)\s*(?:UN|KG|LT|L|PC|CX|MT|M|G)?\s+'
                r'(?:R\$\s*)?(\d+[.,]\d{2})\s*$',
                linha,
                re.IGNORECASE,
            )
            if m_item:
                itens.append({
                    "descricao": m_item.group(1).strip(),
                    "qtd": m_item.group(2),
                    "valor": m_item.group(3),
                })
            else:
                # Linha simples com descrição e valor no final
                m_simples = re.match(r'^(.+?)\s+(?:R\$\s*)?(\d+[.,]\d{2})\s*$', linha)
                if m_simples:
                    descr = m_simples.group(1).strip()
                    letras2 = re.sub(r'[\d\s.,/*()\-]', '', descr)
                    if len(letras2) >= 3:
                        itens.append({
                            "descricao": descr,
                            "qtd": "",
                            "valor": m_simples.group(2),
                        })
        i += 1

    if itens:
        data["itens"] = itens

# ── mapa UF → sigla ─────────────────────────────────────────────────────────
_UF_MAP = {
    "11": "RO", "12": "AC", "13": "AM", "14": "RR", "15": "PA",
    "16": "AP", "17": "TO", "21": "MA", "22": "PI", "23": "CE",
    "24": "RN", "25": "PB", "26": "PE", "27": "AL", "28": "SE",
    "29": "BA", "31": "MG", "32": "ES", "33": "RJ", "35": "SP",
    "41": "PR", "42": "SC", "43": "RS", "50": "MS", "51": "MT",
    "52": "GO", "53": "DF",
}

_MESES = {
    "01": "Jan", "02": "Fev", "03": "Mar", "04": "Abr",
    "05": "Mai", "06": "Jun", "07": "Jul", "08": "Ago",
    "09": "Set", "10": "Out", "11": "Nov", "12": "Dez",
}


def _parse_chave(chave: str) -> dict:
    """
    Decodifica os campos embutidos na chave de acesso NFe/NFCe (44 dígitos).
    Não requer rede nem OCR — instantâneo.
    Estrutura: cUF(2) AAMM(4) CNPJ(14) Mod(2) Série(3) nNF(9) tpEmis(1) cNF(8) cDV(1)
    """
    if len(chave) != 44 or not chave.isdigit():
        return {}
    uf = _UF_MAP.get(chave[0:2], chave[0:2])
    aa, mm = chave[2:4], chave[4:6]
    cnpj_raw = chave[6:20]
    cnpj = f"{cnpj_raw[:2]}.{cnpj_raw[2:5]}.{cnpj_raw[5:8]}/{cnpj_raw[8:12]}-{cnpj_raw[12:14]}"
    modelo_cod = chave[20:22]
    modelo = "NFCe" if modelo_cod == "65" else ("NFe" if modelo_cod == "55" else modelo_cod)
    serie = str(int(chave[22:25]))
    numero = str(int(chave[25:34]))
    mes_nome = _MESES.get(mm, mm)
    return {
        "uf": uf,
        "emissao": f"{mes_nome}/20{aa}",
        "cnpj_emitente": cnpj,
        "modelo": modelo,
        "serie": serie,
        "numero_nf": numero,
    }


def _render_nfce_as_image(
    ocr_data: dict,
    nfe_data: dict,
    chave: str = "",
    chave_data: dict | None = None,
) -> tuple[bytes | None, str | None]:
    """
    Gera imagem DANFE-style da NFCe renderizando um HTML local com todos os dados
    disponíveis (chave_data + ocr_data + nfe_data).
    Não depende de rede nem do SEFAZ — screenshot confiável sem problemas de animação.
    """
    if not _PLAYWRIGHT_OK:
        return None, "playwright não instalado"

    chave_data = chave_data or {}

    # ── Campos de identidade ──────────────────────────────────────────────────
    emitente  = nfe_data.get("emitente")  or ocr_data.get("emitente")  or "—"
    cnpj      = chave_data.get("cnpj_emitente") or nfe_data.get("cnpj") or ocr_data.get("cnpj") or "—"
    endereco  = nfe_data.get("endereco")  or ocr_data.get("endereco")  or ""
    numero_nf = chave_data.get("numero_nf") or nfe_data.get("numero_nf") or ocr_data.get("numero_nf") or "—"
    serie     = chave_data.get("serie")   or nfe_data.get("serie")    or ocr_data.get("serie")    or "—"
    uf        = chave_data.get("uf")      or ""
    emissao   = chave_data.get("emissao") or nfe_data.get("data_emissao") or ocr_data.get("data_emissao") or ""
    hora      = ocr_data.get("hora") or nfe_data.get("hora") or ""
    data_hora = f"{emissao} {hora}".strip()
    total     = nfe_data.get("total")    or ocr_data.get("total") or ocr_data.get("valor_total") or "—"
    desconto  = nfe_data.get("desconto") or ocr_data.get("desconto") or ""
    troco     = nfe_data.get("troco")    or ocr_data.get("troco")    or ""
    chave_fmt = " ".join(chave[i:i+4] for i in range(0, len(chave), 4)) if len(chave) == 44 else chave

    def _esc(v):
        return str(v).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # ── Itens ────────────────────────────────────────────────────────────────
    itens = nfe_data.get("itens") or ocr_data.get("itens") or []
    rows_html = ""
    for it in itens:
        desc = _esc(it.get("descricao") or it.get("produto") or it.get("nome") or "")
        qtd  = _esc(it.get("qtd") or it.get("quantidade") or it.get("qtde") or "")
        un   = _esc(it.get("unidade") or it.get("un") or "")
        vun  = _esc(it.get("valor_unitario") or it.get("vl_unit") or "")
        vtot = _esc(it.get("valor") or it.get("valor_total") or it.get("vl_total") or "")
        rows_html += (
            f"<tr><td>{desc}</td><td class='c'>{qtd} {un}</td>"
            f"<td class='r'>{vun}</td><td class='r'>{vtot}</td></tr>\n"
        )

    # ── Totais extras + pagamento + protocolo ────────────────────────────────
    pagamentos   = nfe_data.get("pagamentos") or ocr_data.get("pagamentos") or []
    protocolo    = nfe_data.get("protocolo")  or ocr_data.get("protocolo") or ""
    data_prot    = nfe_data.get("data_protocolo") or ""

    extras_html = ""
    if desconto:
        extras_html += f"<div class='extra'>Desconto: R$ {_esc(desconto)}</div>"
    if troco:
        extras_html += f"<div class='extra'>Troco: R$ {_esc(troco)}</div>"

    pag_html = ""
    if pagamentos:
        pag_html = "<div class='section-title'>FORMA DE PAGAMENTO</div>"
        for pg in pagamentos:
            pag_html += f"<div class='pag-line'>{_esc(pg)}</div>"

    prot_html = ""
    if protocolo:
        prot_html = (
            f"<div class='section-title'>PROTOCOLO DE AUTORIZAÇÃO</div>"
            f"<div class='prot'>Nº {_esc(protocolo)}"
            + (f" &nbsp;—&nbsp; {_esc(data_prot)}" if data_prot else "")
            + "</div>"
        )

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<style>
  body {{ font-family: 'Courier New', monospace; font-size: 13px;
         background: #fff; color: #000; margin: 0; padding: 20px 24px; width: 700px; }}
  .danfe-label {{ text-align: center; font-size: 11px; letter-spacing: 1px;
                  font-weight: bold; border: 2px solid #000; padding: 4px 6px;
                  margin-bottom: 8px; }}
  h2 {{ text-align: center; font-size: 15px; margin: 4px 0; text-transform: uppercase; }}
  .sub {{ text-align: center; font-size: 11px; color: #333; margin: 2px 0; }}
  .meta {{ display: flex; justify-content: space-between; font-size: 11px;
           border: 1px solid #ccc; padding: 4px 8px; margin: 8px 0;
           background: #f5f5f5; }}
  .chave {{ font-size: 9.5px; word-break: break-all; text-align: center;
            border: 1px solid #888; padding: 4px; margin: 8px 0;
            background: #fafafa; letter-spacing: 0.5px; }}
  .section-title {{ font-weight: bold; font-size: 11px; background: #222;
                    color: #fff; padding: 3px 6px; margin: 10px 0 0; }}
  table {{ width: 100%; border-collapse: collapse; }}
  th {{ background: #444; color: #fff; padding: 4px 6px;
        text-align: left; font-size: 11px; }}
  td {{ padding: 3px 6px; border-bottom: 1px solid #eee; font-size: 11px; }}
  td.c {{ text-align: center; }}
  td.r {{ text-align: right; }}
  .totals {{ margin-top: 6px; }}
  .total-line {{ display: flex; justify-content: space-between;
                 font-size: 12px; padding: 2px 6px; }}
  .total-line.main {{ font-weight: bold; font-size: 15px;
                      border-top: 2px solid #000; margin-top: 4px; padding-top: 4px; }}
  .extra {{ text-align: right; font-size: 11px; color: #555; padding: 1px 6px; }}
  .pag-line {{ font-size: 11px; padding: 2px 6px; }}
  .prot {{ font-size: 10px; padding: 3px 6px; background: #f9f9f9; }}
  .footer {{ font-size: 9px; color: #666; text-align: center; margin-top: 12px;
             border-top: 1px dashed #aaa; padding-top: 8px; }}
</style>
</head>
<body>
<div class="danfe-label">DOCUMENTO AUXILIAR DA NOTA FISCAL DE CONSUMIDOR ELETRÔNICA</div>
<h2>{_esc(emitente)}</h2>
<div class="sub">CNPJ: {_esc(cnpj)}</div>
{"<div class='sub'>" + _esc(endereco) + ("&nbsp;&nbsp;" + _esc(uf) if uf else "") + "</div>" if endereco else ""}
<div class="meta">
  <span>NF-e Nº {_esc(numero_nf)}</span>
  <span>Série {_esc(serie)}</span>
  {"<span>Emissão: " + _esc(data_hora) + "</span>" if data_hora else ""}
</div>
<div class="chave"><strong>Chave de Acesso:</strong> {_esc(chave_fmt) if chave_fmt else "—"}</div>
<div class="section-title">DETALHAMENTO DOS PRODUTOS / SERVIÇOS</div>
<table>
  <tr><th>Descrição</th><th>Qtd/UN</th><th>Vl. Unit.</th><th>Vl. Total</th></tr>
  {rows_html if rows_html else "<tr><td colspan='4' style='text-align:center;padding:8px'>— itens não disponíveis —</td></tr>"}
</table>
<div class="totals">
  {extras_html}
  <div class="total-line main"><span>VALOR TOTAL</span><span>R$ {_esc(total)}</span></div>
</div>
{pag_html}
{prot_html}
<div class="footer">
  Consulte a autenticidade em www.nfce.fazenda.sp.gov.br<br>
  Cópia gerada pelo sistema de controle de abastecimento
</div>
</body>
</html>"""

    import sys as _sys
    import subprocess as _sub
    import tempfile as _tmp
    import os as _os

    # Salva o HTML em arquivo temp
    _html_file = None
    _png_file = None
    _script_file = None
    try:
        with _tmp.NamedTemporaryFile(
            mode="w", suffix=".html", delete=False, encoding="utf-8"
        ) as _fh:
            _fh.write(html)
            _html_file = _fh.name

        _png_file = _html_file.replace(".html", ".png")

        _script = "\n".join([
            "import sys",
            "html_path = sys.argv[1]",
            "png_path = sys.argv[2]",
            "try:",
            "    from playwright.sync_api import sync_playwright",
            "    with sync_playwright() as p:",
            "        browser = p.chromium.launch(headless=True, args=['--no-sandbox'])",
            "        page = browser.new_page(viewport={'width': 720, 'height': 200})",
            "        page.goto('file:///' + html_path.replace('\\\\', '/'))",
            "        page.wait_for_timeout(500)",
            "        data = page.screenshot(full_page=True)",
            "        browser.close()",
            "        open(png_path, 'wb').write(data)",
            "        sys.stdout.buffer.write(data)",
            "        sys.stdout.buffer.flush()",
            "except Exception:",
            "    import traceback",
            "    sys.stderr.write(traceback.format_exc())",
            "    sys.exit(1)",
            "",
        ])

        with _tmp.NamedTemporaryFile(
            mode="w", suffix=".py", delete=False, encoding="utf-8"
        ) as _fs:
            _fs.write(_script)
            _script_file = _fs.name

        proc = _sub.run(
            [_sys.executable, _script_file, _html_file, _png_file],
            capture_output=True,
            timeout=30,
        )
        stderr_txt = proc.stderr.decode("utf-8", errors="replace").strip()
        if proc.returncode == 0 and proc.stdout:
            return proc.stdout, None
        return None, stderr_txt or f"Processo encerrou com código {proc.returncode}"
    except _sub.TimeoutExpired:
        return None, "Timeout ao gerar imagem local"
    except Exception:
        import traceback as _tb
        return None, _tb.format_exc()
    finally:
        for _f in (_html_file, _script_file):
            if _f and _os.path.exists(_f):
                try:
                    _os.unlink(_f)
                except Exception:
                    pass


def _print_nfce_pdf(url: str, timeout_ms: int = 30000) -> tuple[bytes | None, str | None]:
    """
    Usa Playwright em modo impressão (emulate_media='print') para gerar o PDF
    oficial da NFCe — equivalente a clicar "Imprimir > Salvar como PDF" no browser.
    O CSS @media print força todo o conteúdo visível, sem animações AngularJS.
    Retorna (pdf_bytes, log_str).
    """
    if not _PLAYWRIGHT_OK:
        return None, "playwright não instalado"

    import sys as _sys
    import subprocess as _sub
    import tempfile as _tmp
    import os as _os

    _script_body = "\n".join([
        "import sys",
        "url = sys.argv[1]",
        "timeout_ms = int(sys.argv[2])",
        "try:",
        "    from playwright.sync_api import sync_playwright",
        "    with sync_playwright() as p:",
        "        browser = p.chromium.launch(",
        "            headless=True,",
        "            args=['--no-sandbox', '--disable-blink-features=AutomationControlled'],",
        "        )",
        "        ctx = browser.new_context(",
        "            viewport={'width': 1280, 'height': 1800},",
        "            user_agent=(",
        "                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '",
        "                'AppleWebKit/537.36 (KHTML, like Gecko) '",
        "                'Chrome/124.0.0.0 Safari/537.36'",
        "            ),",
        "        )",
        "        page = ctx.new_page()",
        "        page.add_init_script(\"Object.defineProperty(navigator, 'webdriver', {get: () => undefined})\")",
        "        page.goto(url, wait_until='load', timeout=timeout_ms)",
        "        # Aguarda Angular carregar dados",
        "        try:",
        "            page.wait_for_function('typeof angular !== \"undefined\" && angular.element(document.body).injector() !== undefined', timeout=12000)",
        "            page.wait_for_function('angular.element(document.body).injector().get(\"$http\").pendingRequests.length === 0', timeout=20000)",
        "        except Exception:",
        "            pass",
        "        page.wait_for_timeout(2000)",
        "        body_text = page.evaluate('document.body.innerText')[:200]",
        "        sys.stderr.write('PAGE_TEXT:' + body_text + chr(10))",
        "        # Ativa modo impressão — CSS @media print força tudo visível",
        "        page.emulate_media(media='print')",
        "        page.wait_for_timeout(1000)",
        "        pdf_bytes = page.pdf(",
        "            format='A4',",
        "            print_background=True,",
        "            margin={'top': '10mm', 'bottom': '10mm', 'left': '10mm', 'right': '10mm'},",
        "        )",
        "        browser.close()",
        "        sys.stderr.write('PDF_SIZE:' + str(len(pdf_bytes)) + chr(10))",
        "        sys.stdout.buffer.write(pdf_bytes)",
        "        sys.stdout.buffer.flush()",
        "except Exception:",
        "    import traceback",
        "    sys.stderr.write(traceback.format_exc())",
        "    sys.exit(1)",
        "",
    ])

    _tmp_file = None
    try:
        with _tmp.NamedTemporaryFile(mode="w", suffix=".py", delete=False, encoding="utf-8") as _f:
            _f.write(_script_body)
            _tmp_file = _f.name

        proc = _sub.run(
            [_sys.executable, _tmp_file, url, str(timeout_ms)],
            capture_output=True,
            timeout=(timeout_ms / 1000) + 60,
        )
        stderr_txt = proc.stderr.decode("utf-8", errors="replace").strip()
        if proc.returncode == 0 and proc.stdout:
            return proc.stdout, stderr_txt
        return None, stderr_txt or f"Código {proc.returncode}"
    except _sub.TimeoutExpired:
        return None, "Timeout ao gerar PDF"
    except Exception:
        import traceback as _tb
        return None, _tb.format_exc()
    finally:
        if _tmp_file and _os.path.exists(_tmp_file):
            try:
                _os.unlink(_tmp_file)
            except Exception:
                pass


def _print_nfce_pdf_batch(
    url_list: list[str],
    timeout_ms: int = 30_000,
    progress_cb=None,
) -> list[tuple[bytes | None, str | None]]:
    """
    Gera PDFs de múltiplas NFCe em UMA única sessão de browser Chromium.
    O browser sobe apenas 1x, poupando ~2-3s por cupom vs chamadas individuais.
    Retorna lista de (pdf_bytes | None, log) na mesma ordem de url_list.
    progress_cb(atual, total): chamado a cada cupom concluído (no loop principal).
    """
    if not _PLAYWRIGHT_OK:
        return [(None, "playwright não instalado")] * len(url_list)
    if not url_list:
        return []

    import sys as _sys
    import subprocess as _sub
    import tempfile as _tmp
    import os as _os
    import json as _json
    import base64 as _b64

    _batch_script = "\n".join([
        "import sys, json, base64, traceback",
        "from playwright.sync_api import sync_playwright",
        "",
        "urls = json.loads(open(sys.argv[1], encoding='utf-8').read())",
        "timeout_ms = int(sys.argv[2])",
        "results = []",
        "",
        "def _fetch_one(browser, url):",
        "    ctx = browser.new_context(",
        "        viewport={'width': 1280, 'height': 1800},",
        "        user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124.0.0.0 Safari/537.36',",
        "    )",
        "    page = ctx.new_page()",
        "    page.add_init_script(\"Object.defineProperty(navigator, 'webdriver', {get: () => undefined})\")",
        "    page.goto(url, wait_until='load', timeout=timeout_ms)",
        "    try:",
        "        page.wait_for_function('typeof angular !== \"undefined\"', timeout=8000)",
        "        page.wait_for_function('angular.element(document.body).injector().get(\"$http\").pendingRequests.length === 0', timeout=15000)",
        "    except Exception:",
        "        pass",
        "    page.wait_for_timeout(600)",
        "    page.emulate_media(media='print')",
        "    page.wait_for_timeout(200)",
        "    pdf = page.pdf(format='A4', print_background=True, margin={'top':'10mm','bottom':'10mm','left':'10mm','right':'10mm'})",
        "    ctx.close()",
        "    return pdf",
        "",
        "try:",
        "    with sync_playwright() as p:",
        "        browser = p.chromium.launch(headless=True, args=['--no-sandbox','--disable-blink-features=AutomationControlled'])",
        "        for i, url in enumerate(urls):",
        "            sys.stderr.write(f'PROGRESS:{i+1}/{len(urls)}\\n'); sys.stderr.flush()",
        "            try:",
        "                pdf = _fetch_one(browser, url)",
        "                results.append({'ok': True, 'pdf': base64.b64encode(pdf).decode()})",
        "            except Exception:",
        "                results.append({'ok': False, 'err': traceback.format_exc()})",
        "        browser.close()",
        "except Exception:",
        "    results.extend([{'ok': False, 'err': traceback.format_exc()}] * (len(urls) - len(results)))",
        "",
        "sys.stdout.write(json.dumps(results))",
        "sys.stdout.flush()",
    ])

    _tmp_script = _tmp_urls = None
    try:
        with _tmp.NamedTemporaryFile(mode="w", suffix=".py", delete=False, encoding="utf-8") as _f:
            _f.write(_batch_script)
            _tmp_script = _f.name
        with _tmp.NamedTemporaryFile(mode="w", suffix=".json", delete=False, encoding="utf-8") as _f:
            _json.dump(url_list, _f)
            _tmp_urls = _f.name

        total_timeout = (timeout_ms / 1000 + 5) * len(url_list) + 40
        proc = _sub.Popen(
            [_sys.executable, _tmp_script, _tmp_urls, str(timeout_ms)],
            stdout=_sub.PIPE, stderr=_sub.PIPE,
        )

        # Leitura síncrona do stderr linha-por-linha no loop principal
        # (Streamlit só aceita atualizações de widgets na thread principal)
        stderr_lines: list[str] = []
        stdout_chunks: list[bytes] = []

        # Windows não suporta select() em pipes — usamos threads separadas:
        # - _t_err lê stderr (PROGRESS) sem bloquear o loop principal
        # - _t_out lê stdout em paralelo para evitar deadlock de pipe
        #   (o subprocesso escreve MBs de PDF em base64 no stdout; se nada ler,
        #    o buffer enche, o proc trava e nunca termina)
        import threading as _threading
        _progress_events: list[tuple[int, int]] = []

        def _read_stderr_sync():
            for raw in proc.stderr:
                line = raw.decode("utf-8", errors="replace").strip()
                stderr_lines.append(line)
                if line.startswith("PROGRESS:"):
                    try:
                        cur, tot = line[9:].split("/")
                        _progress_events.append((int(cur), int(tot)))
                    except Exception:
                        pass

        def _read_stdout():
            stdout_chunks.append(proc.stdout.read())

        _t_err = _threading.Thread(target=_read_stderr_sync, daemon=True)
        _t_out = _threading.Thread(target=_read_stdout, daemon=True)
        _t_err.start()
        _t_out.start()

        import time as _t_mod
        _deadline = _t_mod.time() + total_timeout
        _last_reported = 0

        while proc.poll() is None:
            _t_mod.sleep(0.3)
            # Repassa eventos de progresso capturados pela thread
            while len(_progress_events) > _last_reported:
                cur, tot = _progress_events[_last_reported]
                if progress_cb:
                    progress_cb(cur, tot)
                _last_reported += 1
            if _t_mod.time() > _deadline:
                proc.kill()
                return [(None, "Timeout")] * len(url_list)

        # Drena eventos finais
        _t_err.join(timeout=3)
        _t_out.join(timeout=10)
        while len(_progress_events) > _last_reported:
            cur, tot = _progress_events[_last_reported]
            if progress_cb:
                progress_cb(cur, tot)
            _last_reported += 1

        stdout_data = b"".join(stdout_chunks)

        if proc.returncode == 0 and stdout_data:
            items = _json.loads(stdout_data.decode("utf-8"))
            log = "\n".join(stderr_lines)
            return [
                (_b64.b64decode(it["pdf"]), log) if it.get("ok")
                else (None, it.get("err", ""))
                for it in items
            ]
        err = "\n".join(stderr_lines)
        return [(None, err or f"Código {proc.returncode}")] * len(url_list)
    except Exception:
        import traceback as _tb
        return [(None, _tb.format_exc())] * len(url_list)
    finally:
        for _f in (_tmp_script, _tmp_urls):
            if _f and _os.path.exists(_f):
                try:
                    _os.unlink(_f)
                except Exception:
                    pass


def _pdf_first_page_to_png(pdf_bytes: bytes) -> bytes | None:
    """
    Converte a primeira página de um PDF em PNG usando pypdf + reportlab
    ou, se disponível, pdf2image. Retorna png_bytes ou None.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[0]
        mat = fitz.Matrix(2.0, 2.0)  # 2x = ~144dpi
        pix = page.get_pixmap(matrix=mat, alpha=False)
        png = pix.tobytes("png")
        doc.close()
        return png
    except Exception:
        pass
    try:
        from pdf2image import convert_from_bytes as _c
        imgs = _c(pdf_bytes, dpi=150, first_page=1, last_page=1)
        if imgs:
            buf = _io_mod.BytesIO()
            imgs[0].save(buf, format="PNG")
            return buf.getvalue()
    except Exception:
        pass
    return None


def _build_consultation_url(chave: str, qr_url: str | None = None) -> str:
    """
    Monta a URL de consulta pública NFCe SEFAZ SP.

    Prioridade:
    1. URL do QR code capturada na imagem  → contém chave|versão|ambiente|tipo|hash (aceita por ConsultaQrCode.aspx).
    2. Fallback: consulta por chave de acesso (chNFe) → aceita apenas os 44 dígitos,
       sem precisar do hash do QR code.  ConsultaQrCode.aspx com ?p=<chave44> é
       rejeitado pelo SEFAZ ("Formato de QR-Code não suportado").
    """
    if qr_url and qr_url.startswith("http"):
        return qr_url
    # Consulta por chave de acesso — funciona sem o hash do QR
    return (
        "https://www.nfce.fazenda.sp.gov.br/NFCeConsultaPublica"
        f"/Paginas/ConsultaNFCe.aspx?chNFe={chave}"
    )


def _fetch_nfce_pagetext(url: str, timeout_ms: int = 25000) -> tuple[str | None, str | None]:
    """
    Usa Playwright para acessar a página SEFAZ e retornar o texto completo renderizado pelo Angular.
    Retorna (page_text, log_str).
    """
    if not _PLAYWRIGHT_OK:
        return None, "playwright não instalado"

    import sys as _sys
    import subprocess as _sub
    import tempfile as _tmp
    import os as _os

    _script = "\n".join([
        "import sys",
        "url = sys.argv[1]",
        "timeout_ms = int(sys.argv[2])",
        "try:",
        "    from playwright.sync_api import sync_playwright",
        "    with sync_playwright() as p:",
        "        browser = p.chromium.launch(",
        "            headless=True,",
        "            args=['--no-sandbox', '--disable-blink-features=AutomationControlled'],",
        "        )",
        "        ctx = browser.new_context(",
        "            viewport={'width': 1280, 'height': 900},",
        "            user_agent=(",
        "                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '",
        "                'AppleWebKit/537.36 (KHTML, like Gecko) '",
        "                'Chrome/124.0.0.0 Safari/537.36'",
        "            ),",
        "        )",
        "        page = ctx.new_page()",
        "        page.add_init_script(\"Object.defineProperty(navigator, 'webdriver', {get: () => undefined})\")",
        "        page.goto(url, wait_until='load', timeout=timeout_ms)",
        "        # Aguarda Angular carregar os dados",
        "        try:",
        "            page.wait_for_function('typeof angular !== \"undefined\" && angular.element(document.body).injector() !== undefined', timeout=12000)",
        "            page.wait_for_function('angular.element(document.body).injector().get(\"$http\").pendingRequests.length === 0', timeout=20000)",
        "        except Exception:",
        "            pass",
        "        page.wait_for_timeout(3000)",
        "        txt = page.evaluate('document.body.innerText')",
        "        sys.stdout.write(txt)",
        "        sys.stdout.flush()",
        "        browser.close()",
        "except Exception:",
        "    import traceback",
        "    sys.stderr.write(traceback.format_exc())",
        "    sys.exit(1)",
        "",
    ])

    _tmp_file = None
    try:
        with _tmp.NamedTemporaryFile(mode="w", suffix=".py", delete=False, encoding="utf-8") as _f:
            _f.write(_script)
            _tmp_file = _f.name

        proc = _sub.run(
            [_sys.executable, _tmp_file, url, str(timeout_ms)],
            capture_output=True,
            timeout=(timeout_ms / 1000) + 60,
        )
        stderr_txt = proc.stderr.decode("utf-8", errors="replace").strip()
        stdout_txt = proc.stdout.decode("utf-8", errors="replace").strip()
        if proc.returncode == 0 and stdout_txt:
            return stdout_txt, stderr_txt
        return None, stderr_txt or f"Código {proc.returncode}"
    except _sub.TimeoutExpired:
        return None, "Timeout ao acessar SEFAZ"
    except Exception:
        import traceback as _tb
        return None, _tb.format_exc()
    finally:
        if _tmp_file and _os.path.exists(_tmp_file):
            try:
                _os.unlink(_tmp_file)
            except Exception:
                pass


def _parse_nfce_pagetext(text: str) -> dict:
    """
    Parseia o texto extraído do DANFE renderizado pela página SEFAZ SP.
    Retorna dict com emitente, itens, totais, pagamentos e protocolo.
    """
    import re

    lines = [l.strip() for l in text.splitlines()]
    nonempty = [l for l in lines if l]
    result: dict = {}

    # Emitente — primeira linha não-vazia após o cabeçalho DANFE
    for i, l in enumerate(nonempty):
        if "DOCUMENTO AUXILIAR" in l.upper():
            if i + 1 < len(nonempty):
                result["emitente"] = nonempty[i + 1]
            break

    # CNPJ
    m = re.search(r'CNPJ[:\s]+([\d./-]{14,})', text)
    if m:
        result["cnpj"] = m.group(1).strip()

    # Endereço — linha imediatamente após a linha do CNPJ
    for i, l in enumerate(nonempty):
        if re.match(r'CNPJ', l, re.IGNORECASE):
            if i + 1 < len(nonempty):
                result["endereco"] = nonempty[i + 1]
            break

    # Nº nota / série / emissão
    m = re.search(r'N[º°o][\s.:]*(\d+)', text, re.IGNORECASE)
    if m:
        result["numero_nf"] = m.group(1).zfill(6)
    m = re.search(r'S[eé]rie[\s.:]*(\d+)', text, re.IGNORECASE)
    if m:
        result["serie"] = m.group(1)
    m = re.search(r'Emiss[aã]o[\s.:]*(\d{2}/\d{2}/\d{4})', text, re.IGNORECASE)
    if m:
        result["data_emissao"] = m.group(1)
    m = re.search(r'Emiss[aã]o[\s.:]*\d{2}/\d{2}/\d{4}\s+(\d{2}:\d{2}:\d{2})', text, re.IGNORECASE)
    if m:
        result["hora"] = m.group(1)

    # Itens — trabalha no texto "achatado" para cruzar quebras de linha
    flat = re.sub(r'\s+', ' ', text)
    item_pat = re.compile(
        r'(.+?)\s*\(C[oó]digo:\s*\d+\s*\)\s*'
        r'Qtde\.?\s*:?\s*([\d.,]+)\s+UN\s*:?\s*(\S+)\s+'
        r'Vl\.?\s*Unit\.?\s*:?\s*[◆R$\s]*([\d.,]+)\s+'
        r'Vl\.?\s*Total\s*([\d.,]+)',
        re.IGNORECASE,
    )
    itens = []
    for m in item_pat.finditer(flat):
        desc = re.sub(r'^\d+\s+', '', m.group(1).strip())
        itens.append({
            "descricao":      desc,
            "qtd":            m.group(2).replace(',', '.'),
            "unidade":        m.group(3),
            "valor_unitario": m.group(4).replace(',', '.'),
            "valor":          m.group(5).replace(',', '.'),
        })
    if itens:
        result["itens"] = itens

    # Total
    m = re.search(r'Valor\s+total\s+R\$\s*([\d.,]+)', text, re.IGNORECASE)
    if m:
        result["total"] = m.group(1)

    # Desconto
    m = re.search(r'Desconto\s+R\$\s*([\d.,]+)', text, re.IGNORECASE)
    if m:
        try:
            if float(m.group(1).replace('.', '').replace(',', '.')) > 0:
                result["desconto"] = m.group(1)
        except Exception:
            pass

    # Troco
    m = re.search(r'Troco\s+R\$\s*([\d.,]+)', text, re.IGNORECASE)
    if m:
        try:
            if float(m.group(1).replace('.', '').replace(',', '.')) > 0:
                result["troco"] = m.group(1)
        except Exception:
            pass

    # Pagamentos — secção entre "Forma de pagamento" e próximo bloco
    pag_sec = re.search(
        r'Forma\s+de\s+pagamento.+?Valor\s+pago\s*\n(.+?)(?:\n\s*(?:Consulta|CHAVE|Protocolo))',
        text, re.IGNORECASE | re.DOTALL,
    )
    pagamentos = []
    if pag_sec:
        for pl in pag_sec.group(1).splitlines():
            pl = pl.strip()
            mp = re.match(r'(.+?)\s{2,}R\$\s*([\d.,]+)', pl)
            if mp:
                pagamentos.append(f"{mp.group(1).strip()}: R$ {mp.group(2)}")
            else:
                mp2 = re.match(r'(.+?)\s+R\$\s*([\d.,]+)', pl)
                if mp2:
                    pagamentos.append(f"{mp2.group(1).strip()}: R$ {mp2.group(2)}")
    if pagamentos:
        result["pagamentos"] = pagamentos

    # Protocolo de Autorização
    m = re.search(r'Protocolo\s+de\s+Autoriza[çc][aã]o\s*[:\s]*([\d]+)', text, re.IGNORECASE)
    if m:
        result["protocolo"] = m.group(1)
    m = re.search(r'Data/Hora\s*[:\s]*(\d{2}/\d{2}/\d{4}\s+\d{2}:\d{2}:\d{2})', text)
    if m:
        result["data_protocolo"] = m.group(1)

    return result


def _gerar_dossie_docx(
    chave: str,
    chave_data: dict | None = None,
    ocr_data: dict | None = None,
    nfe_data: dict | None = None,
    img_bytes: bytes | None = None,
    qr_url: str | None = None,
    nfce_screenshot: bytes | None = None,
) -> bytes:
    """
    Gera dossiê DOCX usando modelo.docx como template (header + footer + margens).
    Página 1: foto do cupom + dados extraídos.
    Página 2: dados NFCe / SEFAZ.
    """
    from pathlib import Path as _Path
    from docx import Document as _Document
    from docx.shared import Cm as _Cm, Pt as _Pt, RGBColor as _RGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlEl
    import io as _io_mod
    from datetime import datetime as _dt

    ocr_data = ocr_data or {}
    nfe_data = nfe_data or {}
    chave_data = chave_data or {}

    MODELO = _Path(__file__).parent.parent / "modelo.docx"
    if MODELO.exists():
        doc = _Document(str(MODELO))
        sec = doc.sections[0]

        # Extrai o logo (image2.jpg = rId6) e move para o header como imagem inline.
        # No original era imagem flutuante com posição absoluta no corpo —
        # isso causava desalinhamento. No header fica correto em todas as páginas.
        _logo_rel = doc.part.rels.get("rId6")
        _logo_bytes_tmp = _logo_rel.target_part.blob if _logo_rel else None

        hdr = sec.header
        hdr.paragraphs[0].clear()
        hdr_p = hdr.paragraphs[0]
        hdr_p.paragraph_format.space_before = _Pt(0)
        hdr_p.paragraph_format.space_after = _Pt(0)
        if _logo_bytes_tmp:
            hdr_run = hdr_p.add_run()
            hdr_run.add_picture(_io_mod.BytesIO(_logo_bytes_tmp), width=_Cm(16))

        # Ajusta margens: header_distance 0.5cm, top_margin 2.8cm
        # (logo mede ~1.9cm de altura + 0.5cm offset + espaço = 2.8cm até o corpo)
        sec.header_distance = _Cm(0.5)
        sec.top_margin = _Cm(2.8)

        # Remove todo o conteúdo do corpo (imagem flutuante + texto da carta)
        body = doc.element.body
        for child in list(body):
            if child.tag != _qn("w:sectPr"):
                body.remove(child)
    else:
        doc = _Document()

    # ── helper: adiciona parágrafo com estilo ─────────────────────────────────
    def _add_p(text: str, bold: bool = False, size: int = 11,
               align=_WD_ALIGN.LEFT, color=None, space_before: int = 0, space_after: int = 6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = _Pt(space_before)
        p.paragraph_format.space_after = _Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = _Pt(size)
        if color:
            run.font.color.rgb = _RGB(*color)
        return p

    # ── helper: aplica bordas simples via XML em todas as células da tabela ──
    def _set_table_borders(tbl):
        """Aplica bordas finas a todas as células da tabela via XML (independe de estilo)."""
        for row in tbl.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                # Remove tcBorders existente para não duplicar
                for old in tc_pr.findall(_qn("w:tcBorders")):
                    tc_pr.remove(old)
                borders = _OxmlEl("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    b = _OxmlEl(f"w:{side}")
                    b.set(_qn("w:val"), "single")
                    b.set(_qn("w:sz"), "4")
                    b.set(_qn("w:space"), "0")
                    b.set(_qn("w:color"), "BBDEFB")
                    borders.append(b)
                tc_pr.append(borders)

    # ── helper: tabela de dados (label | valor) ───────────────────────────────
    def _add_data_table(rows_data: list[tuple[str, str]]):
        if not rows_data:
            return
        t = doc.add_table(rows=len(rows_data), cols=2)
        for i, (lbl, val) in enumerate(rows_data):
            c0 = t.cell(i, 0)
            c1 = t.cell(i, 1)
            c0.width = _Cm(5)
            c1.width = _Cm(11)
            r0 = c0.paragraphs[0].add_run(lbl)
            r0.bold = True
            r0.font.size = _Pt(9)
            r1 = c1.paragraphs[0].add_run(str(val))
            r1.font.size = _Pt(9)
            # fundo azul claro na coluna label
            tc_pr = c0._tc.get_or_add_tcPr()
            shd = _OxmlEl("w:shd")
            shd.set(_qn("w:val"), "clear")
            shd.set(_qn("w:color"), "auto")
            shd.set(_qn("w:fill"), "E3F2FD")
            tc_pr.append(shd)
        _set_table_borders(t)
        return t

    # ── helper: linha separadora via XML ─────────────────────────────────────
    def _add_hline(color_hex: str = "1565C0", thickness: str = "6"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = _Pt(2)
        p.paragraph_format.space_after = _Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = _OxmlEl("w:pBdr")
        bottom = _OxmlEl("w:bottom")
        bottom.set(_qn("w:val"), "single")
        bottom.set(_qn("w:sz"), thickness)
        bottom.set(_qn("w:space"), "1")
        bottom.set(_qn("w:color"), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ═══════════════════════════════════════════════════════════════════════════
    # PÁGINA 1 — Foto do cupom + dados
    # ═══════════════════════════════════════════════════════════════════════════

    # Linha separadora logo abaixo do logo da prefeitura
    _add_hline("1565C0", "12")

    _add_p("DOSSIÊ — COMPROVANTE FISCAL / NFCe", bold=True, size=13,
           align=_WD_ALIGN.CENTER, color=(21, 101, 192), space_after=2)
    _add_p(f"Gerado em {_dt.now().strftime('%d/%m/%Y %H:%M')}",
           size=8, align=_WD_ALIGN.CENTER, color=(120, 120, 120), space_after=8)

    # Foto do cupom
    if img_bytes:
        try:
            from PIL import Image as _PIL_Image
            _pil = _PIL_Image.open(_io_mod.BytesIO(img_bytes))
            if _pil.mode not in ("RGB", "L"):
                _pil = _pil.convert("RGB")
            _buf_img = _io_mod.BytesIO()
            _pil.save(_buf_img, format="JPEG", quality=85)
            _buf_img.seek(0)
            p_img = doc.add_paragraph()
            p_img.alignment = _WD_ALIGN.CENTER
            p_img.paragraph_format.space_after = _Pt(8)
            run_img = p_img.add_run()
            run_img.add_picture(_buf_img, width=_Cm(8))
        except Exception:
            _add_p("[Imagem do cupom não disponível]", size=9, color=(150, 150, 150))
    else:
        _add_p("[Nenhuma imagem carregada]", size=9, color=(150, 150, 150))

    doc.add_paragraph()  # espaço

    # Dados da chave
    if chave:
        _add_p("Chave de Acesso", bold=True, size=10, color=(21, 101, 192), space_after=2)
        chave_fmt = " ".join(chave[i:i+4] for i in range(0, len(chave), 4))
        p_chave = doc.add_paragraph()
        r = p_chave.add_run(chave_fmt)
        r.font.name = "Courier New"
        r.font.size = _Pt(8)
        p_chave.paragraph_format.space_after = _Pt(6)

    if chave_data:
        _add_p("Dados Identificados (da chave)", bold=True, size=10,
               color=(21, 101, 192), space_after=2)
        rows_kd = [
            ("UF", chave_data.get("uf", "—")),
            ("Emissão", chave_data.get("emissao", "—")),
            ("CNPJ Emitente", chave_data.get("cnpj_emitente", "—")),
            ("Modelo", chave_data.get("modelo", "—")),
            ("Série", chave_data.get("serie", "—")),
            ("Número NF", chave_data.get("numero_nf", "—")),
        ]
        _add_data_table(rows_kd)
        doc.add_paragraph()

    if ocr_data:
        _add_p("Dados Extraídos — OCR", bold=True, size=10,
               color=(21, 101, 192), space_after=2)
        rows_ocr_meta = []
        for k, lbl in [("cnpj", "CNPJ"), ("data_emissao", "Data"), ("hora", "Hora"), ("valor_total", "Total")]:
            if k in ocr_data:
                rows_ocr_meta.append((lbl, str(ocr_data[k])))
        _add_data_table(rows_ocr_meta)

        if ocr_data.get("itens"):
            doc.add_paragraph()
            _add_p("Produtos (OCR)", bold=True, size=10, color=(21, 101, 192), space_after=2)
            t_itens = doc.add_table(rows=1, cols=3)
            hdr_cells = t_itens.rows[0].cells
            for ci, txt in enumerate(["Descrição", "Qtd", "Valor (R$)"]):
                r = hdr_cells[ci].paragraphs[0].add_run(txt)
                r.bold = True
                r.font.size = _Pt(8)
                tc_pr = hdr_cells[ci]._tc.get_or_add_tcPr()
                shd = _OxmlEl("w:shd")
                shd.set(_qn("w:val"), "clear")
                shd.set(_qn("w:color"), "auto")
                shd.set(_qn("w:fill"), "1565C0")
                tc_pr.append(shd)
                hdr_cells[ci].paragraphs[0].runs[0].font.color.rgb = _RGB(255, 255, 255)
            for item in ocr_data["itens"]:
                row = t_itens.add_row().cells
                row[0].paragraphs[0].add_run(str(item.get("descricao", ""))).font.size = _Pt(8)
                row[1].paragraphs[0].add_run(str(item.get("qtd", ""))).font.size = _Pt(8)
                row[2].paragraphs[0].add_run(str(item.get("valor", ""))).font.size = _Pt(8)
            _set_table_borders(t_itens)
        doc.add_paragraph()

    # Dados SEFAZ (se vieram via extração HTML)
    if nfe_data:
        if nfe_data.get("emitente"):
            _add_p("Emitente (SEFAZ)", bold=True, size=10, color=(21, 101, 192), space_after=2)
            _add_data_table([("Emitente", nfe_data["emitente"])])
            doc.add_paragraph()

        if nfe_data.get("itens"):
            _add_p("Itens da Nota (SEFAZ)", bold=True, size=10,
                   color=(21, 101, 192), space_after=2)
            t_nfe = doc.add_table(rows=1, cols=3)
            for ci, txt in enumerate(["Descrição", "Qtd", "Valor"]):
                r = t_nfe.rows[0].cells[ci].paragraphs[0].add_run(txt)
                r.bold = True
                r.font.size = _Pt(8)
            for item in nfe_data["itens"]:
                row = t_nfe.add_row().cells
                row[0].paragraphs[0].add_run(str(item.get("descricao", ""))).font.size = _Pt(8)
                row[1].paragraphs[0].add_run(str(item.get("qtd", ""))).font.size = _Pt(8)
                row[2].paragraphs[0].add_run(str(item.get("valor", ""))).font.size = _Pt(8)
            _set_table_borders(t_nfe)

        for k in ("total", "desconto", "troco"):
            if k in nfe_data:
                _add_data_table([(k.title(), str(nfe_data[k]))])
        doc.add_paragraph()

    # Link SEFAZ
    url_sefaz = qr_url or (
        f"https://www.nfce.fazenda.sp.gov.br/NFCeConsultaPublica"
        f"/Paginas/ConsultaQrCode.aspx?p={chave}"
        if chave else ""
    )
    if url_sefaz:
        _add_p("Consulta Pública SEFAZ SP", bold=True, size=10,
               color=(21, 101, 192), space_after=2)
        # Hyperlink via XML
        p_link = doc.add_paragraph()
        r_id = p_link.part.relate_to(url_sefaz,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        hlink = _OxmlEl("w:hyperlink")
        hlink.set(_qn("r:id"), r_id)
        r_xml = _OxmlEl("w:r")
        rpr = _OxmlEl("w:rPr")
        style_el = _OxmlEl("w:rStyle")
        style_el.set(_qn("w:val"), "Hyperlink")
        rpr.append(style_el)
        r_xml.append(rpr)
        t_xml = _OxmlEl("w:t")
        t_xml.text = url_sefaz[:120]
        r_xml.append(t_xml)
        hlink.append(r_xml)
        p_link._p.append(hlink)

    # ═══════════════════════════════════════════════════════════════════════════
    # PÁGINA 2 — NFC-e (DANFE) em tamanho 100%
    # ═══════════════════════════════════════════════════════════════════════════
    _nfce_img = nfce_screenshot
    if not _nfce_img and _PLAYWRIGHT_OK and (ocr_data or nfe_data or chave):
        _nfce_img, _ = _render_nfce_as_image(
            ocr_data=ocr_data,
            nfe_data=nfe_data,
            chave=chave,
            chave_data=chave_data,
        )

    if _nfce_img:
        doc.add_page_break()

        _add_hline("1565C0", "12")
        _add_p("NFC-e — DANFE (Cópia SEFAZ SP)", bold=True, size=13,
               align=_WD_ALIGN.CENTER, color=(21, 101, 192), space_after=6)

        try:
            from PIL import Image as _PIL_Image
            _pil_ss = _PIL_Image.open(_io_mod.BytesIO(_nfce_img))
            if _pil_ss.mode not in ("RGB", "L"):
                _pil_ss = _pil_ss.convert("RGB")

            # Auto-crop: remove bordas brancas (threshold leniente: só pixels >250 são "fundo")
            _gray = _pil_ss.convert("L")
            _mask = _gray.point(lambda px: 0 if px > 250 else 255)
            _bbox = _mask.getbbox()
            if _bbox:
                _l = max(0, _bbox[0] - 12)
                _t = max(0, _bbox[1] - 12)
                _r = min(_pil_ss.width,  _bbox[2] + 12)
                _b = min(_pil_ss.height, _bbox[3] + 12)
                _pil_ss = _pil_ss.crop((_l, _t, _r, _b))

            _buf_ss = _io_mod.BytesIO()
            _pil_ss.save(_buf_ss, format="JPEG", quality=95)
            _buf_ss.seek(0)
            p_ss = doc.add_paragraph()
            p_ss.alignment = _WD_ALIGN.CENTER
            p_ss.paragraph_format.space_before = _Pt(0)
            p_ss.paragraph_format.space_after = _Pt(0)
            # Usa largura natural da imagem (144dpi = fitz 2x de 72dpi), máx 17cm
            _px_per_cm = 144 / 2.54
            _w_cm = min(_pil_ss.width / _px_per_cm, 17.0)
            p_ss.add_run().add_picture(_buf_ss, width=_Cm(_w_cm))
        except Exception:
            _add_p("[Imagem NFCe não disponível]", size=9, color=(150, 150, 150))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  ABAS
# ══════════════════════════════════════════════════════════════════════════════

tab1, tab2, tab3 = st.tabs(
    ["📷 Leitura do Cupom", "📄 Dados & Dossiê Word", "📦 Lote"]
)

# ─────────────────────────────────────────────────────────────────────────────
# ABA 1 — Leitura do Cupom
# ─────────────────────────────────────────────────────────────────────────────
with tab1:
    st.subheader("📷 Upload e extração do cupom fiscal")

    if not _RAPIDOCR_OK:
        st.error(
            "rapidocr-onnxruntime não instalado. "
            "Execute `pip install rapidocr-onnxruntime` e reinicie o Streamlit."
        )
    else:
        uploaded = st.file_uploader(
            "Envie a foto do cupom fiscal",
            type=["png", "jpg", "jpeg", "webp"],
            help="Foto nítida do cupom, de frente, com QR Code visível.",
        )

        if uploaded:
            img = Image.open(uploaded)
            # salva bytes para o PDF
            uploaded.seek(0)
            st.session_state.lab_img_bytes = uploaded.read()

            # ── Detecção do QR ANTES de renderizar os campos, para que os valores
            #    já estejam no session_state quando os widgets forem desenhados.
            _qr_detected = _detect_qr(img)
            if _qr_detected and _qr_detected != st.session_state.lab_qr_url:
                st.session_state.lab_qr_url = _qr_detected
                _chave_qr_auto = _chave_from_qr(_qr_detected)
                if _chave_qr_auto:
                    st.session_state.lab_chave = _chave_qr_auto

            col_img, col_result = st.columns([1, 1], gap="large")

            with col_img:
                st.image(img, caption="Imagem carregada", use_container_width=True)

            # ── Campos de chave / QR — preenchidos automaticamente pelo QR ou manualmente
            with st.container(border=True):
                _fi_col1, _fi_col2 = st.columns(2)
                with _fi_col1:
                    if st.session_state.lab_chave:
                        st.session_state["fi_chave_input"] = st.session_state.lab_chave
                    _fi_chave = st.text_input(
                        "🔑 Chave de acesso (44 dígitos):",
                        max_chars=44,
                        placeholder="Ex: 35240100000000000000550010000000001000000001",
                        key="fi_chave_input",
                    )
                    _fi_chave = re.sub(r"\s+", "", _fi_chave)
                    if _fi_chave:
                        if len(_fi_chave) == 44 and _fi_chave.isdigit():
                            st.session_state.lab_chave = _fi_chave
                            st.success("✅ Chave válida (44).")
                        else:
                            st.warning(f"Deve ter 44 dígitos numéricos (atual: {len(_fi_chave)}).")
                with _fi_col2:
                    if st.session_state.lab_qr_url:
                        st.session_state["fi_qr_input"] = st.session_state.lab_qr_url
                    _fi_qr = st.text_input(
                        "📷 URL do QR Code NFCe:",
                        placeholder="https://www.nfce.fazenda.sp.gov.br/qrcode?p=...",
                        key="fi_qr_input",
                    )
                    if _fi_qr and _fi_qr.startswith("http"):
                        st.session_state.lab_qr_url = _fi_qr
                        if not st.session_state.lab_chave:
                            _chave_from_qr_manual = _extract_chave_from_url(_fi_qr)
                            if _chave_from_qr_manual:
                                st.session_state.lab_chave = _chave_from_qr_manual
                        st.caption("✅ URL salva.")

            with col_result:
                st.markdown("#### Extração automática")

                # 1. QR Code — resultado da detecção já feita acima
                qr_data = _qr_detected
                if qr_data:
                    st.success("QR Code detectado!")
                    chave_qr = st.session_state.lab_chave
                    if chave_qr:
                        st.info(f"**Chave:** `{chave_qr}`")
                        # Dados da chave — instantâneo, sem OCR nem rede
                        kd = _parse_chave(chave_qr)
                        if kd:
                            cols_k = st.columns(3)
                            cols_k[0].metric("Emissão", kd.get("emissao", "—"))
                            cols_k[1].metric("Nº NF", kd.get("numero_nf", "—"))
                            cols_k[2].metric("Série", kd.get("serie", "—"))
                            st.write(f"**CNPJ Emitente:** {kd.get('cnpj_emitente', '—')}")
                            st.write(f"**UF:** {kd.get('uf', '—')} &nbsp;|&nbsp; **Modelo:** {kd.get('modelo', '—')}")
                    else:
                        st.warning("QR detectado, mas chave não extraída.")
                        st.code(qr_data, language=None)
                elif not _CV2_OK and not _ZXING_OK:
                    st.info("opencv / zxingcpp não instalados — detecção de QR ignorada.")
                else:
                    st.warning("QR Code não detectado. Tente OCR ou insira a chave manualmente.")

                # 2. OCR — opcional, só roda ao clicar
                st.divider()
                rodar_ocr = st.button("🔎 Executar OCR no cupom", key="btn_ocr",
                                      help="Extrai valor total e produtos via RapidOCR (ONNX). Rápido e preciso.")
                if rodar_ocr:
                    with st.spinner("Executando OCR..."):
                        ocr_text = ""
                        try:
                            import numpy as _np
                            _engine = _get_rapid_ocr()
                            _result, _ = _engine(_np.array(img))
                            if _result:
                                ocr_text = "\n".join(item[1] for item in _result if item and len(item) >= 2)
                        except Exception as e:
                            st.error(f"Erro OCR: {e}")

                    st.session_state.lab_ocr_text = ocr_text

                    if ocr_text and not st.session_state.lab_chave:
                        chave_ocr = _chave_from_text(ocr_text)
                        if chave_ocr:
                            st.session_state.lab_chave = chave_ocr
                            st.info(f"**Chave (OCR):** `{chave_ocr}`")

                    ocr_data = _parse_ocr_fields(ocr_text)
                    st.session_state.lab_ocr_data = ocr_data

                    if ocr_data:
                        c_data, c_hora, c_total = st.columns(3)
                        c_data.metric("Data", ocr_data.get("data_emissao", "—"))
                        c_hora.metric("Hora", ocr_data.get("hora", "—"))
                        c_total.metric("Valor Total", ocr_data.get("valor_total", "—"))
                        if ocr_data.get("emitente"):
                            st.write(f"**Emitente:** {ocr_data['emitente']}")
                        if ocr_data.get("itens"):
                            st.markdown("**Produtos identificados:**")
                            st.dataframe(
                                pd.DataFrame(ocr_data["itens"]).rename(columns={
                                    "descricao": "Descrição",
                                    "qtd": "Qtd",
                                    "valor_unitario": "Vl. Unit. (R$)",
                                    "valor": "Total (R$)",
                                }),
                                use_container_width=True,
                                hide_index=True,
                            )
                        else:
                            st.info("Itens não detectados — verifique o texto bruto abaixo.")
                    elif ocr_text:
                        st.warning("OCR executado, mas nenhum campo reconhecido.")

                    if ocr_text:
                        with st.expander("Texto bruto do OCR"):
                            st.text(ocr_text)
                else:
                    # Mostra dados OCR de execuções anteriores
                    ocr_data = st.session_state.lab_ocr_data or {}
                    if ocr_data:
                        c_data, c_hora, c_total = st.columns(3)
                        c_data.metric("Data", ocr_data.get("data_emissao", "—"))
                        c_hora.metric("Hora", ocr_data.get("hora", "—"))
                        c_total.metric("Valor Total", ocr_data.get("valor_total", "—"))
                        if ocr_data.get("itens"):
                            st.dataframe(
                                pd.DataFrame(ocr_data["itens"]).rename(columns={
                                    "descricao": "Descrição", "qtd": "Qtd", "valor": "Valor (R$)",
                                }),
                                use_container_width=True, hide_index=True,
                            )



# ─────────────────────────────────────────────────────────────────────────────
# ABA 2 — Dados & Dossiê PDF
# ─────────────────────────────────────────────────────────────────────────────
with tab2:
    st.subheader("📄 Dados da NFe & Geração de Dossiê")

    ocr_data = st.session_state.lab_ocr_data or {}
    nfe_data = st.session_state.lab_nfe_data or {}

    chave = st.session_state.lab_chave

    # ── Dados da chave (instantâneo, sem rede) ───────────────────────────
    kd2 = _parse_chave(chave) if chave else {}
    if kd2:
        st.markdown("#### Dados da Nota (da chave)")
        ck1, ck2, ck3, ck4 = st.columns(4)
        ck1.metric("Emissão", kd2.get("emissao", "—"))
        ck2.metric("Nº NF", kd2.get("numero_nf", "—"))
        ck3.metric("Série", kd2.get("serie", "—"))
        ck4.metric("Modelo", kd2.get("modelo", "—"))
        st.write(f"**CNPJ Emitente:** {kd2.get('cnpj_emitente', '—')} &nbsp;|&nbsp; **UF:** {kd2.get('uf', '—')}")

    # ── Dados OCR ────────────────────────────────────────────────────────
    if ocr_data:
        st.markdown("#### Dados do OCR")
        _lbl = {"data_emissao": "Data", "hora": "Hora", "valor_total": "Total"}
        cm1, cm2, cm3 = st.columns(3)
        for col, (k, lbl) in zip([cm1, cm2, cm3], _lbl.items()):
            col.metric(lbl, ocr_data.get(k, "—"))
        if ocr_data.get("itens"):
            st.markdown("**Produtos (OCR):**")
            st.dataframe(
                pd.DataFrame(ocr_data["itens"]).rename(columns={
                    "descricao": "Descrição", "qtd": "Qtd", "valor": "Valor (R$)",
                }),
                use_container_width=True, hide_index=True,
            )

    # ── Consulta / PDF SEFAZ ─────────────────────────────────────────────
    st.divider()
    st.markdown("#### Cópia Oficial NFCe")

    if not chave:
        st.warning("Informe a chave de acesso acima para consultar o SEFAZ SP.")
    else:
        url_consulta = _build_consultation_url(chave, st.session_state.lab_qr_url)
        st.markdown(f"🔗 [Abrir no SEFAZ SP]({url_consulta})")

        # Se já temos PDF gerado
        _pdf_stored = st.session_state.get("lab_nfce_xml")  # reutiliza a chave
        _sefaz_data = st.session_state.get("lab_nfe_data") or {}
        if _pdf_stored and isinstance(_pdf_stored, bytes) and _pdf_stored[:4] == b'%PDF':
            st.success("✅ PDF da NFCe gerado — cópia oficial disponível para download.")
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                st.download_button(
                    "📄 Baixar PDF da NFCe",
                    data=_pdf_stored,
                    file_name=f"nfce_{chave[:8]}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                )
            with col_dl2:
                if st.button("🗑️ Limpar / Regerar", use_container_width=True):
                    st.session_state.lab_nfce_xml = None
                    st.session_state.lab_nfe_data = None
                    st.rerun()
            # Preview como imagem no dossiê
            _png_preview = _pdf_first_page_to_png(_pdf_stored)
            if _png_preview:
                st.session_state.lab_nfce_screenshot = _png_preview
                with st.expander("👁️ Preview da cópia NFCe", expanded=False):
                    st.image(_png_preview, use_container_width=True)
            else:
                st.caption("ℹ️ Para visualizar preview instale PyMuPDF: `pip install pymupdf`")
            if _sefaz_data.get("itens") and _PANDAS_OK:
                with st.expander(f"📋 Itens ({len(_sefaz_data['itens'])} produtos)", expanded=False):
                    st.dataframe(
                        pd.DataFrame(_sefaz_data["itens"]).rename(columns={
                            "descricao": "Descrição", "qtd": "Qtd",
                            "unidade": "UN", "valor_unitario": "Vl. Unit.", "valor": "Vl. Total",
                        }),
                        use_container_width=True, hide_index=True,
                    )
        else:
            col_btn1, col_btn2, _sp = st.columns([2, 1.5, 1])
            with col_btn1:
                gerar_pdf = st.button(
                    "🖨️ Gerar PDF via impressão (~30s)",
                    type="primary",
                    use_container_width=True,
                    disabled=not _PLAYWRIGHT_OK,
                    help=(
                        "Simula 'Imprimir > Salvar como PDF' no browser. "
                        "O modo @media print força todo o DANFE visível."
                    ),
                )
            with col_btn2:
                limpar = st.button("🗑️ Limpar dados NFe", use_container_width=True)

            if limpar:
                st.session_state.lab_nfe_data = None
                st.session_state.lab_nfce_xml = None
                st.session_state.lab_nfce_screenshot = None
                st.rerun()

            if gerar_pdf:
                with st.spinner("Abrindo SEFAZ em modo impressão (~30s)..."):
                    _pdf_bytes, _pdf_log = _print_nfce_pdf(
                        url_consulta, timeout_ms=30000
                    )
                if _pdf_bytes and _pdf_bytes[:4] == b'%PDF':
                    st.session_state.lab_nfce_xml = _pdf_bytes
                    # Extrai texto da página para estruturar dados
                    _pt, _ = _fetch_nfce_pagetext(url_consulta, timeout_ms=25000)
                    if _pt:
                        _parsed = _parse_nfce_pagetext(_pt)
                        if _parsed.get("emitente") or _parsed.get("itens"):
                            st.session_state.lab_nfe_data = _parsed
                    st.success(f"✅ PDF gerado! ({len(_pdf_bytes):,} bytes)")
                    st.rerun()
                else:
                    st.error("Não foi possível gerar o PDF.")
                    with st.expander("Detalhes do erro", expanded=True):
                        st.code(_pdf_log or "sem mensagem", language="text")

        # Preview da imagem NFCe no dossiê (se disponível)
        _ss = st.session_state.lab_nfce_screenshot
        if _ss and not (_pdf_stored and isinstance(_pdf_stored, bytes) and _pdf_stored[:4] == b'%PDF'):
            st.markdown("##### Cópia NFCe no dossiê")
            st.image(_ss, use_container_width=True)

    # ── Dossiê ────────────────────────────────────────────────────────────
    st.divider()
    _img_bytes_for_doc = (
        st.session_state.lab_img_bytes
        if "lab_img_bytes" in st.session_state
        else None
    )
    _chave_data_for_doc = _parse_chave(chave) if chave else {}
    # nfe_data já é lab_nfe_data do session_state
    _nfe_for_doc = nfe_data

    docx_bytes = _gerar_dossie_docx(
        chave=chave or "",
        chave_data=_chave_data_for_doc,
        ocr_data=ocr_data,
        nfe_data=_nfe_for_doc,
        img_bytes=_img_bytes_for_doc,
        qr_url=st.session_state.lab_qr_url,
        nfce_screenshot=st.session_state.lab_nfce_screenshot,
    )
    fname_docx = f"dossie_nfce_{chave[:8]}.docx" if chave else "dossie_nfce.docx"
    st.download_button(
        label="⬇️ Baixar Dossiê (Word)",
        data=docx_bytes,
        file_name=fname_docx,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
        )

# ─────────────────────────────────────────────────────────────────────────────
# ABA 3 — NFe Original (SEFAZ iframe)
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# ABA 3 — Processamento em Lote
# ─────────────────────────────────────────────────────────────────────────────
with tab3:
    st.subheader("📦 Processamento em Lote")
    st.caption(
        "Envie até 30 fotos de cupons fiscais. "
        "QR Code e OCR são detectados automaticamente. "
        "O PDF do SEFAZ é gerado em lote com **um único browser** — muito mais rápido."
    )

    # ── Upload múltiplo ───────────────────────────────────────────────────
    uploaded_batch = st.file_uploader(
        "Fotos dos cupons fiscais",
        type=["png", "jpg", "jpeg", "webp"],
        accept_multiple_files=True,
        help="Selecione vários arquivos de uma vez (Ctrl+clique ou Shift+clique).",
        key="batch_uploader",
    )

    # Sincroniza uploads com o state
    if uploaded_batch:
        existing_names = {it["filename"] for it in st.session_state.lab_batch_items}
        added = 0
        for uf in uploaded_batch:
            if uf.name not in existing_names:
                # getvalue() é o método correto para UploadedFile no Streamlit
                raw = bytes(uf.getvalue())
                # Detecção rápida de QR logo na inserção
                try:
                    from PIL import Image as _PILImage
                    import io as _io
                    _img = _PILImage.open(_io.BytesIO(raw))
                    _qr = _detect_qr(_img)
                    _chave_auto = _chave_from_qr(_qr) if _qr else None
                except Exception:
                    _qr = None
                    _chave_auto = None

                st.session_state.lab_batch_items.append({
                    "filename":   uf.name,
                    "img_bytes":  raw,
                    "qr_img_bytes": None,  # foto separada só para leitura do QR
                    "chave":      _chave_auto,
                    "qr_url":     _qr,
                    "ocr_text":   None,
                    "ocr_data":   None,
                    "pdf_bytes":  None,
                    "status":     "qr_ok" if _chave_auto else "aguardando",
                    "pdf_error":  None,
                    "docx_bytes": None,
                })
                added += 1
        if added:
            st.rerun()

    items: list[dict] = st.session_state.lab_batch_items

    if not items:
        st.info("Nenhum cupom carregado ainda. Use o uploader acima.")
    else:
        # ── Barra de ações ────────────────────────────────────────────────
        n_total    = len(items)
        n_ocr      = sum(1 for it in items if it.get("ocr_data"))
        n_chave    = sum(1 for it in items if it.get("chave"))
        n_sem_chave = n_total - n_chave
        n_pdf_ok   = sum(1 for it in items if it.get("pdf_bytes") and it["pdf_bytes"][:4] == b"%PDF")
        n_docx_ok  = sum(1 for it in items if it.get("docx_bytes"))
        col_a, col_b, col_c, col_d = st.columns([2, 2, 2, 1])
        with col_a:
            rodar_ocr_lote = st.button(
                f"🔎 OCR em todos ({n_total - n_ocr} pendentes)",
                use_container_width=True,
                disabled=n_ocr == n_total,
            )
        with col_b:
            gerar_pdf_lote = st.button(
                f"🖨️ Gerar PDFs ({n_chave} com chave)",
                use_container_width=True,
                type="primary",
                disabled=not _PLAYWRIGHT_OK or n_chave == 0,
                help="Um único browser para todos — muito mais rápido.",
            )
        with col_c:
            # Download ZIP com PDFs + Dossiês disponíveis
            import io as _io2
            import zipfile as _zipfile
            _n_zip = n_pdf_ok + n_docx_ok
            if _n_zip > 0:
                _zip_buf = _io2.BytesIO()
                with _zipfile.ZipFile(_zip_buf, "w", _zipfile.ZIP_DEFLATED) as _zf:
                    for _zi, _it in enumerate(items):
                        # Usa índice como prefixo para evitar colisões de nome
                        _slug = f"{_zi + 1:02d}_{(_it.get('chave') or _it['filename'])[:20]}"
                        if _it.get("pdf_bytes") and _it["pdf_bytes"][:4] == b"%PDF":
                            _zf.writestr(f"pdfs/nfce_{_slug}.pdf", _it["pdf_bytes"])
                        if _it.get("docx_bytes"):
                            _zf.writestr(f"dossies/dossie_{_slug}.docx", _it["docx_bytes"])
                st.download_button(
                    f"📥 ZIP ({_n_zip} arquivos)",
                    data=_zip_buf.getvalue(),
                    file_name="nfce_lote.zip",
                    mime="application/zip",
                    use_container_width=True,
                )
            else:
                st.button("📥 Baixar ZIP", disabled=True, use_container_width=True,
                          help="Gere PDFs ou Dossiês primeiro.")
        with col_d:
            if st.button("🗑️ Limpar", use_container_width=True):
                st.session_state.lab_batch_items = []
                st.rerun()

        # ── segunda linha: busca de chave ─────────────────────────────────
        col_ck, _sp_ck = st.columns([3, 3])
        with col_ck:
            buscar_chaves_lote = st.button(
                f"🔑 Buscar chaves ({n_sem_chave} sem chave) — QR → OCR",
                use_container_width=True,
                disabled=n_sem_chave == 0,
                help="Tenta QR Code; nos que falhar, usa OCR para encontrar a chave de 44 dígitos.",
            )

        # ── OCR em lote ───────────────────────────────────────────────────
        if buscar_chaves_lote:
            import numpy as _np_ck
            from PIL import Image as _PILImage_ck
            import io as _io_ck
            _sem_chave_items = [(i, it) for i, it in enumerate(items) if not it.get("chave")]
            _ck_bar = st.progress(0, text="Buscando chaves...")
            _ocr_engine_ck = None
            _ck_found_count = 0
            for _ci, (_orig_idx, _it) in enumerate(_sem_chave_items):
                _label = _it["filename"]
                # 1. Tenta QR code — usa foto dedicada ao QR se disponível
                try:
                    _qr_src = _it.get("qr_img_bytes") or _it["img_bytes"]
                    _img_ck = _PILImage_ck.open(_io_ck.BytesIO(_qr_src))
                    _qr_ck = _detect_qr(_img_ck)
                    if _qr_ck:
                        _chave_ck = _chave_from_qr(_qr_ck)
                        if _chave_ck:
                            _it["chave"] = _chave_ck
                            _it["qr_url"] = _qr_ck
                            _it["status"] = "qr_ok"
                            _ck_found_count += 1
                            _ck_bar.progress(
                                (_ci + 1) / len(_sem_chave_items),
                                text=f"✅ QR {_ci+1}/{len(_sem_chave_items)}: {_label}",
                            )
                            continue
                except Exception:
                    pass
                # 2. OCR → busca chave de 44 dígitos no texto
                try:
                    if _ocr_engine_ck is None:
                        _ocr_engine_ck = _get_rapid_ocr()
                    _img_ck2 = _PILImage_ck.open(_io_ck.BytesIO(_it["img_bytes"]))
                    _res_ck, _ = _ocr_engine_ck(_np_ck.array(_img_ck2))
                    _txt_ck = "\n".join(r[1] for r in (_res_ck or []) if r and len(r) >= 2)
                    if _txt_ck and not _it.get("ocr_text"):
                        _it["ocr_text"] = _txt_ck
                    _chave_ck2 = _chave_from_text(_txt_ck) if _txt_ck else None
                    if _chave_ck2:
                        _it["chave"] = _chave_ck2
                        _it["status"] = "ocr_ok"
                        _ck_found_count += 1
                except Exception:
                    pass
                _ck_bar.progress(
                    (_ci + 1) / len(_sem_chave_items),
                    text=f"OCR {_ci+1}/{len(_sem_chave_items)}: {_label}",
                )
            _ck_bar.empty()
            st.success(
                f"✅ Chaves encontradas: **{_ck_found_count}** de **{len(_sem_chave_items)}** "
                f"({'QR+OCR' if _ocr_engine_ck else 'apenas QR'})"
            )
            st.rerun()

        if rodar_ocr_lote:
            import time as _time_ocr
            _ocr_engine = _get_rapid_ocr()
            import numpy as _np2
            _ocr_t0 = _time_ocr.perf_counter()
            _ocr_bar = st.progress(0, text="Executando OCR...")
            _ocr_timer_ph = st.empty()
            _ocr_tempos: list[float] = []
            for _idx, _it in enumerate(items):
                if _it.get("ocr_data"):
                    continue
                _t_item = _time_ocr.perf_counter()
                try:
                    from PIL import Image as _PILImage2
                    import io as _io3
                    _img2 = _PILImage2.open(_io3.BytesIO(_it["img_bytes"]))
                    _res, _ = _ocr_engine(_np2.array(_img2))
                    _txt = "\n".join(r[1] for r in (_res or []) if r and len(r) >= 2)
                    _od = _parse_ocr_fields(_txt)
                    _it["ocr_text"] = _txt
                    _it["ocr_data"] = _od
                    # Tenta extrair chave do texto se ainda não tem
                    if not _it.get("chave") and _txt:
                        _ck = _chave_from_text(_txt)
                        if _ck:
                            _it["chave"] = _ck
                    _it["status"] = "ocr_ok"
                except Exception as _e:
                    _it["status"] = "ocr_erro"
                _dt_ocr = _time_ocr.perf_counter() - _t_item
                _ocr_tempos.append(_dt_ocr)
                _elapsed_ocr = _time_ocr.perf_counter() - _ocr_t0
                _media_ocr = _elapsed_ocr / len(_ocr_tempos)
                _rest_ocr = _media_ocr * (n_total - _idx - 1)
                _ocr_bar.progress(
                    (_idx + 1) / n_total,
                    text=f"OCR {_idx + 1}/{n_total} | último: {_dt_ocr:.1f}s | restante: ~{_rest_ocr:.0f}s",
                )
                _ocr_timer_ph.caption(
                    f"⏱️ Decorrido: **{_elapsed_ocr:.1f}s** | "
                    f"Média/cupom: **{_media_ocr:.1f}s** | "
                    f"Estimativa restante: **{_rest_ocr:.0f}s**"
                )
            _ocr_total = _time_ocr.perf_counter() - _ocr_t0
            _ocr_bar.empty()
            _ocr_timer_ph.success(
                f"✅ OCR concluído em **{_ocr_total:.1f}s** "
                f"({_ocr_total / max(len(_ocr_tempos), 1):.1f}s/cupom)"
            )
            st.rerun()

        # ── PDF em lote ───────────────────────────────────────────────────
        if gerar_pdf_lote:
            import time as _time
            _itens_com_chave = [(i, it) for i, it in enumerate(items) if it.get("chave")]
            _urls_batch = [
                _build_consultation_url(it["chave"], it.get("qr_url"))
                for _, it in _itens_com_chave
            ]

            _pdf_t0 = _time.perf_counter()
            _pdf_timer_ph = st.empty()
            _pdf_status = st.status(
                f"Gerando {len(_urls_batch)} PDFs (browser único)...", expanded=True
            )
            _pdf_bar = st.progress(0)

            def _upd_progress(cur, tot):
                _el = _time.perf_counter() - _pdf_t0
                _media_pdf = _el / max(cur, 1)
                _rest = _media_pdf * (tot - cur)
                _pdf_bar.progress(
                    cur / tot,
                    text=f"PDF {cur}/{tot} | decorrido: {_el:.0f}s | restante: ~{_rest:.0f}s",
                )
                _pdf_timer_ph.caption(
                    f"⏱️ Decorrido: **{_el:.0f}s** | "
                    f"Média por NFC-e: **{_media_pdf:.1f}s** | "
                    f"Estimativa restante: **{_rest:.0f}s**"
                )

            with _pdf_status:
                _resultados = _print_nfce_pdf_batch(_urls_batch, timeout_ms=30000, progress_cb=_upd_progress)

            _pdf_total = _time.perf_counter() - _pdf_t0
            _pdf_bar.empty()
            _ok_count = 0
            for (_orig_idx, _it), (_pdf, _log) in zip(_itens_com_chave, _resultados):
                if _pdf and _pdf[:4] == b"%PDF":
                    items[_orig_idx]["pdf_bytes"] = _pdf
                    items[_orig_idx]["status"] = "pdf_ok"
                    _ok_count += 1
                else:
                    items[_orig_idx]["status"] = "pdf_erro"
                    items[_orig_idx]["pdf_error"] = _log

            _pdf_status.update(
                label=f"✅ {_ok_count}/{len(_urls_batch)} PDFs em {_pdf_total:.0f}s "
                      f"({_pdf_total / max(len(_urls_batch), 1):.1f}s/NFC-e).",
                state="complete",
            )
            _pdf_timer_ph.success(
                f"✅ PDF concluído em **{_pdf_total:.0f}s** — "
                f"{_pdf_total / max(len(_urls_batch), 1):.1f}s por NFC-e"
            )
            # Gera dossiês automaticamente para quem teve PDF — em paralelo
            _d_count = 0
            _d_errors: list[str] = []
            # Coleta dados na thread principal antes de passar para workers
            _itens_pdf_ok = [
                (orig_idx, dict(st.session_state.lab_batch_items[orig_idx]))
                for orig_idx, _it in _itens_com_chave
                if st.session_state.lab_batch_items[orig_idx].get("pdf_bytes", b"")[:4] == b"%PDF"
            ]

            import concurrent.futures as _cf
            import traceback as _dtb

            def _gerar_um_dossie(args):
                # Recebe cópia dos dados — sem acessar st.session_state na thread
                orig_idx, item_data = args
                try:
                    nfce_ss = _pdf_first_page_to_png(item_data["pdf_bytes"])
                    db = _gerar_dossie_docx(
                        chave=item_data.get("chave", ""),
                        chave_data=_parse_chave(item_data["chave"]) if item_data.get("chave") else {},
                        ocr_data=item_data.get("ocr_data"),
                        nfe_data=None,
                        img_bytes=item_data.get("img_bytes"),
                        qr_url=item_data.get("qr_url"),
                        nfce_screenshot=nfce_ss,
                    )
                    return orig_idx, db, None
                except Exception:
                    return orig_idx, None, _dtb.format_exc()[:200]

            _workers = min(4, len(_itens_pdf_ok))
            with _cf.ThreadPoolExecutor(max_workers=_workers) as _pool:
                for _orig_idx, _docx, _err in _pool.map(_gerar_um_dossie, _itens_pdf_ok):
                    _fname = st.session_state.lab_batch_items[_orig_idx]["filename"]
                    if _docx:
                        st.session_state.lab_batch_items[_orig_idx]["docx_bytes"] = _docx
                        st.session_state.lab_batch_items[_orig_idx]["status"] = "dossie_ok"
                        _d_count += 1
                    elif _err:
                        _d_errors.append(f"{_fname}: {_err}")
            if _d_count:
                _total_geral = _time.perf_counter() - _pdf_t0
                st.toast(f"📋 {_d_count} dossiês gerados em {_total_geral:.0f}s total.")
            if _d_errors:
                st.warning("Erros na geração de dossiês:\n" + "\n".join(_d_errors))
            st.rerun()

        # ── Tabela de resultados ──────────────────────────────────────────
        st.divider()
        _STATUS_EMOJI = {
            "aguardando": "⏳",
            "qr_ok":      "📷✅",
            "ocr_ok":     "🔎✅",
            "ocr_erro":   "🔎❌",
            "pdf_ok":     "📄✅",
            "pdf_erro":   "📄❌",
            "dossie_ok":  "📋✅",
        }

        for _idx, _it in enumerate(items):
            with st.container(border=True):
                _c1, _c2, _c3, _c4, _c5 = st.columns([0.5, 2.5, 3, 1.5, 1.5])
                _c1.markdown(f"**{_idx + 1}**")
                _c2.markdown(f"📁 `{_it['filename']}`  \n{_STATUS_EMOJI.get(_it['status'], '?')} **{_it['status']}**")
                _chave_disp = _it.get("chave") or ""
                if _chave_disp:
                    _c3.code(_chave_disp, language=None)
                else:
                    _c3.markdown("🔑 `—`")
                if _it.get("qr_url"):
                    _c3.code(_it["qr_url"], language=None)
                # Aviso explícito quando não tem chave e não foi tentado
                if not _it.get("chave") and _it.get("status") in ("aguardando", "ocr_erro", None):
                    _c3.caption("⚠️ Sem chave — use **Buscar chaves** ou preencha abaixo")

                with _c4:
                    if _it.get("pdf_bytes") and _it["pdf_bytes"][:4] == b"%PDF":
                        st.download_button(
                            "📄 PDF",
                            data=_it["pdf_bytes"],
                            file_name=f"nfce_{(_it.get('chave') or _it['filename'])[:12]}.pdf",
                            mime="application/pdf",
                            key=f"dl_pdf_{_idx}",
                            use_container_width=True,
                        )
                    elif _it.get("status") == "pdf_erro":
                        with st.expander("❌ PDF"):
                            st.code(_it.get("pdf_error", "")[:300], language="text")

                with _c5:
                    if _it.get("docx_bytes"):
                        _slug2 = (_it.get("chave") or _it["filename"])[:12]
                        st.download_button(
                            "📋 Dossiê",
                            data=_it["docx_bytes"],
                            file_name=f"dossie_{_slug2}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_docx_{_idx}",
                            use_container_width=True,
                        )

                # ── Entrada manual de chave/QR para itens sem chave ──────
                if not _it.get("chave"):
                    with st.expander("✏️ Informar chave / foto do QR"):
                        # ── Foto dedicada ao QR ───────────────────────────
                        st.caption("**Foto do QR Code** (zoom no QR — usada só para leitura da chave)")
                        _qr_upload = st.file_uploader(
                            "Foto com zoom no QR Code",
                            type=["png", "jpg", "jpeg", "webp"],
                            key=f"qr_upload_{_idx}",
                            label_visibility="collapsed",
                        )
                        if _qr_upload is not None:
                            _qr_raw = bytes(_qr_upload.getvalue())
                            _it["qr_img_bytes"] = _qr_raw
                            # Tenta ler QR imediatamente
                            try:
                                from PIL import Image as _PILImg2
                                import io as _io2
                                _qr_res = _detect_qr(_PILImg2.open(_io2.BytesIO(_qr_raw)))
                                if _qr_res:
                                    _chave_res = _chave_from_qr(_qr_res)
                                    if _chave_res:
                                        _it["chave"] = _chave_res
                                        _it["qr_url"] = _qr_res
                                        _it["status"] = "qr_ok"
                                        st.success(f"✅ Chave encontrada: `{_chave_res}`")
                                        st.rerun()
                                    else:
                                        st.warning("QR lido mas chave não extraída — tente 'Buscar chaves'.")
                                else:
                                    st.warning("QR não detectado nesta foto — tente 'Buscar chaves' ou informe manualmente.")
                            except Exception:
                                st.warning("Erro ao ler a foto do QR.")

                        st.divider()
                        # ── Entrada manual ────────────────────────────────
                        st.caption("**Ou informe manualmente:**")
                        _fm1, _fm2 = st.columns([3, 2])
                        with _fm1:
                            _chave_manual = st.text_input(
                                "Chave de acesso (44 dígitos)",
                                key=f"chave_manual_{_idx}",
                                placeholder="Cole aqui os 44 dígitos da chave",
                                max_chars=44,
                            )
                        with _fm2:
                            _qr_manual = st.text_input(
                                "URL do QR Code (opcional)",
                                key=f"qr_manual_{_idx}",
                                placeholder="Cole a URL do QR Code se tiver",
                            )
                        if st.button("💾 Salvar", key=f"salvar_chave_{_idx}"):
                            _digits_only = re.sub(r'\D', '', _chave_manual)
                            if len(_digits_only) == 44:
                                _it["chave"] = _digits_only
                                _it["qr_url"] = _qr_manual.strip() or None
                                _it["status"] = "ocr_ok"
                                st.rerun()
                            else:
                                st.error(f"Chave inválida — encontrados {len(_digits_only)} dígitos (esperado 44).")

                # ── Link SEFAZ quando tem chave mas ainda não tem PDF ────
                if _it.get("chave") and not (_it.get("pdf_bytes") and _it["pdf_bytes"][:4] == b"%PDF"):
                    _url_sefaz = _build_consultation_url(_it["chave"], _it.get("qr_url"))
                    st.markdown(f"🔗 [Consultar no SEFAZ SP]({_url_sefaz})", unsafe_allow_html=False)

                # Mostra dados OCR se disponível
                if _it.get("ocr_data"):
                    _od2 = _it["ocr_data"]
                    _mc1, _mc2, _mc3, _mc4 = st.columns(4)
                    _mc1.caption(f"Emitente: {_od2.get('emitente', '—')[:30]}")
                    _mc2.caption(f"Data: {_od2.get('data_emissao', '—')}")
                    _mc3.caption(f"Total: R$ {_od2.get('valor_total', '—')}")
                    _mc4.caption(f"Itens: {len(_od2.get('itens', []))}")
                    # Miniatura do cupom
                    with st.expander("🖼️ Ver foto do cupom"):
                        st.image(_it["img_bytes"], use_container_width=True)
